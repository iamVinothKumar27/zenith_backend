# server.py
from flask import Flask, jsonify, request, send_from_directory, send_file
from flask_cors import CORS
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from supadata import Supadata

from PyPDF2 import PdfReader

import requests
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai

import os, json, re, io, tempfile
from html import escape as _html_escape
import time
import uuid
import math

import hashlib
from collections import deque
from datetime import datetime, timezone, timedelta

# ---- datetime helpers (avoid naive vs aware issues) ----
def _to_utc_aware(dt):
    """Ensure a datetime is timezone-aware in UTC (Mongo may return naive UTC)."""
    if not dt or not isinstance(dt, datetime):
        return dt
    if dt.tzinfo is None:
        return dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


from werkzeug.utils import secure_filename

from flask import g
from pymongo import MongoClient
import gridfs
from bson.objectid import ObjectId
import firebase_admin
from firebase_admin import credentials as fb_credentials
from firebase_admin import auth as fb_auth
load_dotenv()  

# ----------------- MOCKTEST DEDUPE CACHES -----------------
# Keep a short rolling window to reduce repeated questions across sessions.
RECENT_CODING_TITLES = deque(maxlen=200)
RECENT_MCQ_HASHES = deque(maxlen=400)


# ----------------- EMAIL (SMTP: Titan/GoDaddy etc.) -----------------
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import re
# ----------------- PASSWORD REUSE PREVENTION -----------------
PASSWORD_PEPPER = os.getenv("PASSWORD_PEPPER", "").strip()  # set this in .env (random long string)

def _pw_hash(uid: str, password: str) -> str:
    """
    Hash password so we can detect reuse.
    Uses SHA-256 with uid + server-side pepper.
    NOTE: This is NOT used for authentication; only reuse detection.
    """
    uid = (uid or "").strip()
    password = (password or "").strip()
    raw = f"{uid}|{PASSWORD_PEPPER}|{password}".encode("utf-8")
    return hashlib.sha256(raw).hexdigest()

def _yt_id(url: str) -> str:
    """Extract YouTube ID from watch?v= or youtu.be links."""
    if not url:
        return ""
    m = re.search(r"(?:v=|youtu\.be/)([A-Za-z0-9_-]{6,})", url)
    return m.group(1) if m else ""


def derive_video_number_from_course(db, uid: str, course_title: str, video_url: str):
    """
    Returns (video_no, total_videos) like (4, 32) by searching course_states.videos
    Supports shapes:
      1) [{"Week 1": [{"topic":..., "video": URL}, ...]}, {"Week 2": [...]}]   ✅ your current build_weekly_json
      2) [{"week":"Week 1","videos":[{"videoUrl":...}, ...]}]
      3) flat list of urls or dicts
    """
    if not uid or not course_title or not video_url:
        return (None, None)

    target_id = _yt_id(video_url) or video_url.strip()

    state = db.course_states.find_one(
        {"uid": uid, "courseTitle": course_title},
        {"videos": 1, "weeks": 1, "course": 1}
    )
    if not state:
        return (None, None)

    def _get_url(obj):
        if obj is None:
            return ""
        if isinstance(obj, str):
            return obj.strip()
        if isinstance(obj, dict):
            return (obj.get("videoUrl")
                    or obj.get("video_url")
                    or obj.get("url")
                    or obj.get("video")
                    or "").strip()
        return str(obj).strip()

    def _flatten_video_urls(node):
        out = []
        if node is None:
            return out

        # string url
        if isinstance(node, str):
            u = node.strip()
            if u:
                out.append(u)
            return out

        # list -> recurse each
        if isinstance(node, list):
            for it in node:
                out.extend(_flatten_video_urls(it))
            return out

        # dict -> handle multiple schema
        if isinstance(node, dict):
            # A) {"videos": [ ... ]}
            if isinstance(node.get("videos"), list):
                for v in node.get("videos"):
                    u = _get_url(v)
                    if u:
                        out.append(u)
                    else:
                        out.extend(_flatten_video_urls(v))

            # B) build_weekly_json: {"Week 1": [ {topic, video}, ... ]}
            for k, v in node.items():
                if isinstance(v, list):
                    for item in v:
                        u = _get_url(item)  # reads "video" key too
                        if u:
                            out.append(u)
                        else:
                            out.extend(_flatten_video_urls(item))
                else:
                    out.extend(_flatten_video_urls(v))
            return out

        # fallback: unknown type
        return out

    # ✅ pull weeks/videos from all possible places
    videos_root = state.get("videos")
    if not videos_root and isinstance(state.get("course"), dict):
        videos_root = state["course"].get("videos") or state["course"].get("weeks")
    if not videos_root:
        videos_root = state.get("weeks")

    flattened = _flatten_video_urls(videos_root)
    if not flattened:
        return (None, None)

    total = len(flattened)

    # match by youtube id (preferred) else full url
    for idx, vurl in enumerate(flattened, start=1):
        vid = _yt_id(vurl) or vurl.strip()
        if vid and target_id and vid == target_id:
            return (idx, total)

    return (None, total)

def derive_video_meta_from_course(db, uid: str, course_title: str, video_url: str):
    """Like derive_video_number_from_course, but also returns the video's title (if available).

    Returns (video_no, total_videos, video_title)
    """
    if not uid or not course_title or not video_url:
        return (None, None, None)

    target_id = _yt_id(video_url) or video_url.strip()

    state = db.course_states.find_one(
        {"uid": uid, "courseTitle": course_title},
        {"videos": 1, "weeks": 1, "course": 1}
    )
    if not state:
        return (None, None, None)

    videos_root = state.get("videos")
    if not videos_root and isinstance(state.get("course"), dict):
        videos_root = state["course"].get("videos") or state["course"].get("weeks")
    if not videos_root:
        videos_root = state.get("weeks")

    entries = flatten_course_videos_with_titles(videos_root)
    if not entries:
        return (None, None, None)

    total = len(entries)
    for idx, ent in enumerate(entries, start=1):
        vurl = (ent.get("url") or "").strip()
        vid = _yt_id(vurl) or vurl
        if vid and target_id and vid == target_id:
            return (idx, total, (ent.get("title") or f"Video {idx}"))

    return (None, total, None)


def flatten_course_videos_with_titles(videos_root):
    """Flatten course videos into an ordered list of {url, title}.

    Supports multiple schemas:
      - build_weekly_json: [{"Week 1":[{"topic":..., "video":...}, ...]}, ...]
      - {"week":"Week 1","videos":[{"title":..., "videoUrl":...}, ...]}
      - flat list of strings / dicts
    """
    def _get_url(obj):
        if obj is None:
            return ""
        if isinstance(obj, str):
            return obj.strip()
        if isinstance(obj, dict):
            return (obj.get("videoUrl")
                    or obj.get("video_url")
                    or obj.get("url")
                    or obj.get("video")
                    or "").strip()
        return str(obj).strip()

    def _get_title(obj):
        if obj is None:
            return ""
        if isinstance(obj, dict):
            return (obj.get("title")
                    or obj.get("topic")
                    or obj.get("name")
                    or obj.get("videoTitle")
                    or obj.get("heading")
                    or "").strip()
        return ""

    out = []

    def _push(url, title):
        u = (url or "").strip()
        if not u:
            return
        t = (title or "").strip()
        out.append({"url": u, "title": t})

    def _walk(node, inherited_title=""):
        if node is None:
            return
        if isinstance(node, str):
            _push(node, inherited_title)
            return
        if isinstance(node, list):
            for it in node:
                _walk(it, inherited_title)
            return
        if isinstance(node, dict):
            # If this dict itself is a video entry
            u = _get_url(node)
            t = _get_title(node) or inherited_title
            if u:
                _push(u, t)
                return

            # Schema: {"videos":[...]}
            if isinstance(node.get("videos"), list):
                for v in node.get("videos"):
                    _walk(v, inherited_title=_get_title(v) or inherited_title)
                return

            # build_weekly_json: {"Week 1":[...]}
            for k, v in node.items():
                if isinstance(v, list):
                    for item in v:
                        _walk(item, inherited_title=_get_title(item) or inherited_title)
                else:
                    _walk(v, inherited_title=inherited_title)
            return

    _walk(videos_root, "")
    # Fill missing titles with Video N
    for i, ent in enumerate(out, start=1):
        if not ent.get("title"):
            ent["title"] = f"Video {i}"
    return out

SMTP_HOST = os.getenv("SMTP_HOST", "").strip()
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "").strip()
SMTP_PASS = os.getenv("SMTP_PASS", "").strip()

# ✅ Sender aliases (Google Workspace / Gmail "Send mail as")
# These should be configured as aliases for your primary mailbox (SMTP_USER).
MAIL_FROM_DEFAULT = (os.getenv("MAIL_FROM_DEFAULT") or os.getenv("MAIL_FROM") or os.getenv("SMTP_FROM") or SMTP_USER).strip()
MAIL_FROM_AUTH = (os.getenv("MAIL_FROM_AUTH") or "authentication@zenithlearning.site").strip()
MAIL_FROM_COURSES = (os.getenv("MAIL_FROM_COURSES") or "courses@zenithlearning.site").strip()
MAIL_FROM_PROFILE = (os.getenv("MAIL_FROM_PROFILE") or "profile@zenithlearning.site").strip()
MAIL_FROM_ADMIN = (os.getenv("MAIL_FROM_ADMIN") or "admin@zenithlearning.site").strip()
MAIL_FROM_CONTACT = (os.getenv("MAIL_FROM_CONTACT") or "contact@zenithlearning.site").strip()
MAIL_FROM_MOCKTEST = (os.getenv("MAIL_FROM_MOCKTEST") or os.getenv("MAIL_FROM_TESTS") or "mock-tests@zenithlearning.site").strip()
MAIL_FROM_ATS = (os.getenv("MAIL_FROM_ATS") or os.getenv("MAIL_FROM_ATS_INTELLIGENCE") or "ats-intelligence@zenithlearning.site").strip()

# Optional: use SendGrid API instead of SMTP (kept for safety, but Mailgun removed)
SENDGRID_API_KEY = os.getenv("SENDGRID_API_KEY", "").strip()
SENDGRID_FROM = os.getenv("SENDGRID_FROM", "").strip()

FRONTEND_BASE_URL = os.getenv("FRONTEND_BASE_URL", os.getenv("FRONTEND_URL", "")).strip().rstrip("/")
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "admin@zenithlearning.site").strip()
ADMIN_EMAILS = os.getenv("ADMIN_EMAILS", "").strip()


def _admin_email_set():
    """Return a lowercase set of admin emails.

    Supports ADMIN_EMAILS as comma-separated list. Falls back to ADMIN_EMAIL.
    """
    s=set()
    if ADMIN_EMAIL: s.add(ADMIN_EMAIL.strip().lower())
    if ADMIN_EMAILS:
        for e in ADMIN_EMAILS.split(','):
            e=e.strip().lower()
            if e: s.add(e)
    return s

CONTACT_INBOX = os.getenv("CONTACT_INBOX", "contact@zenithlearning.site").strip()


def _smtp_ready() -> bool:
    return bool(SMTP_HOST and SMTP_PORT and SMTP_USER and SMTP_PASS)


def _mailgun_ready() -> bool:
    api_key = (os.getenv("MAILGUN_API_KEY") or "").strip()
    domain = (os.getenv("MAILGUN_DOMAIN") or "").strip()
    return bool(api_key and domain)

def _pick_from(kind: str = "") -> str:
    k = (kind or "").strip().lower()
    if k in ("auth", "authentication", "login"):
        return MAIL_FROM_AUTH or MAIL_FROM_DEFAULT
    if k in ("courses", "course", "quiz"):
        return MAIL_FROM_COURSES or MAIL_FROM_DEFAULT
    if k in ("profile",):
        return MAIL_FROM_PROFILE or MAIL_FROM_DEFAULT
    if k in ("admin",):
        return MAIL_FROM_ADMIN or MAIL_FROM_DEFAULT
    if k in ("contact",):
        return MAIL_FROM_CONTACT or MAIL_FROM_DEFAULT
    if k in ("mocktest","mock-tests","tests","mock_test"):
        return MAIL_FROM_MOCKTEST or MAIL_FROM_DEFAULT
    if k in ("ats","ats-intelligence","ats_intelligence"):
        return MAIL_FROM_ATS or MAIL_FROM_DEFAULT
    return MAIL_FROM_DEFAULT or SMTP_USER


def _is_local_dev() -> bool:
    """Best-effort detection: treat localhost/dev as local."""
    try:
        from flask import has_request_context, request
        if has_request_context():
            host = (request.host or "").lower()
            if host.startswith("127.0.0.1") or host.startswith("localhost"):
                return True
    except Exception:
        pass

    env = (os.getenv("APP_ENV") or os.getenv("FLASK_ENV") or "").lower()
    if env in ("dev", "development", "local"):
        return True

    web = (os.getenv("FRONTEND_BASE_URL") or os.getenv("PUBLIC_BASE_URL") or "").lower()
    if "localhost" in web or "127.0.0.1" in web:
        return True

    # Render/Vercel/etc set env vars; if present, assume hosted.
    if os.getenv("RENDER") or os.getenv("VERCEL") or os.getenv("RAILWAY_STATIC_URL") or os.getenv("FLY_APP_NAME"):
        return False

    return False

def _should_use_mailgun() -> bool:
    return _preferred_mail_provider() == "mailgun"


def _send_email_sync_smtp(
    to_email: str,
    subject: str,
    html_body: str,
    text_body: str = "",
    *,
    from_email: str = None,
    reply_to: str = None,
) -> bool:
    """Send email via SMTP. Returns True if sent, False otherwise.

    Notes:
    - When using Google Workspace SMTP, authenticate with SMTP_USER (primary mailbox),
      and set From to one of its verified aliases (courses@, profile@, etc.).
    - For contact form, prefer setting Reply-To to the user's email.
    """
    if not to_email:
        return False
    if not _smtp_ready():
        print("[MAIL] SMTP not configured. Skipping send to:", to_email, "subject:", subject)
        return False

    sender = (from_email or MAIL_FROM_DEFAULT or SMTP_USER).strip()

    msg = MIMEMultipart("alternative")
    msg["From"] = f"Zenith Learning <{sender}>"
    msg["To"] = to_email
    msg["Subject"] = subject
    if reply_to:
        msg["Reply-To"] = reply_to

    if text_body:
        msg.attach(MIMEText(text_body, "plain", "utf-8"))
    msg.attach(MIMEText(html_body, "html", "utf-8"))

    # Keep timeouts low to avoid Gunicorn worker timeouts on hosted deployments.
    connect_timeout = int(os.getenv("SMTP_TIMEOUT", "8"))

    server = None
    try:
        if int(SMTP_PORT) == 465:
            server = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=connect_timeout)
            server.ehlo()
        else:
            server = smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=connect_timeout)
            server.ehlo()
            try:
                server.starttls()
                server.ehlo()
            except Exception:
                pass

        server.login(SMTP_USER, SMTP_PASS)
        server.sendmail(sender, [to_email], msg.as_string())
        server.quit()
        print("[MAIL] SMTP Sent:", subject, "->", to_email, "| from:", sender)
        return True
    except Exception as e:
        print("[MAIL] SMTP Send failed:", e)
        try:
            if server:
                server.quit()
        except Exception:
            pass
        return False

def _send_email_sync_sendgrid(to_email: str, subject: str, html_body: str, text_body: str = "") -> bool:
    """Send email via SendGrid Web API. Returns True if accepted by SendGrid."""
    api_key = os.getenv("SENDGRID_API_KEY", "").strip()
    if not api_key:
        return False

    from_email = os.getenv("SENDGRID_FROM", "").strip() or MAIL_FROM
    if not from_email:
        return False

    try:
        import requests  # type: ignore
    except Exception:
        requests = None

    payload = {
        "personalizations": [{
            "to": [{"email": to_email}],
            "subject": subject,
        }],
        "from": {"email": from_email, "name": "Zenith Learning"},
        "content": [
            {"type": "text/plain", "value": text_body or " "},
            {"type": "text/html", "value": html_body},
        ],
    }

    try:
        if requests:
            r = requests.post(
                "https://api.sendgrid.com/v3/mail/send",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json=payload,
                timeout=10,
            )
            if 200 <= r.status_code < 300:
                print("[MAIL] SendGrid accepted:", subject, "->", to_email)
                return True
            print("[MAIL] SendGrid failed:", r.status_code, getattr(r, "text", ""))
            return False
        else:
            # Minimal fallback without requests
            import urllib.request
            import urllib.error

            req = urllib.request.Request(
                "https://api.sendgrid.com/v3/mail/send",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                ok = 200 <= resp.status < 300
                print("[MAIL] SendGrid accepted:", ok, subject, "->", to_email)
                return ok
    except Exception as e:
        print("[MAIL] SendGrid send failed:", e)
        return False


def _send_email_sync_mailgun(
    to_email: str,
    subject: str,
    html_body: str,
    text_body: str = "",
    *,
    from_email: str = None,
    reply_to: str = None,
) -> bool:
    """Send email via Mailgun HTTP API. Returns True if Mailgun accepted the message."""
    api_key = os.getenv("MAILGUN_API_KEY", "").strip()
    domain = os.getenv("MAILGUN_DOMAIN", "").strip()
    if not api_key or not domain:
        return False

    # API base: US (default) or EU (set MAILGUN_BASE_URL=https://api.eu.mailgun.net)
    # Some deployments set MAILGUN_BASE_URL with /v3 included (ex: https://api.mailgun.net/v3).
    base_url = (os.getenv("MAILGUN_BASE_URL", "https://api.mailgun.net").strip().rstrip("/"))
    if base_url.endswith("/v3"):
        endpoint = f"{base_url}/{domain}/messages"
    else:
        endpoint = f"{base_url}/v3/{domain}/messages"

    from_email = ((from_email or "").strip() or (os.getenv("MAILGUN_FROM", "").strip() or os.getenv("MAIL_FROM", "").strip() or os.getenv("SMTP_FROM", "").strip() or SMTP_USER).strip())
    if not from_email:
        return False

    data = {
        "from": f"Zenith Learning <{from_email}>",
        "to": [to_email],
        "subject": subject,
        "text": text_body or " ",
        "html": html_body,
    }
    if reply_to:
        data["h:Reply-To"] = reply_to

    try:
        import requests  # type: ignore
    except Exception:
        requests = None

    try:
        if requests:
            r = requests.post(
                endpoint,
                auth=("api", api_key),
                data=data,
                timeout=10,
            )
            if 200 <= r.status_code < 300:
                print("[MAIL] Mailgun accepted:", subject, "->", to_email)
                return True
            print("[MAIL] Mailgun failed:", r.status_code, getattr(r, "text", ""))
            return False
        else:
            # Minimal fallback without requests
            import urllib.request
            import urllib.parse
            import base64

            encoded = urllib.parse.urlencode({k: (v[0] if isinstance(v, list) else v) for k, v in data.items()}).encode("utf-8")
            auth_header = base64.b64encode(f"api:{api_key}".encode("utf-8")).decode("ascii")
            req = urllib.request.Request(
                endpoint,
                data=encoded,
                headers={"Authorization": f"Basic {auth_header}", "Content-Type": "application/x-www-form-urlencoded"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                code = getattr(resp, "status", 200)
                if 200 <= code < 300:
                    print("[MAIL] Mailgun accepted:", subject, "->", to_email)
                    return True
                print("[MAIL] Mailgun failed:", code)
                return False
    except Exception as e:
        print("[MAIL] Mailgun send failed:", e)
        return False



def _is_render_env() -> bool:
    """Detect Render hosted environment."""
    return bool(
        os.getenv("RENDER")  # sometimes "true"
        or os.getenv("RENDER_SERVICE_ID")
        or os.getenv("RENDER_EXTERNAL_URL")
        or os.getenv("RENDER_INSTANCE_ID")
    )

def _preferred_mail_provider() -> str:
    """Choose provider: local -> SMTP, hosted -> Mailgun if configured, else SMTP.

    Override with MAIL_PROVIDER=smtp|mailgun|sendgrid or USE_MAILGUN=1.
    """
    override = (os.getenv("MAIL_PROVIDER") or "").strip().lower()
    if override in ("smtp", "mailgun", "sendgrid"):
        return override
    if (os.getenv("USE_MAILGUN") or "").strip().lower() in ("1", "true", "yes"):
        return "mailgun" if _mailgun_ready() else "smtp"
    if _is_local_dev():
        return "smtp"
    if _mailgun_ready():
        return "mailgun"
    if os.getenv("SENDGRID_API_KEY", "").strip():
        return "sendgrid"
    return "smtp"

def _send_email(
    to_email: str,
    subject: str,
    html_body: str,
    text_body: str = "",
    *,
    kind: str = "",
    from_email: str = None,
    reply_to: str = None,
):
    """Send email on a background thread (SMTP only)."""
    if not to_email:
        return

    from_email = (from_email or _pick_from(kind))

    def _job():
        provider = _preferred_mail_provider()

        if provider == "sendgrid":
            if _send_email_sync_sendgrid(to_email, subject, html_body, text_body):
                return
            provider = "mailgun" if _mailgun_ready() else "smtp"

        if provider == "mailgun":
            if _send_email_sync_mailgun(to_email, subject, html_body, text_body, from_email=from_email, reply_to=reply_to):
                return
            provider = "smtp"

        _send_email_sync_smtp(to_email, subject, html_body, text_body, from_email=from_email, reply_to=reply_to)

    try:
        import threading
        t = threading.Thread(target=_job, daemon=True)
        t.start()
    except Exception as e:
        print("[MAIL] Could not spawn mail thread:", e)
        _job()

    return True

def _brand_email(
    title: str,
    preheader: str = "",
    body_html: str = "",
    primary_cta: dict = None,
    secondary_cta: dict = None,
    # Backward/forward compatible args used in some parts of the app
    subtitle: str = "",
    cta_url: str = "",
    cta_text: str = "",
    cta2_url: str = "",
    cta2_text: str = "",
    **_ignore,
):
    """Return a professional, branded email HTML.

    Supports both styles:
    - primary_cta / secondary_cta dicts
    - subtitle + cta_url/cta_text (legacy)
    """
    # Map legacy CTA args into the newer dict form
    try:
        if (not primary_cta) and cta_url and cta_text:
            u = (cta_url or "").strip()
            if u.startswith("/"):
                u = _safe_public_url(u)
            primary_cta = {"url": u, "label": cta_text}
        if (not secondary_cta) and cta2_url and cta2_text:
            u2 = (cta2_url or "").strip()
            if u2.startswith("/"):
                u2 = _safe_public_url(u2)
            secondary_cta = {"url": u2, "label": cta2_text}
    except Exception:
        pass

    # Subtitle is what we want to SHOW below the title. Preheader is still used as the hidden inbox preview.
    visible_subtitle = (subtitle or preheader or "").strip()
    hidden_preheader = (preheader or subtitle or "").strip()
    logo_text = "Zenith Learning"
    year = datetime.now().year
    primary_btn = ""
    if primary_cta and primary_cta.get("url") and primary_cta.get("label"):
        primary_btn = f"""
        <a href="{primary_cta.get('url')}" style="display:inline-block;background:#2563eb;color:#ffffff;text-decoration:none;padding:12px 18px;border-radius:12px;font-weight:700;font-size:14px;">
          {primary_cta.get('label')}
        </a>"""
    secondary_btn = ""
    if secondary_cta and secondary_cta.get("url") and secondary_cta.get("label"):
        secondary_btn = f"""
        <a href="{secondary_cta.get('url')}" style="display:inline-block;background:#111827;color:#ffffff;text-decoration:none;padding:12px 18px;border-radius:12px;font-weight:700;font-size:14px;">
          {secondary_cta.get('label')}
        </a>"""

    cta_row = ""
    if primary_btn or secondary_btn:
        cta_row = f"<tr><td style='padding:12px 24px 22px 24px;'>{primary_btn} &nbsp; {secondary_btn}</td></tr>"

    primary_url = (primary_cta.get('url') if primary_cta else "") or ""

    html = f"""<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>{title}</title>
  </head>
  <body style="margin:0;padding:0;background:#f6f7fb;font-family:Inter,system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;color:#111827;">
    <div style="display:none;max-height:0;overflow:hidden;opacity:0;color:transparent;">{hidden_preheader}</div>

    <table width="100%" cellspacing="0" cellpadding="0" style="background:#f6f7fb;padding:28px 10px;">
      <tr>
        <td align="center">
          <table width="640" cellspacing="0" cellpadding="0" style="background:#ffffff;border-radius:18px;overflow:hidden;box-shadow:0 10px 30px rgba(17,24,39,0.08);">
            <tr>
              <td style="background:linear-gradient(135deg,#0ea5e9,#2563eb);padding:22px 24px;">
                <div style="font-size:16px;font-weight:800;color:#ffffff;letter-spacing:0.2px;">{logo_text}</div>
                <div style="font-size:12px;color:rgba(255,255,255,0.9);margin-top:4px;">Smart learning • Roadmaps • Quizzes • Progress</div>
              </td>
            </tr>

            <tr>
              <td style="padding:26px 24px 6px 24px;">
                <div style="font-size:20px;font-weight:800;line-height:1.25;">{title}</div>
                <div style="font-size:13px;color:#6b7280;margin-top:8px;">{visible_subtitle}</div>
              </td>
            </tr>

            <tr>
              <td style="padding:10px 24px 6px 24px;font-size:14px;line-height:1.65;color:#111827;">
                {body_html}
              </td>
            </tr>

            {cta_row}

            <tr>
              <td style="padding:0 24px 22px 24px;">
                <div style="border-top:1px solid #e5e7eb;padding-top:14px;font-size:12px;color:#6b7280;line-height:1.6;">
                  Need help? Reply to this email or contact our team from the <b>Contact</b> page in the app.<br/>
                  <span style="color:#9ca3af;">© {year} Zenith Learning. All rights reserved.</span>
                </div>
              </td>
            </tr>
          </table>

          <div style="width:640px;font-size:12px;color:#9ca3af;line-height:1.5;margin-top:10px;text-align:center;">
            If the button doesn't work, copy and paste this link in your browser:<br/>
            <span style="word-break:break-all;">{primary_url}</span>
          </div>
        </td>
      </tr>
    </table>
  </body>
</html>"""
    return html

def _safe_public_url(path: str) -> str:
    """
    Always return an ABSOLUTE URL for emails.
    Priority:
    1) FRONTEND_BASE_URL env (recommended)
    2) FRONTEND_URL env
    3) Fallback to request Origin (if available)
    4) Last resort: return path (will not work in email clients)
    """
    base = (os.getenv("FRONTEND_BASE_URL") or os.getenv("FRONTEND_URL") or "").strip().rstrip("/")

    # best-effort fallback (sometimes request has Origin)
    if not base:
        try:
            origin = (request.headers.get("Origin") or "").strip().rstrip("/")
            if origin:
                base = origin
        except Exception:
            pass

    if not path.startswith("/"):
        path = "/" + path

    return (base + path) if base else path

IST = timezone(timedelta(hours=5, minutes=30))

def _now_ist_str():
    # Example: 2026-02-06 01:10 AM IST
    return datetime.now(IST).strftime("%Y-%m-%d %I:%M %p IST")



# ----------------- MONGODB + FIREBASE AUTH -----------------
MONGODB_URI = os.getenv("MONGODB_URI", "")
MONGODB_DB = os.getenv("MONGODB_DB", "zenith")
_mongo_client = None

def get_db():
    global _mongo_client
    if not MONGODB_URI:
        raise RuntimeError("MONGODB_URI is not set")
    if _mongo_client is None:
        _mongo_client = MongoClient(MONGODB_URI, serverSelectionTimeoutMS=5000)
    return _mongo_client[MONGODB_DB]

def _init_firebase_admin():
    if firebase_admin._apps:
        return
    # Option A: Path to service account json
    sa_path = os.getenv("FIREBASE_SERVICE_ACCOUNT_PATH", "").strip()
    sa_json = os.getenv("FIREBASE_SERVICE_ACCOUNT_JSON", "").strip()
    if sa_path:
        cred = fb_credentials.Certificate(sa_path)
        firebase_admin.initialize_app(cred)
        return
    if sa_json:
        cred = fb_credentials.Certificate(json.loads(sa_json))
        firebase_admin.initialize_app(cred)
        return
    raise RuntimeError("Firebase Admin is not configured. Set FIREBASE_SERVICE_ACCOUNT_PATH or FIREBASE_SERVICE_ACCOUNT_JSON")

def require_user():
    """Verify Firebase ID token from Authorization: Bearer <token>."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return None, ("Missing Authorization Bearer token", 401)
    token = auth_header.split(" ", 1)[1].strip()
    try:
        _init_firebase_admin()
        decoded = fb_auth.verify_id_token(token)
        g.user = decoded
        return decoded, None
    except Exception as e:
        return None, (f"Invalid token: {e}", 401)


def require_admin():
    """Verify token + allow only configured admin mailboxes.

    - ADMIN_EMAIL: single mailbox (legacy)
    - ADMIN_EMAILS: comma-separated list (recommended)

    NOTE: Frontend can also gate access, but backend remains the source of truth.
    """
    user, err = require_user()
    if err:
        return None, err

    try:
        email = (user.get("email") or "").strip().lower()
        if not email:
            return None, ("Admin access required", 403)
        if email not in _admin_email_set():
            return None, ("Admin access required", 403)
        return user, None
    except Exception as e:
        return None, (f"Admin check failed: {e}", 500)


# ----------------- COURSE HOLD (ADMIN CONTROL) -----------------
def is_course_held(uid: str, course_title: str) -> bool:
    """Return True if the given course is on hold for the user."""
    if not uid or not course_title:
        return False
    try:
        db = get_db()
        doc = db.course_holds.find_one({"uid": uid, "courseTitle": course_title}, projection={"_id": 0, "held": 1})
        return bool((doc or {}).get("held"))
    except Exception:
        return False


def block_if_held(uid: str, course_title: str):
    """Return (response, status_code) on block, else None."""
    if is_course_held(uid, course_title):
        return (jsonify({"error": "This course is currently on hold by admin."}), 403)
    return None


def derive_course_title_from_video(uid: str, video_url: str):
    """Best-effort: find which course contains a given video_url for the user."""
    if not uid or not video_url:
        return None
    db = get_db()
    # Only project the minimum fields.
    states = db.course_states.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1, "videos": 1})

    def iter_urls(obj):
        if obj is None:
            return
        if isinstance(obj, list):
            for it in obj:
                yield from iter_urls(it)
        elif isinstance(obj, dict):
            # week wrapper
            if isinstance(obj.get("videos"), list):
                for v in obj.get("videos"):
                    yield from iter_urls(v)
                return
            # single video
            if isinstance(obj.get("video"), str):
                yield obj.get("video")
                return
            # build_weekly_json shape
            for _, v in obj.items():
                yield from iter_urls(v)

    for st in states:
        vids = st.get("videos")
        for u in iter_urls(vids):
            if u == video_url:
                return st.get("courseTitle")
    return None

def formdata_hash(obj) -> str:
    try:
        raw = json.dumps(obj or {}, sort_keys=True, ensure_ascii=False).encode("utf-8")
    except Exception:
        raw = str(obj).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()



# ----------------- QUIZ STATE -----------------
QUIZ_STORE = {}       # quiz_id -> quiz(with answers)
QUIZ_ATTEMPTS = {}    # quiz_id -> attempts_count (0,1,2)
MAX_ATTEMPTS = None  # unlimited attempts until pass
PASS_PERCENT = 0.40  # 40% to pass         # pass score for quiz

# ✅ NEW: cache quiz per video_url
QUIZ_BY_VIDEO = {}    # video_url -> {quiz_id, questions_only, full_quiz}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10MB upload limit (PDF chat)

# ----------------- PROFILE PHOTO UPLOADS -----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROFILE_UPLOAD_DIR = os.path.join(BASE_DIR, "uploads", "profile")
os.makedirs(PROFILE_UPLOAD_DIR, exist_ok=True)

ALLOWED_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}


@app.route("/uploads/profile/<path:filename>")
def serve_profile_upload(filename):
    # Public endpoint to serve profile photos
    return send_from_directory(PROFILE_UPLOAD_DIR, filename)

ALLOWED_ORIGINS = [
    "https://zenith.vinothkumarts.in",
    "https://zenith-frontend-red.vercel.app",   # your vercel domain
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "https://www.zenithlearning.site"
]

from flask_cors import CORS

CORS(
  app,
  resources={r"/*": {"origins": ALLOWED_ORIGINS}},
  supports_credentials=True,
  allow_headers=["Content-Type", "Authorization"],
  methods=["GET","POST","DELETE","OPTIONS"]
)


# ----------------- BASIC HEALTH -----------------
# Uptime monitors often probe HEAD/GET /. Return 200 so deployments don't look "down".
@app.route("/", methods=["GET", "HEAD"])
def root_health():
    return ("OK", 200)


# ----------------- AUTH: SYNC FIREBASE USER TO MONGODB -----------------
@app.route("/auth/firebase", methods=["POST"])
def auth_firebase():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    payload = request.get_json() or {}
    uid = payload.get("uid") or user.get("uid")
    email = payload.get("email") or user.get("email")
    name = payload.get("name") or user.get("name") or ""
    photoURL = payload.get("photoURL") or ""
    providerId = payload.get("providerId") or "firebase"

    if not uid:
        return jsonify({"error": "uid missing"}), 400


    db = get_db()
    now = datetime.now(timezone.utc)

    # ✅ Canonicalize by email to avoid duplicate Mongo users when the same person logs in via password + SSO.
    email_norm = (email or "").strip().lower()
    if email_norm:
        # Upsert by email (unique), then attach latest Firebase uid/provider/name/photo
        res = db.users.update_one(
            {"email": email_norm},
            {"$set": {
                "uid": uid,
                "email": email_norm,
                "name": name,
                "photoURL": photoURL,
                "providerId": providerId,
                "updatedAt": now,
            },
             # role ONLY on first insert
             "$setOnInsert": {"createdAt": now}},
            upsert=True
        )
    else:
        # Fallback (rare): no email available -> upsert by uid
        res = db.users.update_one(
        {"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid},
            {"$set": {
                "uid": uid,
                "email": email_norm,
                "name": name,
                "photoURL": photoURL,
                "providerId": providerId,
                "updatedAt": now,
            },
             "$setOnInsert": {"createdAt": now}},
            upsert=True
        )

    # ✅ Send welcome email for SSO (Google) only for first-time sign-in
    try:
        is_new = bool(getattr(res, "upserted_id", None))
        if is_new and (providerId or "").lower().startswith("google"):
            # Avoid duplicates if user doc already has the flag
            udoc = db.users.find_one({"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid}, projection={"_id": 0, "welcomeSsoSent": 1, "name": 1}) or {}
            if not udoc.get("welcomeSsoSent"):
                home_url = _safe_public_url("/")
                body = f"""
                <p style="margin:0 0 10px 0;">Hi <b>{(name or 'there')}</b>,</p>
                <p style="margin:0 0 10px 0;">
                  Welcome to <b>Zenith Learning</b>! Your Google sign-in is set up and you're ready to start learning.
                </p>
                <p style="margin:0 0 10px 0;">Here are a few quick tips to get the best experience:</p>
                <ul style="margin:10px 0 10px 20px;">
                  <li>Create a roadmap to structure your learning in weeks</li>
                  <li>Complete quizzes after each video to unlock the next content</li>
                  <li>Track progress in <b>My Courses</b> and continue anytime</li>
                </ul>
                """
                html = _brand_email(
                    "Welcome to Zenith Learning 🎉",
                    "Your Google sign-in is ready — begin your roadmap today.",
                    body,
                    primary_cta={"label": "Start learning", "url": home_url},
                    secondary_cta={"label": "Contact support", "url": _safe_public_url("/contact")}
                )
                _send_email(email, "Welcome to Zenith Learning", html, kind="auth")
                db.users.update_one({"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid}, {"$set": {"welcomeSsoSent": True, "updatedAt": now}})
    except Exception as _e:
        pass


    user_doc = db.users.find_one({"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid}, {"_id": 0})
    return jsonify({
        "ok": True,
        "user": user_doc
    })

    
# ----------------- AUTH EMAIL FLOWS (Verification + Password reset) -----------------

def _new_token() -> str:
    return uuid.uuid4().hex + uuid.uuid4().hex  # 64 chars

@app.route("/auth/send-verification", methods=["POST"])
def auth_send_verification():
    """
    Send a custom verification email (SMTP) for password-based signups.
    Requires Firebase ID token. We DO NOT create accounts here; we simply verify/activate.
    """
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    email = user.get("email") or ""
    uid = user.get("uid") or ""
    if not email or not uid:
        return jsonify({"error": "email/uid missing"}), 400

    db = get_db()
    now = datetime.now(timezone.utc)

    # ✅ Ensure Mongo user doc exists (canonical by email to prevent duplicates)
    try:
        email_norm = (email or "").strip().lower()
        filt = {"email": email_norm} if email_norm else {"uid": uid}
        db.users.update_one(
            filt,
            {"$set": {
                "uid": uid,
                "email": email_norm,
                "name": (user.get("name") or ""),
                "photoURL": (user.get("picture") or user.get("photoURL") or ""),
                "providerId": "password",
                "updatedAt": now,
            }, "$setOnInsert": {"createdAt": now}},
            upsert=True
        )
    except Exception as e:
        print(f"[WARN] Mongo upsert during send-verification failed: {e}")

    # If already verified, avoid re-sending
    try:
        _init_firebase_admin()
        u = fb_auth.get_user(uid)
        if getattr(u, "email_verified", False):
            return jsonify({"ok": True, "alreadyVerified": True})
    except Exception:
        pass

    token = _new_token()
    # Store token
    db.email_verifications.insert_one({
        "token": token,
        "uid": uid,
        "email": email,
        "used": False,
        "createdAt": now,
        "expiresAt": now + timedelta(hours=24),
    })

    verify_url = _safe_public_url(f"/verify-email?token={token}")
    body = f"""
    <p style="margin:0 0 10px 0;">Hi <b>{(user.get('name') or 'there')}</b>,</p>
    <p style="margin:0 0 10px 0;">
      Thanks for signing up for <b>Zenith Learning</b>. To activate your account and start saving courses and progress,
      please verify your email address.
    </p>
    <ul style="margin:10px 0 10px 20px;">
      <li>This verification link is valid for <b>24 hours</b>.</li>
      <li>If you didn’t create this account, you can safely ignore this email.</li>
    </ul>
    <p style="margin:10px 0 0 0;">Click the button below to continue.</p>
    """
    html = _brand_email(
        "Verify your email to activate your Zenith account",
        "One quick step to activate your account and start learning.",
        body,
        primary_cta={"label": "Open verification page", "url": verify_url},
    )
    _send_email(email, "Verify your Zenith Learning email", html, kind="auth")
    return jsonify({"ok": True})

@app.route("/auth/verify-email", methods=["POST"])
def auth_verify_email():
    """
    Verify token, then mark Firebase user as emailVerified.
    Also sends a welcome email (manual signup) exactly once.
    """
    data = request.get_json() or {}
    token = (data.get("token") or "").strip()
    if not token:
        return jsonify({"error": "token required"}), 400

    db = get_db()
    now = datetime.now(timezone.utc)


    rec = db.email_verifications.find_one({"token": token})
    if not rec:
        return jsonify({"error": "Invalid or expired token"}), 400
    if rec.get("used"):
        return jsonify({"ok": True, "alreadyUsed": True})
    exp = _to_utc_aware(rec.get("expiresAt"))
    if exp and isinstance(exp, datetime) and exp < now:
        return jsonify({"error": "Token expired"}), 400

    uid = rec.get("uid")
    email = rec.get("email")

    # ✅ Ensure user exists in MongoDB even if the user verified before first login
    db.users.update_one(
        {"uid": uid},
        {"$set": {"uid": uid, "email": email, "updatedAt": now},
         "$setOnInsert": {"createdAt": now, "providerId": "password"}},
        upsert=True
    )

    try:
        _init_firebase_admin()
        fb_auth.update_user(uid, email_verified=True)
    except Exception as e:
        return jsonify({"error": f"Firebase update failed: {e}"}), 500

    db.email_verifications.update_one({"token": token}, {"$set": {"used": True, "usedAt": now}})

    # Welcome email for manual signup (send only once)
    user_doc = db.users.find_one({"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid}, projection={"_id": 0}) or {}
    if not user_doc.get("welcomeManualSent"):
        home_url = _safe_public_url("/")
        body = f"""
        <p style="margin:0 0 10px 0;">Hi <b>{(user_doc.get('name') or 'there')}</b>,</p>
        <p style="margin:0 0 10px 0;">
          Your email is verified ✅. Welcome to <b>Zenith Learning</b>!
        </p>
        <p style="margin:0 0 10px 0;">
          Here’s what you can do next:
        </p>
        <ul style="margin:10px 0 10px 20px;">
          <li>Create a personalized learning roadmap</li>
          <li>Track video progress and quiz performance</li>
          <li>Resume anytime from <b>My Courses</b></li>
          <li>Use Notes to save important points</li>
        </ul>
        <p style="margin:10px 0 0 0;">Ready to continue?</p>
        """
        html = _brand_email(
            "Welcome to Zenith Learning 🎉",
            "Your account is active. Let’s start learning.",
            body,
            primary_cta={"label": "Go to Zenith", "url": home_url},
            secondary_cta={"label": "Contact support", "url": _safe_public_url("/contact")}
        )
        _send_email(email, "Welcome to Zenith Learning", html, kind="auth")
        db.users.update_one({"email": (email or "").strip().lower()} if (email or "").strip() else {"uid": uid}, {"$set": {"welcomeManualSent": True, "updatedAt": now}}, upsert=True)

    return jsonify({"ok": True})

@app.route("/auth/password-reset/request", methods=["POST"])
def auth_password_reset_request():
    """
    Send a password reset email using SMTP.
    Always returns ok to avoid account enumeration.
    """
    data = request.get_json() or {}
    email = (data.get("email") or "").strip().lower()
    if not email:
        return jsonify({"error": "email required"}), 400

    db = get_db()
    now = datetime.now(timezone.utc)


    token = _new_token()

    try:
        _init_firebase_admin()
        u = fb_auth.get_user_by_email(email)
        uid = u.uid
    except Exception:
        # Return ok even if not found
        return jsonify({"ok": True})

    db.password_resets.insert_one({
        "token": token,
        "uid": uid,
        "email": email,
        "used": False,
        "createdAt": now,
        "expiresAt": now + timedelta(hours=1),
    })

    reset_url = _safe_public_url(f"/reset-password?token={token}")
    body = f"""
    <p style="margin:0 0 10px 0;">Hi,</p>
    <p style="margin:0 0 10px 0;">
      We received a request to reset the password for your Zenith Learning account (<b>{email}</b>).
    </p>
    <p style="margin:0 0 10px 0;">
      If you initiated this request, click the button below to set a new password.
      This link is valid for <b>1 hour</b>.
    </p>
    <ul style="margin:10px 0 10px 20px;">
      <li>If you did not request a reset, you can ignore this email.</li>
      <li>For security, never share your password with anyone.</li>
    </ul>
    """
    html = _brand_email(
        "Reset your Zenith Learning password",
        "Use the button below to set a new password (valid for 1 hour).",
        body,
        primary_cta={"label": "Reset password", "url": reset_url},
    )
    _send_email(email, "Reset your Zenith Learning password", html, kind="auth")
    return jsonify({"ok": True})

@app.route("/auth/password-reset/confirm", methods=["POST"])
def auth_password_reset_confirm():
    data = request.get_json() or {}
    token = (data.get("token") or "").strip()
    new_password = (data.get("newPassword") or "").strip()
    if not token or not new_password:
        return jsonify({"error": "token and newPassword required"}), 400
    if len(new_password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400

    db = get_db()
    now = datetime.now(timezone.utc)

    rec = db.password_resets.find_one({"token": token})
    if not rec:
        return jsonify({"error": "Invalid or expired token"}), 400
    if rec.get("used"):
        return jsonify({"error": "Token already used"}), 400
    exp = _to_utc_aware(rec.get("expiresAt"))
    if exp and isinstance(exp, datetime) and exp < now:
        return jsonify({"error": "Token expired"}), 400

    uid = (rec.get("uid") or "").strip()
    email = (rec.get("email") or "").strip().lower()

    if not uid:
        return jsonify({"error": "Invalid reset record (uid missing)"}), 400

    # ✅ Ensure user exists in MongoDB
    db.users.update_one(
        {"uid": uid},
        {"$set": {"uid": uid, "email": email, "updatedAt": now},
         "$setOnInsert": {"createdAt": now, "providerId": "password"}},
        upsert=True
    )

    # ✅ Prevent using the same password again (compare with stored hash)
    try:
        user_doc = db.users.find_one({"uid": uid}, projection={"_id": 0, "passwordHash": 1}) or {}
        old_hash = (user_doc.get("passwordHash") or "").strip()
        new_hash = _pw_hash(uid, new_password)

        # If we have old hash and it's same => block
        if old_hash and old_hash == new_hash:
            return jsonify({
                "error": "New password cannot be the same as your old password. Please choose a different password."
            }), 400
    except Exception:
        # don't fail reset if hash check fails unexpectedly
        pass

    # ✅ Update Firebase password
    try:
        _init_firebase_admin()
        fb_auth.update_user(uid, password=new_password)
    except Exception as e:
        return jsonify({"error": f"Password update failed: {e}"}), 500

    # ✅ Mark token used
    db.password_resets.update_one({"token": token}, {"$set": {"used": True, "usedAt": now}})

    # ✅ Save latest password hash (for next-time reuse detection)
    try:
        db.users.update_one(
            {"uid": uid},
            {"$set": {"passwordHash": _pw_hash(uid, new_password), "updatedAt": now}},
            upsert=True
        )
    except Exception:
        pass

    # ✅ Confirmation email
    body = f"""
    <p style="margin:0 0 10px 0;">Hi,</p>
    <p style="margin:0 0 10px 0;">
      Your Zenith Learning password has been successfully updated for <b>{email}</b>.
    </p>
    <p style="margin:0 0 10px 0;">
      If you didn’t make this change, please reset your password again immediately and contact support.
    </p>
    """
    html = _brand_email(
        "Your password was updated",
        "Your password reset is complete.",
        body,
        primary_cta={"label": "Login to Zenith", "url": _safe_public_url("/login")},
        secondary_cta={"label": "Contact support", "url": _safe_public_url("/contact")}
    )
    _send_email(email, "Zenith Learning — Password updated", html, kind="auth")
    return jsonify({"ok": True})

@app.route("/contact", methods=["POST"])
def contact_send():
    data = request.get_json() or {}
    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip()
    subject = (data.get("subject") or "Contact request").strip()
    message = (data.get("message") or "").strip()

    if not name or not email or not message:
        return jsonify({"error": "name, email and message are required"}), 400

    # Admin mail
    admin_subject = f"[Zenith Contact] {subject}"
    body_admin = f"""
    <p style="margin:0 0 10px 0;">You received a new message from the Zenith Contact form.</p>
    <table style="border-collapse:collapse;font-size:14px;">
      <tr><td style="padding:6px 10px;border:1px solid #e5e7eb;"><b>Name</b></td><td style="padding:6px 10px;border:1px solid #e5e7eb;">{name}</td></tr>
      <tr><td style="padding:6px 10px;border:1px solid #e5e7eb;"><b>Email</b></td><td style="padding:6px 10px;border:1px solid #e5e7eb;">{email}</td></tr>
      <tr><td style="padding:6px 10px;border:1px solid #e5e7eb;"><b>Subject</b></td><td style="padding:6px 10px;border:1px solid #e5e7eb;">{subject}</td></tr>
    </table>
    <p style="margin:12px 0 6px 0;"><b>Message:</b></p>
    <div style="white-space:pre-wrap;background:#f9fafb;border:1px solid #e5e7eb;border-radius:12px;padding:12px;">{message}</div>
    """
    html_admin = _brand_email(
        "New Contact Form Message",
        "A user submitted a message via the Contact page.",
        body_admin,
        primary_cta={"label": "Open Zenith", "url": _safe_public_url("/")},
    )
    _send_email(CONTACT_INBOX, admin_subject, html_admin, kind="contact", reply_to=email)

    # User acknowledgement
    body_user = f"""
    <p style="margin:0 0 10px 0;">Hi <b>{name}</b>,</p>
    <p style="margin:0 0 10px 0;">
      Thanks for reaching out to <b>Zenith Learning</b>. We’ve received your message and our team will respond as soon as possible.
    </p>
    """
    html_user = _brand_email(
        "We received your message",
        "Thanks for contacting Zenith Learning — we’ll reply soon.",
        body_user,
        primary_cta={"label": "Back to Zenith", "url": _safe_public_url("/")},
    )
    _send_email(email, "Zenith Learning — We received your message", html_user, kind="contact")

    return jsonify({"ok": True})


# ----------------- PROFILE (EXTRA FIELDS + LOCAL PHOTO) -----------------
# --- Profile update email helpers ---
_PROFILE_FIELD_LABELS = {
    "name": "Name",
    "dob": "Date of Birth",
    "education": "Education",
    "college": "College",
    "degree": "Degree",
    "department": "Department",
    "yearBatch": "Batch / Year",
    "year": "Year",
    "phone": "Phone",
    "location": "Location",
    "bio": "Bio",
    "photoURL": "Profile Photo (URL)",
    "photoLocalURL": "Profile Photo",
}

def _norm_profile_val(v):
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    # keep numbers/bools as-is but stringify for comparisons
    return str(v).strip()

def _profile_change_rows(before: dict, after_partial: dict):
    """
    Compute change rows ONLY for fields present in `after_partial` (i.e., what client attempted to write).
    Returns list of dicts: {field,label,change,old,new}
    """
    rows = []
    before = before or {}
    for k, new_v in (after_partial or {}).items():
        if k in ("updatedAt",):
            continue
        old_v = _norm_profile_val(before.get(k, ""))
        new_vn = _norm_profile_val(new_v)
        if old_v == new_vn:
            continue

        if old_v and not new_vn:
            change = "Deleted"
        elif (not old_v) and new_vn:
            change = "Added"
        else:
            change = "Updated"

        rows.append({
            "field": k,
            "label": _PROFILE_FIELD_LABELS.get(k, k),
            "change": change,
            "old": old_v,
            "new": new_vn,
        })
    return rows

def _send_profile_update_email(to_email: str, name: str, rows: list, when_ist: str, cta_url: str = ""):
    """Send a profile update email listing only changed fields."""
    if not to_email or not rows:
        return

    # Build rows table
    tr_html = ""
    for r in rows:
        new_disp = r["new"] if r["change"] != "Deleted" else "—"
        # Avoid super-long bio blobs in email tables
        if r["field"] == "bio" and len(new_disp) > 140:
            new_disp = new_disp[:140].rstrip() + "…"
        if r["field"] in ("photoLocalURL", "photoURL"):
            # don't dump long urls; indicate status
            if r["change"] == "Deleted":
                new_disp = "Removed"
            else:
                new_disp = "Updated"

        tr_html += f"""
        <tr>
          <td style="padding:8px 10px;border:1px solid #e5e7eb;"><b>{r['label']}</b></td>
          <td style="padding:8px 10px;border:1px solid #e5e7eb;">{r['change']}</td>
          <td style="padding:8px 10px;border:1px solid #e5e7eb;word-break:break-word;">{new_disp}</td>
        </tr>
        """

    body = f"""
    <p style="margin:0 0 10px 0;">Hi <b>{(name or 'there').strip()}</b>,</p>
    <p style="margin:0 0 12px 0;">Your Zenith profile was updated on <b>{when_ist}</b>. Here’s what changed:</p>

    <table style="border-collapse:collapse;font-size:14px;width:100%;max-width:560px;">
      <tr>
        <td style="padding:8px 10px;border:1px solid #e5e7eb;background:#f9fafb;"><b>Field</b></td>
        <td style="padding:8px 10px;border:1px solid #e5e7eb;background:#f9fafb;"><b>Change</b></td>
        <td style="padding:8px 10px;border:1px solid #e5e7eb;background:#f9fafb;"><b>New value</b></td>
      </tr>
      {tr_html}
    </table>

    <p style="margin:14px 0 0 0;color:#6b7280;font-size:13px;">
      If you didn’t make this change, please reset your password and contact support.
    </p>
    """

    html = _brand_email(
        "Profile updated",
        "Your Zenith profile details were updated.",
        body,
        primary_cta={"label": "Open Profile", "url": (cta_url or _safe_public_url("/profile"))},
    )
    _send_email(to_email, "Zenith Learning — Profile updated", html, kind="profile")


@app.route("/profile/me", methods=["GET"])
def profile_me_get():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code
    db = get_db()
    doc = db.users.find_one({"uid": user["uid"]}, {"_id": 0})
    return jsonify({"ok": True, "user": doc or {}})


@app.route("/profile/me", methods=["POST"])
def profile_me_post():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    # Only allow these fields to be written from the client
    allowed = [
        "name",
        "dob",
        "education",
        "college",
        "degree",
        "department",
        "yearBatch",
        "year",
        "phone",
        "location",
        "bio",
        "photoURL",
        "photoLocalURL",
    ]
    update = {k: (data.get(k) or "") for k in allowed if k in data}

    # Normalize strings
    for k, v in list(update.items()):
        if isinstance(v, str):
            update[k] = v.strip()

    # Backward compatibility: accept `year` from client and also set `yearBatch`
    if update.get("year") and not update.get("yearBatch"):
        update["yearBatch"] = update["year"]

    # Keep empty strings out (optional) except name
    update["name"] = (update.get("name") or "").strip()

    db = get_db()
    uid = user["uid"]

    # Read previous state so we can email ONLY what changed
    before = db.users.find_one({"uid": uid}, {"_id": 0}) or {}

    # Compute change rows BEFORE we add updatedAt
    rows = _profile_change_rows(before, update)

    now = datetime.now(timezone.utc)
    update["updatedAt"] = now

    db.users.update_one({"uid": uid}, {"$set": update}, upsert=True)
    doc = db.users.find_one({"uid": uid}, {"_id": 0})

    # Send profile update email (only if there were changes)
    try:
        to_email = (user.get("email") or (doc or {}).get("email") or before.get("email") or "").strip()
        display_name = (doc or {}).get("name") or update.get("name") or before.get("name") or ""
        if to_email and rows:
            _send_profile_update_email(
                to_email=to_email,
                name=display_name,
                rows=rows,
                when_ist=_now_ist_str(),
                cta_url=_safe_public_url("/profile"),
            )
    except Exception as e:
        print("[MAIL] profile update email failed:", e)

    return jsonify({"ok": True, "user": doc})


def _remove_existing_avatar_files(uid: str):
    """Remove any existing avatar files for the user (uid.*)"""
    try:
        for fn in os.listdir(PROFILE_UPLOAD_DIR):
            if fn.startswith(f"{uid}."):
                try:
                    os.remove(os.path.join(PROFILE_UPLOAD_DIR, fn))
                except:
                    pass
    except:
        pass


def _profile_avatar_url(uid: str):
    host = (request.host_url or "").rstrip("/")
    return f"{host}/profile/photo/{uid}"


@app.route("/profile/photo/<uid>", methods=["GET"])
def profile_photo_get(uid):
    """Serve the user's avatar from MongoDB GridFS if available, else from local uploads."""
    db = get_db()
    fs = gridfs.GridFS(db)

    user_doc = db.users.find_one({"uid": uid}, {"_id": 0, "avatarFileId": 1}) or {}
    fid = user_doc.get("avatarFileId")
    if fid:
        try:
            gf = fs.get(ObjectId(fid))
            data = gf.read()
            ct = getattr(gf, "content_type", None) or getattr(gf, "contentType", None) or "application/octet-stream"
            return (data, 200, {"Content-Type": ct, "Cache-Control": "no-store"})
        except:
            pass

    # Fallback: try to serve local file uid.* if exists
    try:
        for fn in os.listdir(PROFILE_UPLOAD_DIR):
            if fn.startswith(f"{uid}."):
                return send_from_directory(PROFILE_UPLOAD_DIR, fn)
    except:
        pass

    return jsonify({"error": "photo not found"}), 404


@app.route("/profile/photo", methods=["POST"])
def profile_photo_upload():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    if "photo" not in request.files:
        return jsonify({"error": "photo file missing"}), 400

    f = request.files.get("photo")
    if not f or not f.filename:
        return jsonify({"error": "invalid photo"}), 400

    db = get_db()
    fs = gridfs.GridFS(db)
    uid = user["uid"]

    original = f.filename
    _, ext = os.path.splitext(original)
    ext = (ext or "").lower()
    if not ext:
        mt = (f.mimetype or "").lower()
        if "jpeg" in mt or "jpg" in mt:
            ext = ".jpg"
        elif "png" in mt:
            ext = ".png"
        elif "webp" in mt:
            ext = ".webp"

    if ext not in ALLOWED_IMAGE_EXTS:
        return jsonify({"error": "Only JPG / PNG / WEBP allowed"}), 400

    # Read current avatar state (for delete + email diff)
    current = db.users.find_one(
        {"uid": uid},
        {"_id": 0, "avatarFileId": 1, "photoLocalURL": 1, "email": 1, "name": 1}
    ) or {}

    # Replace existing avatar (GridFS) if present
    old_fid = current.get("avatarFileId")
    if old_fid:
        try:
            fs.delete(ObjectId(old_fid))
        except Exception:
            pass

    # Store in GridFS
    content_type = (f.mimetype or "").strip() or "application/octet-stream"
    file_id = fs.put(f.stream.read(), filename=f"avatar_{uid}{ext}", contentType=content_type)
    public_url = _profile_avatar_url(uid)

    # Clean any old local file leftovers
    _remove_existing_avatar_files(uid)

    now = datetime.now(timezone.utc)
    db.users.update_one(
        {"uid": uid},
        {"$set": {
            "avatarFileId": str(file_id),
            "photoLocalURL": public_url,
            "photoURL": "",
            "updatedAt": now
        }},
        upsert=True,
    )

    doc = db.users.find_one({"uid": uid}, {"_id": 0}) or {}

    # ✅ Send email: profile photo updated (only mention photo)
    try:
        before = current or {}
        rows = _profile_change_rows(before, {"photoLocalURL": public_url})

        to_email = (user.get("email") or doc.get("email") or before.get("email") or "").strip()
        display_name = (doc.get("name") or before.get("name") or user.get("name") or "").strip()

        if to_email and rows:
            _send_profile_update_email(
                to_email=to_email,
                name=display_name,
                rows=rows,
                when_ist=_now_ist_str(),
                cta_url=_safe_public_url("/profile"),
            )
    except Exception as e:
        print("[MAIL] profile photo upload email failed:", e)

    return jsonify({"ok": True, "user": doc})



@app.route("/profile/photo", methods=["DELETE"])
def profile_photo_delete():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    uid = user["uid"]
    db = get_db()
    fs = gridfs.GridFS(db)

    # Read current avatar state (for email diff)
    current = db.users.find_one(
        {"uid": uid},
        {"_id": 0, "avatarFileId": 1, "photoLocalURL": 1, "email": 1, "name": 1}
    ) or {}

    # Delete from GridFS if present
    old_fid = current.get("avatarFileId")
    if old_fid:
        try:
            fs.delete(ObjectId(old_fid))
        except Exception:
            pass

    # Also remove any local leftovers
    _remove_existing_avatar_files(uid)

    now = datetime.now(timezone.utc)
    db.users.update_one(
        {"uid": uid},
        {"$set": {"photoLocalURL": "", "photoURL": "", "avatarFileId": "", "updatedAt": now}},
        upsert=True,
    )

    doc = db.users.find_one({"uid": uid}, {"_id": 0}) or {}

    # ✅ Send email: profile photo deleted (only mention photo)
    try:
        before = current or {}
        rows = _profile_change_rows(before, {"photoLocalURL": ""})

        to_email = (user.get("email") or doc.get("email") or before.get("email") or "").strip()
        display_name = (doc.get("name") or before.get("name") or user.get("name") or "").strip()

        if to_email and rows:
            _send_profile_update_email(
                to_email=to_email,
                name=display_name,
                rows=rows,
                when_ist=_now_ist_str(),
                cta_url=_safe_public_url("/profile"),
            )
    except Exception as e:
        print("[MAIL] profile photo delete email failed:", e)

    return jsonify({"ok": True, "user": doc})

# ----------------- COURSE STATE CACHE (ROADMAP + VIDEOS) -----------------
@app.route("/course/state/get", methods=["POST"])
def course_state_get():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    courseTitle = (data.get("courseTitle") or "").strip()

    # If "Other Domains" is used, resolve to actual subject if present
    if courseTitle.lower() in ["other domains", "other domain"] and isinstance(data.get("formData"), dict) and data.get("formData", {}).get("subject"):
        courseTitle = str(data.get("formData", {}).get("subject")).strip()
    formData = data.get("formData")

    if not courseTitle:
        return jsonify({"error": "courseTitle required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    db = get_db()
    q = {"uid": user["uid"], "courseTitle": courseTitle}
    if formData is not None:
        q["formHash"] = formdata_hash(formData)

    doc = db.course_states.find_one(q, sort=[("updatedAt", -1)], projection={"_id": 0})
    if not doc:
        # fallback: latest for the course (any formData)
        doc = db.course_states.find_one({"uid": user["uid"], "courseTitle": courseTitle}, sort=[("updatedAt", -1)], projection={"_id": 0})

    if not doc:
        return jsonify({"found": False}), 200

    return jsonify({"found": True, "state": doc})


@app.route("/course/state/save", methods=["POST"])
def course_state_save():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    courseTitle = (data.get("courseTitle") or "").strip()
    formData = data.get("formData") or {}
    # If user selected "Other Domains", store the actual subject as the course title
    if courseTitle.lower() in ["other domains", "other domain"] and isinstance(formData, dict) and formData.get("subject"):
        courseTitle = str(formData.get("subject")).strip()
    roadmap = data.get("roadmap")
    videos = data.get("videos")

    if not courseTitle or roadmap is None or videos is None:
        return jsonify({"error": "courseTitle, roadmap, videos required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    db = get_db()
    now = datetime.now(timezone.utc)


    h = formdata_hash(formData)

    # ✅ Detect first-time course creation (for enrolled email)
    existing_course = db.course_states.find_one(
        {"uid": user["uid"], "courseTitle": courseTitle, "formHash": h},
        projection={"_id": 1}
    )


    db.course_states.update_one(
        {"uid": user["uid"], "courseTitle": courseTitle, "formHash": h},
        {"$set": {"uid": user["uid"], "courseTitle": courseTitle, "formHash": h, "formData": formData, "roadmap": roadmap, "videos": videos, "updatedAt": now},
         "$setOnInsert": {"createdAt": now}},
        upsert=True
    )


    # ✅ Course enrolled email (send only on first insert)
    if not existing_course:
        try:
            to_email = user.get("email") or ""
            course_id = hashlib.sha1(f"{user['uid']}|{courseTitle}|{h}".encode("utf-8")).hexdigest()[:12]
            body = f"""
            <p style="margin:0 0 10px 0;">Hi <b>{(user.get('name') or 'there')}</b>,</p>
            <p style="margin:0 0 10px 0;">
              You’re enrolled in a new course on <b>Zenith Learning</b>.
              Your roadmap and videos are ready.
            </p>
            <ul style="margin:10px 0 10px 20px;">
              <li><b>Course:</b> {courseTitle}</li>
              <li><b>Course ID:</b> {course_id}</li>
              <li><b>Quiz Result Time (IST):</b> {_now_ist_str()}</li>
            </ul>
            <p style="margin:10px 0 0 0;">Open the course to continue:</p>
            """
            open_url = _safe_public_url(f"/course/{courseTitle}/form")
            html = _brand_email(
                "Course enrolled — your roadmap is ready",
                f"You enrolled in {courseTitle}. Open your roadmap and start learning.",
                body,
                primary_cta={"label": "Open course", "url": open_url},
                secondary_cta={"label": "View My Courses", "url": _safe_public_url("/my-courses")}
            )
            _send_email(to_email, f"Zenith Learning — Enrolled: {courseTitle}", html, kind="courses")
        except Exception:
            pass


    return jsonify({"ok": True})


@app.route("/course/state/delete", methods=["POST"])
def course_state_delete():
    """User-initiated discontinue/remove course.

    Removes saved course state, progress, quiz progress, cached quizzes, and hold entries
    for the current user + course.
    """
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    payload = request.get_json() or {}
    course_title = (payload.get("courseTitle") or "").strip()
    if not course_title:
        return jsonify({"error": "courseTitle missing"}), 400

    uid = user.get("uid")

    # ✅ send unenrolled mail (before deleting data)
    try:
        to_email = (user.get("email") or "").strip()
        uname = user.get("name") or "there"
        when_ist = _now_ist_str()

        body = f"""
        <p style="margin:0 0 10px 0;">Hi <b>{uname}</b>,</p>
        <p style="margin:0 0 10px 0;">
          You have discontinued (unenrolled) from the course below on <b>Zenith Learning</b>.
        </p>
        <ul style="margin:10px 0 10px 20px;">
          <li><b>Course:</b> {course_title}</li>
          <li><b>Unenrolled at:</b> {when_ist}</li>
        </ul>
        <p style="margin:10px 0 0 0;">
          You can enroll again anytime from <b>My Courses</b>.
        </p>
        """

        html = _brand_email(
            "Course discontinued (Unenrolled)",
            "You have been unenrolled from a course on Zenith Learning.",
            body,
            primary_cta={"label": "View My Courses", "url": _safe_public_url("/my-courses")},
            secondary_cta={"label": "Contact support", "url": _safe_public_url("/contact")}
        )
        _send_email(to_email, f"Zenith Learning — Unenrolled: {course_title}", html, kind="courses")
    except Exception:
        pass

    db = get_db()

    # ✅ Remove ALL course-related data for this user + course.
    # NOTE: Some caches (summaries/transcripts/mindmaps) are keyed by uid+video_url
    # and can be shared across courses, so we do not wipe them here.
    db.course_states.delete_many({"uid": uid, "courseTitle": course_title})
    db.progress.delete_many({"uid": uid, "courseTitle": course_title})
    db.quiz_progress.delete_many({"uid": uid, "courseTitle": course_title})
    db.quizzes.delete_many({"uid": uid, "courseTitle": course_title})
    db.course_holds.delete_many({"uid": uid, "courseTitle": course_title})

    # Unlock/quiz flags + attempt history
    db.course_progress.delete_many({"uid": uid, "courseTitle": course_title})
    db.quiz_attempts.delete_many({"uid": uid, "courseTitle": course_title})

    # Course-scoped chat history + notes
    db.chat_sessions.delete_many({"uid": uid, "courseTitle": course_title})
    db.notes.delete_many({"uid": uid, "courseTitle": course_title})

    # Completion mail tracker for this course
    db.course_completion_mails.delete_many({"uid": uid, "courseTitle": course_title})

    return jsonify({"ok": True})
@app.route("/courses/list", methods=["GET"])
def courses_list():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()

    states = list(db.course_states.find(
        {"uid": user["uid"]},
        projection={"_id": 0, "courseTitle": 1, "formData": 1, "videos": 1, "updatedAt": 1, "createdAt": 1}
    ))

    progress_docs = list(db.progress.find(
        {"uid": user["uid"]},
        projection={"_id": 0, "courseTitle": 1, "progressByVideo": 1, "updatedAt": 1}
    ))
    progress_by_course = {d.get("courseTitle"): d for d in progress_docs}

    # ✅ Course unlock / quiz state (passed/completed per video)
    course_progress_docs = list(db.course_progress.find(
        {"uid": user["uid"]},
        projection={"_id": 0, "courseTitle": 1, "quizPassedMap": 1, "quizCompletedMap": 1, "currentGlobalId": 1, "highestUnlockedId": 1, "updatedAt": 1}
    ))
    course_progress_by_course = {d.get("courseTitle"): d for d in course_progress_docs}

    items = []
    for st in states:
        title = st.get("courseTitle")
        videos = st.get("videos") or []
        # ✅ Robust total video counter
        # We support multiple shapes because different parts of the app store videos differently:
        # 1) [{"week": "Week 1", "videos": [..]}]
        # 2) [{"Week 1": [..]}, {"Week 2": [..]}]   (current build_weekly_json output)
        # 3) flat list of video dicts: [{"topic":..., "video":...}]
        def count_videos(obj):
            if obj is None:
                return 0
            # list/array
            if isinstance(obj, list):
                total_local = 0
                for it in obj:
                    total_local += count_videos(it)
                return total_local
            # dict/object
            if isinstance(obj, dict):
                # week wrapper with "videos" array
                if isinstance(obj.get("videos"), list):
                    return len(obj.get("videos"))
                # single video
                if "video" in obj and ("topic" in obj or "title" in obj):
                    return 1
                # build_weekly_json shape: {"Week 1": [..]}
                total_local = 0
                for _, v in obj.items():
                    if isinstance(v, list):
                        total_local += len(v)
                    elif isinstance(v, dict):
                        total_local += count_videos(v)
                return total_local
            return 0

        total = count_videos(videos)

        pd = progress_by_course.get(title) or {}
        pmap = pd.get("progressByVideo") or {}

        cp = course_progress_by_course.get(title) or {}
        qp_map = cp.get("quizPassedMap") or {}
        qc_map = cp.get("quizCompletedMap") or {}

        def _count_true(dct):
            if not isinstance(dct, dict):
                return 0
            return sum(1 for _, v in dct.items() if bool(v))

        passed_quizzes = _count_true(qp_map)
        completed_quizzes = _count_true(qc_map)

        completed = 0
        started = False
        best_resume = None  # (scoreTuple, url)
        for vurl, p in pmap.items():
            if not isinstance(p, dict):
                continue
            percent = float(p.get("percent") or 0)
            current = float(p.get("current") or 0)
            if percent > 0 or current > 0:
                started = True
            is_completed = bool(p.get("completed")) or percent >= 98
            if is_completed:
                completed += 1
            else:
                score = (percent, current)
                if best_resume is None or score > best_resume[0]:
                    best_resume = (score, vurl)

        # ✅ Status: Not Started / Doing / Completed
        # - Completed when all videos completed OR all quizzes passed
        # - Doing when ANY video progress exists OR any quiz completed/passed
        if total > 0 and (completed >= total or passed_quizzes >= total):
            status = "Completed"
        elif started or completed > 0 or passed_quizzes > 0 or completed_quizzes > 0:
            status = "Doing"
        else:
            status = "Not Started"

        held = is_course_held(user["uid"], title)

        items.append({
            "courseTitle": title,
            "displayTitle": (st.get("formData") or {}).get("subject") if (str(title).strip().lower() in ["other domains","other domain"]) else title,
            "totalVideos": total,
            "completedVideos": completed,
            "totalQuizzes": total,
            "passedQuizzes": passed_quizzes,
            "completedQuizzes": completed_quizzes,
            "resumeGlobalId": int(cp.get("currentGlobalId") or 1),
            "highestUnlockedId": int(cp.get("highestUnlockedId") or 1),
            "status": status,
            "held": held,
            "resumeVideoUrl": best_resume[1] if best_resume else None,
            "updatedAt": (st.get("updatedAt") or st.get("createdAt") or pd.get("updatedAt"))
        })

    def _ts(item):
        d = item.get("updatedAt")
        if not d:
            return 0
        try:
            return int(d.timestamp())
        except Exception:
            return 0

    items.sort(key=_ts, reverse=True)
    return jsonify({"items": items})


# ----------------- ADMIN: USERS OVERVIEW -----------------
@app.route("/admin/users", methods=["GET"])
def admin_users():
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()

    def count_videos(obj):
        if obj is None:
            return 0
        if isinstance(obj, list):
            return sum(count_videos(it) for it in obj)
        if isinstance(obj, dict):
            if isinstance(obj.get("videos"), list):
                return len(obj.get("videos"))
            if "video" in obj and ("topic" in obj or "title" in obj):
                return 1
            total_local = 0
            for _, v in obj.items():
                if isinstance(v, list):
                    total_local += len(v)
                elif isinstance(v, dict):
                    total_local += count_videos(v)
            return total_local
        return 0

    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    admin_email_lc = (ADMIN_EMAIL or "").strip().lower()

    users_out = []
    for u in db.users.find({}, projection={"_id": 0, "uid": 1, "email": 1, "name": 1}):
        uid = u.get("uid")
        email = (u.get("email") or "").strip()
        is_admin = email.lower() == admin_email_lc

        # ✅ Do not show learning progress for the admin mailbox
        if is_admin:
            users_out.append({
                "uid": uid,
                "email": email,
                "name": u.get("name"),
                "isAdmin": True,
                "courses": [],
                "overallPercent": 0,
                "heldCourses": [],
            })
            continue

        states = list(db.course_states.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1, "videos": 1}))
        holds = list(db.course_holds.find({"uid": uid, "held": True}, projection={"_id": 0, "courseTitle": 1}))
        held_set = set([h.get("courseTitle") for h in holds if h.get("courseTitle")])

        courses = []
        percents = []
        for st in states:
            title = st.get("courseTitle")
            total = count_videos(st.get("videos"))
            cp = db.course_progress.find_one({"uid": uid, "courseTitle": title}, projection={"_id": 0, "quizPassedMap": 1}) or {}
            passed = count_true(cp.get("quizPassedMap") or {})
            percent = 0
            if total > 0:
                percent = (passed / total) * 100
                percents.append(percent)
            courses.append({
                "courseKey": title,
                "courseTitle": title,
                "totalQuizzes": total,
                "passedQuizzes": passed,
                "percent": round(percent, 2),
                "held": title in held_set,
            })

        overall = round(sum(percents) / len(percents), 2) if percents else 0
        users_out.append({
            "uid": uid,
            "email": email,
            "name": u.get("name"),
            "isAdmin": False,
            "courses": courses,
            "overallPercent": overall,
            "heldCourses": sorted(list(held_set)),
        })

    users_out.sort(key=lambda x: (not x.get("isAdmin"), x.get("email") or ""))
    return jsonify({"ok": True, "adminEmail": ADMIN_EMAIL, "users": users_out})


# ----------------- ADMIN: ROLE MANAGEMENT -----------------
@app.route("/admin/promote", methods=["POST"])
def admin_promote():
    return jsonify({"error": "Role management removed. Admin is fixed to ADMIN_EMAIL."}), 410


@app.route("/admin/demote", methods=["POST"])
def admin_demote():
    return jsonify({"error": "Role management removed. Admin is fixed to ADMIN_EMAIL."}), 410


# ----------------- ADMIN: DELETE USER -----------------
@app.route("/admin/delete-user", methods=["POST"])
def admin_delete_user():
    """Delete a user and all their learning data.

    Body: {"uid": "<firebase uid>"}
    Safety: admin cannot delete themselves.
    """
    admin_user, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    target_uid = (data.get("uid") or "").strip()
    if not target_uid:
        return jsonify({"error": "uid required"}), 400
    if admin_user.get("uid") == target_uid:
        return jsonify({"error": "You cannot delete your own account."}), 400

    db = get_db()

    # ✅ fetch user's email + name before deletion
    target_doc = db.users.find_one(
        {"uid": target_uid},
        projection={"_id": 0, "email": 1, "name": 1}
    ) or {}

    target_email = (target_doc.get("email") or "").strip()
    target_name = (target_doc.get("name") or "there")

    # ✅ delete user data
    # remove user doc (uid is stored on the canonical email doc)
    db.users.delete_one({"uid": target_uid})
    if target_email:
        db.users.delete_one({"email": target_email.strip().lower()})
    db.course_states.delete_many({"uid": target_uid})
    db.course_progress.delete_many({"uid": target_uid})
    db.progress.delete_many({"uid": target_uid})
    db.quizzes.delete_many({"uid": target_uid})
    db.transcripts.delete_many({"uid": target_uid})
    db.summaries.delete_many({"uid": target_uid})
    db.course_holds.delete_many({"uid": target_uid})


    # ✅ delete from Firebase Auth too
    try:
        _init_firebase_admin()
        try:
            fb_auth.delete_user(target_uid)
        except Exception:
            # fallback: if uid not found but email exists, try delete by email
            if target_email:
                u = fb_auth.get_user_by_email(target_email)
                fb_auth.delete_user(u.uid)
    except Exception as e:
        print(f"[WARN] Firebase delete failed for {target_uid}: {e}")

    # ✅ send account removed email (after deletion is ok, but use saved email)
    try:
        when_ist = _now_ist_str()
        body = f"""
        <p style="margin:0 0 10px 0;">Hi <b>{target_name}</b>,</p>
        <p style="margin:0 0 10px 0;">
          Your <b>Zenith Learning</b> account has been removed by the administrator.
        </p>
        <ul style="margin:10px 0 10px 20px;">
          <li><b>Removed at:</b> {when_ist}</li>
        </ul>
        <p style="margin:10px 0 0 0;">
          If you believe this was a mistake, please contact support.
        </p>
        """
        html = _brand_email(
            "Your Zenith account has been removed",
            "Your account was removed by the administrator.",
            body,
            primary_cta={"label": "Contact support", "url": _safe_public_url("/contact")},
            secondary_cta={"label": "Open Zenith", "url": _safe_public_url("/")}
        )
        _send_email(target_email, "Zenith Learning — Account removed", html, kind="mocktest")
    except Exception:
        pass

    return jsonify({"ok": True, "uid": target_uid})



# ----------------- ADMIN: HOLD / UNHOLD COURSE FOR USER -----------------
@app.route("/admin/course-hold", methods=["POST"])
def admin_course_hold():
    """Hold/unhold a course for a specific user.

    Body: {"uid": "...", "courseTitle": "...", "held": true/false}
    """
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    uid = (data.get("uid") or "").strip()
    courseTitle = (data.get("courseTitle") or "").strip()
    held = bool(data.get("held"))

    if not uid or not courseTitle:
        return jsonify({"error": "uid and courseTitle required"}), 400

    db = get_db()
    now = datetime.now(timezone.utc)


    db.course_holds.update_one(
        {"uid": uid, "courseTitle": courseTitle},
        {"$set": {"uid": uid, "courseTitle": courseTitle, "held": held, "updatedAt": now},
         "$setOnInsert": {"createdAt": now}},
        upsert=True,
    )

    # ✅ Notify user (admin action)
    try:
        udoc = db.users.find_one({"uid": uid}, projection={"_id": 0, "email": 1, "name": 1}) or {}
        to_email = (udoc.get("email") or "").strip()
        if to_email:
            uname = udoc.get("name") or "there"
            status = "ON HOLD" if held else "ACTIVE"
            body = f"""
            <p style="margin:0 0 10px 0;">Hi <b>{uname}</b>,</p>
            <p style="margin:0 0 10px 0;">
              Your course access has been updated by the admin on <b>Zenith Learning</b>.
            </p>
            <ul style="margin:10px 0 10px 20px;">
              <li><b>Course:</b> {courseTitle}</li>
              <li><b>Status:</b> <b>{status}</b></li>
              <li><b>Updated at (IST):</b> {_now_ist_str()}</li>
            </ul>
            <p style="margin:10px 0 0 0;">If you think this is a mistake, contact support.</p>
            """
            html = _brand_email(
                f"Course status updated — {status}",
                f"Course: {courseTitle} • Status: {status}",
                body,
                primary_cta={"label": "Open My Courses", "url": _safe_public_url("/my-courses")},
                secondary_cta={"label": "Contact support", "url": _safe_public_url("/contact")},
            )
            _send_email(to_email, f"Zenith Learning — Course {status}: {courseTitle}", html, kind="mocktest")
    except Exception:
        pass

    return jsonify({"ok": True, "uid": uid, "courseTitle": courseTitle, "held": held})


# ----------------- ADMIN: ANALYTICS ENDPOINTS -----------------
@app.route("/admin/courses-studying", methods=["GET"])
def admin_courses_studying():
    """Aggregate: how many users are studying each course."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    # course_states can include multiple formHash docs; group by uid+courseTitle
    pipeline = [
        {"$group": {"_id": {"uid": "$uid", "courseTitle": "$courseTitle"}}},
        {"$group": {"_id": "$_id.courseTitle", "users": {"$sum": 1}}},
        {"$sort": {"users": -1}},
    ]
    rows = list(db.course_states.aggregate(pipeline))

    holds = list(db.course_holds.find({"held": True}, projection={"_id": 0, "uid": 1, "courseTitle": 1}))
    hold_counts = {}
    for h in holds:
        ct = h.get("courseTitle")
        if not ct:
            continue
        hold_counts[ct] = hold_counts.get(ct, 0) + 1

    out = []
    for r in rows:
        ct = r.get("_id")
        out.append({
            "courseTitle": ct,
            "usersStudying": int(r.get("users") or 0),
            "usersOnHold": int(hold_counts.get(ct, 0)),
        })
    return jsonify({"ok": True, "courses": out})


@app.route("/admin/course-progress", methods=["GET"])
def admin_course_progress():
    """Detailed progress per user per course (quiz-passed/total).

    ✅ Requirement: course progress should be calculated from quiz progress.
    """
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()

    # Map course -> total videos (derived from course_states by taking max)
    def count_videos(obj):
        if obj is None:
            return 0
        if isinstance(obj, list):
            return sum(count_videos(it) for it in obj)
        if isinstance(obj, dict):
            if isinstance(obj.get("videos"), list):
                return len(obj.get("videos"))
            if "video" in obj and ("topic" in obj or "title" in obj):
                return 1
            return sum(count_videos(v) for v in obj.values())
        return 0

    totals = {}
    for st in db.course_states.find({}, projection={"_id": 0, "courseTitle": 1, "videos": 1}):
        ct = st.get("courseTitle")
        if not ct:
            continue
        totals[ct] = max(totals.get(ct, 0), count_videos(st.get("videos")))

    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    users = {u["uid"]: u for u in db.users.find({}, projection={"_id": 0, "uid": 1, "email": 1, "name": 1})}

    # quiz progress is stored in course_progress collection
    rows = []
    for cp in db.course_progress.find({}, projection={"_id": 0, "uid": 1, "courseTitle": 1, "quizPassedMap": 1, "updatedAt": 1}):
        uid = cp.get("uid")
        ct = cp.get("courseTitle")
        u = users.get(uid) or {}
        email_lc = (u.get("email") or "").strip().lower()
        if email_lc == (ADMIN_EMAIL or "").strip().lower():
            continue

        total = int(totals.get(ct, 0))
        passed = count_true(cp.get("quizPassedMap") or {})
        pct = int(round((passed / total) * 100)) if total else 0

        rows.append({
            "uid": uid,
            "name": u.get("name"),
            "email": u.get("email"),
                        "courseTitle": ct,
            "totalQuizzes": total,
            "passedQuizzes": passed,
            "percent": pct,
            "held": is_course_held(uid, ct),
            "updatedAt": cp.get("updatedAt"),
        })

    rows.sort(key=lambda x: (x.get("held") is False, x.get("percent") or 0), reverse=True)
    return jsonify({"ok": True, "rows": rows})


@app.route("/admin/quiz-performance", methods=["GET"])
def admin_quiz_performance():
    """Aggregate quiz pass performance per course."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()

    # Total quizzes per course derived same as /courses/list (total videos)
    def count_videos(obj):
        if obj is None:
            return 0
        if isinstance(obj, list):
            return sum(count_videos(it) for it in obj)
        if isinstance(obj, dict):
            if isinstance(obj.get("videos"), list):
                return len(obj.get("videos"))
            if "video" in obj and ("topic" in obj or "title" in obj):
                return 1
            return sum(count_videos(v) for v in obj.values())
        return 0

    totals = {}
    for st in db.course_states.find({}, projection={"_id": 0, "courseTitle": 1, "videos": 1}):
        ct = st.get("courseTitle")
        if not ct:
            continue
        totals[ct] = max(totals.get(ct, 0), count_videos(st.get("videos")))

    # Sum passed counts across users
    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    agg = {}
    for cp in db.course_progress.find({}, projection={"_id": 0, "courseTitle": 1, "quizPassedMap": 1}):
        ct = cp.get("courseTitle")
        if not ct:
            continue
        agg[ct] = agg.get(ct, 0) + count_true(cp.get("quizPassedMap") or {})

    courses = []
    for ct, total in totals.items():
        passed_total = int(agg.get(ct, 0))
        # usersStudying: unique uid in course_states
        usersStudying = len(list(db.course_states.aggregate([
            {"$match": {"courseTitle": ct}},
            {"$group": {"_id": "$uid"}},
        ])))
        # Max possible passed = usersStudying * total
        denom = usersStudying * int(total)
        pass_rate = round((passed_total / denom) * 100, 2) if denom else 0
        courses.append({
            "courseTitle": ct,
            "totalQuizzesPerUser": int(total),
            "usersStudying": int(usersStudying),
            "passedQuizzesTotal": passed_total,
            "passRatePercent": pass_rate,
        })
    courses.sort(key=lambda x: x.get("passRatePercent") or 0, reverse=True)
    return jsonify({"ok": True, "courses": courses})

# ----------------- WATCH PROGRESS -----------------
@app.route("/progress/get", methods=["GET"])
def progress_get():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    courseTitle = (request.args.get("courseTitle") or "").strip()
    if not courseTitle:
        return jsonify({"error": "courseTitle required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    db = get_db()
    doc = db.progress.find_one({"uid": user["uid"], "courseTitle": courseTitle}, projection={"_id": 0})
    if not doc:
        return jsonify({"progressByVideo": {}})

    return jsonify({"progressByVideo": doc.get("progressByVideo", {})})




@app.route("/admin/mocktest-analytics", methods=["GET"])
def admin_mocktest_analytics():
    """Admin-only aggregate analytics for mock test sessions."""
    _, err = require_admin()
    if err:
        return jsonify({"error": err[0]}), err[1]

    try:
        db = get_db()
        col = _mocktest_sessions_col(db)
        # Limit to last 2000 sessions to keep response fast
        docs = list(col.find({}).sort("created_at", -1).limit(2000))

        total_sessions = len(docs)
        user_ids = set()
        submitted = 0
        passed = 0
        total_score_sum = 0
        total_marks_sum = 0

        mode_counts = {}
        day_counts = {}  # YYYY-MM-DD -> count

        def _day_key(dt):
            try:
                if isinstance(dt, str):
                    # keep prefix
                    return dt[:10]
                return (dt.date().isoformat())
            except Exception:
                return None

        for d in docs:
            uid = d.get("uid")
            if uid:
                user_ids.add(uid)

            mode = (d.get("mode") or "all")
            mode_counts[mode] = mode_counts.get(mode, 0) + 1

            dk = _day_key(d.get("created_at") or d.get("createdAt"))
            if dk:
                day_counts[dk] = day_counts.get(dk, 0) + 1

            ts = d.get("total_score")
            tm = d.get("total_marks")
            status = (d.get("status") or "").lower()

            if ts is not None and tm is not None and int(tm or 0) > 0:
                submitted += 1
                try:
                    ts_i = int(ts or 0)
                    tm_i = int(tm or 0)
                except Exception:
                    ts_i, tm_i = 0, 0
                total_score_sum += ts_i
                total_marks_sum += tm_i
                # Define pass as >= 60%
                if tm_i > 0 and (ts_i / tm_i) >= 0.6:
                    passed += 1
            elif status in ("submitted", "completed"):
                submitted += 1

        avg_score_pct = round((total_score_sum / total_marks_sum) * 100, 2) if total_marks_sum else 0.0
        pass_rate = round((passed / submitted) * 100, 2) if submitted else 0.0

        # Recent 14 days timeline
        today = datetime.now(timezone.utc).date()
        timeline = []
        for i in range(13, -1, -1):
            day = (today - timedelta(days=i)).isoformat()
            timeline.append({"day": day, "count": int(day_counts.get(day, 0))})

        # Mode breakdown
        mode_breakdown = [{"mode": k, "count": int(v)} for k, v in sorted(mode_counts.items(), key=lambda x: (-x[1], x[0]))]

        return jsonify({
            "ok": True,
            "totalSessions": total_sessions,
            "uniqueUsers": len(user_ids),
            "submittedSessions": submitted,
            "passedSessions": passed,
            "avgScorePercent": avg_score_pct,
            "passRate": pass_rate,
            "modeBreakdown": mode_breakdown,
            "timeline": timeline,
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/admin/mocktest-analytics/users", methods=["GET"])
def admin_mocktest_analytics_users():
    """Admin-only user-wise analytics for mock test sessions."""
    _, err = require_admin()
    if err:
        return jsonify({"error": err[0]}), err[1]

    try:
        db = get_db()
        col = _mocktest_sessions_col(db)
        # Limit to last 5000 sessions for performance
        docs = list(col.find({}).sort("created_at", -1).limit(5000))

        users = {u.get("uid"): u for u in db.users.find({}, projection={"_id": 0, "uid": 1, "email": 1, "name": 1, "photoURL": 1, "photoLocalURL": 1})}

        agg = {}
        for d in docs:
            uid = d.get("uid")
            if not uid:
                continue
            a = agg.setdefault(uid, {
                "uid": uid,
                "sessions": 0,
                "submitted": 0,
                "passed": 0,
                "scoreSum": 0,
                "marksSum": 0,
                "modeCounts": {},
                "lastAttempt": None,
            })
            a["sessions"] += 1

            mode = (d.get("mode") or "all")
            a["modeCounts"][mode] = a["modeCounts"].get(mode, 0) + 1

            created = d.get("created_at") or d.get("createdAt")
            # keep most recent
            try:
                if a["lastAttempt"] is None:
                    a["lastAttempt"] = created
                else:
                    # compare strings safely; ISO compares lexicographically
                    if isinstance(created, str) and isinstance(a["lastAttempt"], str):
                        if created > a["lastAttempt"]:
                            a["lastAttempt"] = created
            except Exception:
                pass

            ts = d.get("total_score")
            tm = d.get("total_marks")
            status = (d.get("status") or "").lower()
            if ts is not None and tm is not None and int(tm or 0) > 0:
                a["submitted"] += 1
                try:
                    ts_i = int(ts or 0)
                    tm_i = int(tm or 0)
                except Exception:
                    ts_i, tm_i = 0, 0
                a["scoreSum"] += ts_i
                a["marksSum"] += tm_i
                if tm_i > 0 and (ts_i / tm_i) >= 0.6:
                    a["passed"] += 1
            elif status in ("submitted", "completed"):
                a["submitted"] += 1

        rows = []
        for uid, a in agg.items():
            u = users.get(uid) or {}
            email_lc = (u.get("email") or "").strip().lower()
            # hide admin account
            if email_lc == (ADMIN_EMAIL or "").strip().lower():
                continue

            avg_score_pct = round((a["scoreSum"] / a["marksSum"]) * 100, 2) if a["marksSum"] else 0.0
            pass_rate = round((a["passed"] / a["submitted"]) * 100, 2) if a["submitted"] else 0.0
            rows.append({
                "uid": uid,
                "name": u.get("name"),
                "email": u.get("email"),
                "photoURL": u.get("photoURL"),
                "photoLocalURL": u.get("photoLocalURL"),
                "sessions": int(a["sessions"]),
                "submitted": int(a["submitted"]),
                "passed": int(a["passed"]),
                "avgScorePercent": avg_score_pct,
                "passRate": pass_rate,
                "modeBreakdown": [{"mode": k, "count": int(v)} for k, v in sorted(a["modeCounts"].items(), key=lambda x: (-x[1], x[0]))],
                "lastAttempt": a.get("lastAttempt"),
            })

        rows.sort(key=lambda r: (r.get("passRate") or 0, r.get("avgScorePercent") or 0, r.get("sessions") or 0), reverse=True)
        return jsonify({"ok": True, "rows": rows})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/admin/quiz-analytics/users", methods=["GET"])
def admin_quiz_analytics_users():
    """Admin-only user-wise analytics for quizzes (across all courses)."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        db = get_db()

        # Map course -> total quizzes (videos count)
        def count_videos(obj):
            if obj is None:
                return 0
            if isinstance(obj, list):
                return sum(count_videos(it) for it in obj)
            if isinstance(obj, dict):
                if isinstance(obj.get("videos"), list):
                    return len(obj.get("videos"))
                if "video" in obj and ("topic" in obj or "title" in obj):
                    return 1
                return sum(count_videos(v) for v in obj.values())
            return 0

        totals = {}
        for st in db.course_states.find({}, projection={"_id": 0, "courseTitle": 1, "videos": 1}):
            ct = st.get("courseTitle")
            if not ct:
                continue
            totals[ct] = max(totals.get(ct, 0), count_videos(st.get("videos")))

        users = {u.get("uid"): u for u in db.users.find({}, projection={"_id": 0, "uid": 1, "email": 1, "name": 1, "photoURL": 1, "photoLocalURL": 1})}

        def count_true(dct):
            if not isinstance(dct, dict):
                return 0
            return sum(1 for _, v in dct.items() if bool(v))

        agg = {}
        # For each user-course progress, sum totals/passed
        for cp in db.course_progress.find({}, projection={"_id": 0, "uid": 1, "courseTitle": 1, "quizPassedMap": 1, "updatedAt": 1}):
            uid = cp.get("uid")
            ct = cp.get("courseTitle")
            if not uid or not ct:
                continue
            total = int(totals.get(ct, 0))
            passed = int(count_true(cp.get("quizPassedMap") or {}))
            a = agg.setdefault(uid, {
                "uid": uid,
                "courses": 0,
                "totalQuizzes": 0,
                "passedQuizzes": 0,
                "lastUpdated": None,
            })
            a["courses"] += 1
            a["totalQuizzes"] += total
            a["passedQuizzes"] += min(passed, total)
            lu = cp.get("updatedAt")
            try:
                if a["lastUpdated"] is None:
                    a["lastUpdated"] = lu
                else:
                    if isinstance(lu, str) and isinstance(a["lastUpdated"], str) and lu > a["lastUpdated"]:
                        a["lastUpdated"] = lu
            except Exception:
                pass

        rows = []
        for uid, a in agg.items():
            u = users.get(uid) or {}
            email_lc = (u.get("email") or "").strip().lower()
            if email_lc == (ADMIN_EMAIL or "").strip().lower():
                continue

            total = int(a.get("totalQuizzes") or 0)
            passed = int(a.get("passedQuizzes") or 0)
            avg_percent = round((passed / total) * 100, 2) if total else 0.0

            rows.append({
                "uid": uid,
                "name": u.get("name"),
                "email": u.get("email"),
                "photoURL": u.get("photoURL"),
                "photoLocalURL": u.get("photoLocalURL"),
                "courses": int(a.get("courses") or 0),
                "totalQuizzes": total,
                "passedQuizzes": passed,
                "avgPercent": avg_percent,
                "lastUpdated": a.get("lastUpdated"),
            })

        rows.sort(key=lambda r: (r.get("avgPercent") or 0, r.get("passedQuizzes") or 0), reverse=True)
        return jsonify({"ok": True, "rows": rows})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/progress/upsert", methods=["POST"])
def progress_upsert():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    courseTitle = (data.get("courseTitle") or "").strip()
    videoUrl = (data.get("videoUrl") or "").strip()
    progress = data.get("progress") or {}

    if not courseTitle or not videoUrl:
        return jsonify({"error": "courseTitle and videoUrl required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    db = get_db()
    now = datetime.now(timezone.utc)



    db.progress.update_one(
        {"uid": user["uid"], "courseTitle": courseTitle},
        {"$set": {f"progressByVideo.{videoUrl}": progress, "updatedAt": now},
         "$setOnInsert": {"createdAt": now, "uid": user["uid"], "courseTitle": courseTitle}},
        upsert=True
    )

    return jsonify({"ok": True})


# ----------------- COURSE RESUME / QUIZ UNLOCK STATE -----------------
# Stores where the learner left off + which videos are unlocked/passed.
# This fixes:
# 1) If user exits and re-opens the course, it should resume.
# 2) Quiz/videos should remain unlocked across sessions/devices.

@app.route("/course/progress/get", methods=["POST"])
def course_progress_get():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    courseTitle = (data.get("courseTitle") or "").strip()
    if not courseTitle:
        return jsonify({"error": "courseTitle required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    db = get_db()
    doc = db.course_progress.find_one({"uid": user["uid"], "courseTitle": courseTitle}, projection={"_id": 0})
    if not doc:
        return jsonify({"found": False, "progress": None}), 200

    # Ensure required keys exist
    progress = {
        "currentGlobalId": int(doc.get("currentGlobalId") or 1),
        "highestUnlockedId": int(doc.get("highestUnlockedId") or 1),
        "quizPassedMap": doc.get("quizPassedMap") or {},
        "quizSubmittedMap": doc.get("quizSubmittedMap") or {},
        "quizCompletedMap": doc.get("quizCompletedMap") or {},
        "updatedAt": doc.get("updatedAt"),
    }
    return jsonify({"found": True, "progress": progress}), 200


@app.route("/course/progress/save", methods=["POST"])
def course_progress_save():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    courseTitle = (data.get("courseTitle") or "").strip()
    progress = data.get("progress") or {}

    if not courseTitle:
        return jsonify({"error": "courseTitle required"}), 400

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    blocked = block_if_held(user.get("uid"), courseTitle)
    if blocked:
        return blocked

    # sanitize
    def _obj(x):
        return x if isinstance(x, dict) else {}

    currentGlobalId = int(progress.get("currentGlobalId") or 1)
    highestUnlockedId = int(progress.get("highestUnlockedId") or 1)
    quizPassedMap = _obj(progress.get("quizPassedMap"))
    quizSubmittedMap = _obj(progress.get("quizSubmittedMap"))
    quizCompletedMap = _obj(progress.get("quizCompletedMap"))

    db = get_db()
    now = datetime.now(timezone.utc)



    db.course_progress.update_one(
        {"uid": user["uid"], "courseTitle": courseTitle},
        {"$set": {
            "uid": user["uid"],
            "courseTitle": courseTitle,
            "currentGlobalId": currentGlobalId,
            "highestUnlockedId": highestUnlockedId,
            "quizPassedMap": quizPassedMap,
            "quizSubmittedMap": quizSubmittedMap,
            "quizCompletedMap": quizCompletedMap,
            "updatedAt": now,
        },
         "$setOnInsert": {"createdAt": now}},
        upsert=True,
    )


    # ✅ Course completion email (send once when all quizzes/videos are completed)
    try:
        # determine total videos from course_states
        cs = db.course_states.find_one({"uid": user["uid"], "courseTitle": courseTitle}, projection={"_id": 0, "videos": 1}) or {}
        def _count_videos(obj):
            if obj is None:
                return 0
            if isinstance(obj, list):
                return sum(_count_videos(it) for it in obj)
            if isinstance(obj, dict):
                if isinstance(obj.get("videos"), list):
                    return len(obj.get("videos"))
                if "video" in obj and ("topic" in obj or "title" in obj):
                    return 1
                return sum(_count_videos(v) for v in obj.values())
            return 0
        total_videos = _count_videos(cs.get("videos"))
        passed_count = sum(1 for k,v in (quizPassedMap or {}).items() if v)
        completed_count = sum(1 for k,v in (quizCompletedMap or {}).items() if v)

        is_complete = (total_videos > 0) and (passed_count >= total_videos or completed_count >= total_videos)
        if is_complete:
            # Avoid duplicates
            course_key = hashlib.sha1(f"{user['uid']}|{courseTitle}".encode("utf-8")).hexdigest()
            already = db.course_completion_mails.find_one({"courseKey": course_key}, projection={"_id": 1})
            if not already:
                db.course_completion_mails.insert_one({"courseKey": course_key, "uid": user["uid"], "courseTitle": courseTitle, "createdAt": now})
                to_email = user.get("email") or ""
                body = f"""
                <p style="margin:0 0 10px 0;">Hi <b>{(user.get('name') or 'there')}</b>,</p>
                <p style="margin:0 0 10px 0;">
                  Congratulations — you have completed the course <b>{courseTitle}</b> on Zenith Learning 🎉
                </p>
                <ul style="margin:10px 0 10px 20px;">
                  <li><b>Total videos:</b> {total_videos}</li>
                  <li><b>Quizzes passed:</b> {passed_count}</li>
                  <li><b>Completed items:</b> {completed_count}</li>
                </ul>
                <p style="margin:10px 0 10px 0;">
                  Next steps:
                </p>
                <ul style="margin:10px 0 10px 20px;">
                  <li>Review your notes and revise weak areas</li>
                  <li>Try a new roadmap with a higher difficulty or faster pace</li>
                  <li>Share feedback through the Contact page</li>
                </ul>
                """
                html = _brand_email(
                    "Course completed ✅",
                    f"You completed {courseTitle}. Keep the momentum going.",
                    body,
                    primary_cta={"label": "View My Courses", "url": _safe_public_url("/my-courses")},
                    secondary_cta={"label": "Start a new course", "url": _safe_public_url("/")}
                )
                _send_email(to_email, f"Zenith Learning — Course Completed: {courseTitle}", html, kind="courses")
    except Exception:
        pass


    return jsonify({"ok": True})



# ----------------- ENV / CLIENTS -----------------
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")
youtube = build("youtube", "v3", developerKey=YOUTUBE_API_KEY)

# ✅ Supadata
supadata = Supadata(api_key=os.getenv("SUPADATA_KEY", ""))
# ----------------- HELPERS -----------------
def is_quota_error(e: Exception) -> bool:
    msg = str(e).lower()
    return ("429" in msg) or ("quota" in msg) or ("rate limit" in msg) or ("resource exhausted" in msg)


class QuotaExceededError(Exception):
    def __init__(self, message: str = 'Quota exceeded', retry_after: float | None = None):
        super().__init__(message)
        self.retry_after = retry_after


def _extract_retry_after_seconds(msg: str) -> float | None:
    try:
        m = re.search(r"retry in ([0-9]+(?:\.[0-9]+)?)s", msg, flags=re.I)
        if m:
            return float(m.group(1))
    except Exception:
        return None
    return None

def _parse_iso8601_duration_to_seconds(iso: str) -> int:
    """
    Convert YouTube ISO 8601 duration (e.g., PT1H2M10S, PT45S) to seconds.
    Returns 0 if parsing fails.
    """
    if not iso:
        return 0
    iso = iso.strip().upper()
    m = re.match(r"^PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?$", iso)
    if not m:
        return 0
    h = int(m.group(1) or 0)
    mi = int(m.group(2) or 0)
    s = int(m.group(3) or 0)
    return h * 3600 + mi * 60 + s


def _looks_like_shorts(title: str, duration_sec: int, video_url: str = "") -> bool:
    t = (title or "").lower()
    u = (video_url or "").lower()
    # Heuristics: YouTube Shorts are typically <= 60s, and often tagged in title
    if "/shorts/" in u:
        return True
    if "#shorts" in t or " shorts" in t or t.endswith("shorts"):
        return True
    if duration_sec and duration_sec <= 60:
        return True
    return False


def _extract_topic_keywords(topic: str) -> set:
    t = re.sub(r"[^a-z0-9\s]", " ", (topic or "").lower())
    words = [w for w in t.split() if len(w) > 2]
    stop = {
        "the", "and", "for", "with", "from", "that", "this", "into", "your", "you",
        "are", "was", "were", "can", "will", "what", "why", "how", "when", "where",
        "about", "using", "use", "tutorial", "course", "learn", "learning",
        "beginner", "advanced", "english", "in", "to", "of", "on", "a", "an"
    }
    return {w for w in words if w not in stop}

def get_gemini_response(input_prompt: str) -> str:
    # NOTE:
    # Hosted deployments (e.g., Render + Gunicorn) can hit worker timeouts if LLM calls
    # take too long. Keep latency predictable by enforcing an API timeout and limiting
    # output tokens.
    model = genai.GenerativeModel("gemini-2.5-flash")
    try:
        response = model.generate_content(
            input_prompt,
            generation_config={
                "temperature": 0.2,
                "max_output_tokens": 1400,
            },
            # google-generativeai supports request_options with timeout (seconds)
            request_options={"timeout": 22},
        )
    except Exception as e:
        # Surface quota/rate errors with a consistent message so the UI can show it.
        if is_quota_error(e):
            raise RuntimeError("quota exceeded") from e
        raise

    if not getattr(response, "candidates", None):
        raise ValueError("Gemini API returned no candidates.")

    try:
        return response.candidates[0].content.parts[0].text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting or parsing response: {e}")


# ---------------- MINDMAP (JSON TREE) ----------------
def _clean_json_text(s: str) -> str:
    """Best-effort cleanup if the model wraps JSON in code fences or adds leading text."""
    if not s:
        return ""
    txt = s.strip()
    # remove markdown fences
    txt = re.sub(r"^```(?:json)?\s*", "", txt, flags=re.IGNORECASE)
    txt = re.sub(r"```\s*$", "", txt)
    # If the model returns extra text before/after JSON, try to extract the first JSON object.
    # This is intentionally simple and safe.
    first = txt.find("{")
    last = txt.rfind("}")
    if first != -1 and last != -1 and last > first:
        txt = txt[first:last+1]
    # Remove trailing commas before } or ] (common LLM mistake)
    txt = re.sub(r",\s*([}\]])", r"\1", txt)
    return txt.strip()

import json, re
from typing import Any

def _strip_code_fences(s: str) -> str:
    if not s:
        return s
    s = re.sub(r"^```(?:json)?\s*", "", s.strip(), flags=re.IGNORECASE)
    s = re.sub(r"\s*```$", "", s.strip())
    return s.strip()

def _remove_trailing_commas(s: str) -> str:
    return re.sub(r",\s*([}\]])", r"\1", s)

def _find_first_json_span(text: str) -> str | None:
    if not text:
        return None

    start = None
    for i, ch in enumerate(text):
        if ch == '{' or ch == '[':
            start = i
            break
    if start is None:
        return None

    stack = []
    in_str = False
    esc = False

    for j in range(start, len(text)):
        ch = text[j]

        if in_str:
            if esc:
                esc = False
            elif ch == '\\':
                esc = True
            elif ch == '"':
                in_str = False
            continue

        if ch == '"':
            in_str = True
            continue

        if ch in '{[':
            stack.append(ch)
        elif ch in '}]':
            if not stack:
                return None
            top = stack.pop()
            if (top == '{' and ch != '}') or (top == '[' and ch != ']'):
                return None
            if not stack:
                return text[start:j+1]

    return None

def safe_json_parse(model_text: str) -> Any:
    """
    Defensive JSON parsing for LLM outputs.
    - Strips code fences
    - Extracts the first JSON object/array by brace/bracket balancing (respects strings/escapes)
    - Repairs trailing commas
    - Parses with json.loads (with light quote repairs fallback)
    """
    if not model_text or not str(model_text).strip():
        raise ValueError("Empty model output")

    raw = _strip_code_fences(str(model_text))

    span = _find_first_json_span(raw)
    if span is None:
        # try removing any leading text before first bracket
        raw2 = re.sub(r"^[^\[{]*", "", raw, flags=re.DOTALL)
        span = _find_first_json_span(raw2)

    if span is None:
        raise ValueError("Could not parse mindmap JSON from model output.")

    span = _remove_trailing_commas(span)

    try:
        return json.loads(span)
    except json.JSONDecodeError:
        repaired = span.replace("“", '"').replace("”", '"').replace("’", "'")
        repaired = _remove_trailing_commas(repaired)
        return json.loads(repaired)

def repair_json_with_gemini(bad_text: str) -> str:
    """Ask Gemini to repair/convert arbitrary output into STRICT valid JSON mindmap only."""
    bad_text = (bad_text or "").strip()
    if not bad_text:
        return ""
    prompt = f"""You are a strict JSON repair tool.

Return ONLY valid JSON. No markdown. No commentary. No code fences.

Target schema (recursive):
{{"name": string, "children": [ ... ]}}

Rules:
- Use double quotes for all keys/strings.
- No trailing commas.
- Ensure all brackets/braces are closed.
- Remove any non-JSON text.
- If input is truncated, COMPLETE the JSON in the most likely way.
- Keep labels short (2–8 words).

Input:
{bad_text}
"""
    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        resp = model.generate_content(prompt)
        text = (getattr(resp, "text", None) or "").strip()
        return _strip_code_fences(text)
    except Exception:
        return ""


def build_fallback_mindmap(transcript: str, safe_title: str) -> dict:
    # Heuristic fallback (works even without API key):
    # Build a NotebookLM-like structured mindmap by extracting key phrases around common headings.
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", transcript) if s.strip()]

    def _norm(s: str) -> str:
        s = re.sub(r"\s+", " ", (s or "").strip())
        return s

    def _clean_label(s: str, max_words: int = 8) -> str:
        s = _norm(s)
        s = re.sub(r"^(hello|hi)\b.*?\b(tutorial|video)\b[:,\-\s]*", "", s, flags=re.I)
        s = re.sub(r"\b(please\s+subscribe|thanks\s+for\s+watching|hit\s+the\s+bell)\b.*$", "", s, flags=re.I)
        words = re.findall(r"[A-Za-z0-9\-]+", s)
        if not words:
            return "Item"
        return " ".join(words[:max_words])

    def _pick_points(keywords, max_items=7, max_words=7):
        kws = [k.lower() for k in keywords if k]
        out = []
        seen = set()
        for s in sentences:
            sl = s.lower()
            if any(k in sl for k in kws):
                lab = _clean_label(s, max_words=max_words)
                if lab.lower() in seen:
                    continue
                seen.add(lab.lower())
                out.append(lab)
            if len(out) >= max_items:
                break
        return out

    def _branch(name: str, points):
        return {"name": name, "children": [{"name": p, "children": []} for p in points]}

    # Detect main topic words (simple frequency)
    text_l = transcript.lower()
    has_gd = "gradient descent" in text_l
    has_cost = "cost function" in text_l
    has_lr = "learning rate" in text_l
    has_nn = "neural network" in text_l or "weights" in text_l or "bias" in text_l

    children = []

    # Overview / Purpose
    overview = _pick_points(
        ["in this video", "you will learn", "let's", "we will", "today"],
        max_items=4,
        max_words=8,
    )
    purpose = _pick_points(
        ["used to", "purpose", "optimiz", "train", "minimiz", "improve accuracy"],
        max_items=5,
        max_words=7,
    )
    if overview:
        children.append(_branch("Overview", overview))
    if purpose:
        children.append(_branch("Definition and Purpose", purpose))

    # Core concepts (topic-aware)
    core_points = []
    if has_gd:
        core_points += ["Optimization algorithm", "Move downhill (minimize loss)"]
    if has_cost:
        core_points += ["Cost function (prediction error)"]
    if has_lr:
        core_points += ["Learning rate (step size)"]
    if has_nn:
        core_points += ["Weights and biases", "Training on labeled data"]
    core_points += _pick_points(
        ["cost function", "learning rate", "weights", "biases", "minimize", "reduce the cost"],
        max_items=6,
        max_words=7,
    )
    core_points_dedup = []
    seen = set()
    for p in core_points:
        pl = p.lower()
        if pl in seen:
            continue
        seen.add(pl)
        core_points_dedup.append(p)
    if core_points_dedup:
        children.append(_branch("Key Concepts", core_points_dedup[:10]))

    # How it works / Process
    how_points = _pick_points(
        ["take small steps", "direction", "downhill", "gradient", "update", "adjust", "repeat"],
        max_items=8,
        max_words=8,
    )
    if how_points:
        children.append(_branch("How it works", how_points))

    # Examples / Analogies
    example_points = _pick_points(
        ["example", "let's consider", "house", "sold for", "squiggle", "number", "mountain", "maze"],
        max_items=7,
        max_words=8,
    )
    if example_points:
        children.append(_branch("Examples and Analogies", example_points))

    # Types / Variants
    types_points = []
    if "batch" in text_l:
        types_points.append("Batch Gradient Descent")
    if "stochastic" in text_l:
        types_points.append("Stochastic Gradient Descent")
    if "mini" in text_l and "batch" in text_l:
        types_points.append("Mini-batch Gradient Descent")
    if types_points:
        children.append({"name": "Types of Gradient Descent", "children": [{"name": t, "children": []} for t in types_points]})

    # Challenges / Pitfalls
    chall_points = _pick_points(
        ["non-convex", "global minimum", "saddle", "vanishing", "exploding", "unstable", "struggle"],
        max_items=8,
        max_words=8,
    )
    if chall_points:
        children.append(_branch("Challenges", chall_points))

    # Takeaways
    takeaways = _pick_points(
        ["powerful", "commonly used", "despite", "in summary", "end of this", "used to train"],
        max_items=4,
        max_words=8,
    )
    if takeaways:
        children.append(_branch("Summary / Takeaways", takeaways))

    if not children:
        children = [
            {"name": "Overview", "children": []},
            {"name": "Key Concepts", "children": []},
            {"name": "How it works", "children": []},
        ]

    return {"name": safe_title, "children": children}

    # ---------------- SUMMARY ----------------

# REPLACED

def generate_mindmap_tree(transcript: str, title: str = "Mindmap") -> dict:
    """Generate a detailed, concept-focused JSON tree mindmap from transcript.

    Returns schema: {"name": <root>, "children": [ ... ]}
    """
    safe_title = (title or "Mindmap").strip()[:80]
    transcript = (transcript or "").strip()
    if not transcript:
        return {"name": safe_title, "children": []}

    prompt = f"""
You are an expert instructor. Read the transcript, infer the underlying concepts, and produce a DETAILED mindmap that helps a student study.

Output rules (STRICT):
- Output ONLY valid JSON (no markdown, no commentary, no backticks).
- Schema (recursive): {{"name": string, "children": [schema]}}
- Root MUST be exactly: {{"name": "{safe_title}", "children": [...]}}
- Labels must be short: 2–8 words (no long sentences).
- Remove filler/intro/outro like: "hello everyone", "subscribe", "thanks for watching".
- Prefer concepts, definitions, steps, formulas, comparisons, examples, pitfalls.
- Expand acronyms when possible (e.g., SGD -> Stochastic Gradient Descent).
- Depth target: 4–6 levels.
- Breadth target: 6–10 top-level branches; each branch should have 3–10 children (when possible).
- If transcript is shallow, enrich with standard subtopics (1–2 nodes per concept) WITHOUT hallucinating facts.

Recommended top-level branches (adapt as relevant):
- Definition & Purpose
- Key Components / Concepts
- Workflow / Steps
- Examples / Applications
- Tools / Methods / Variants
- Comparisons (if any)
- Metrics / Evaluation (if any)
- Pitfalls / Challenges
- Summary / Takeaways

Transcript:
{transcript}
""".strip()

    def _strip_fences(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.I)
        s = re.sub(r"\s*```$", "", s)
        return s.strip()

    def _clean_labels(node: dict) -> dict:
        if not isinstance(node, dict):
            return {"name": safe_title, "children": []}
        name = str(node.get("name", "")).strip()
        name = re.sub(r"^(hello|hi)\b.*$", "", name, flags=re.I).strip() or name
        name = re.sub(r"\b(thanks for watching|please subscribe|hit the bell)\b.*$", "", name, flags=re.I).strip() or name
        words = re.findall(r"[A-Za-z0-9][A-Za-z0-9\-]*", name)
        name = " ".join(words[:10]) if words else (node.get("name") or "Item")

        kids = node.get("children") or []
        if not isinstance(kids, list):
            kids = []
        cleaned_children = []
        for ch in kids:
            if isinstance(ch, dict):
                cleaned_children.append(_clean_labels(ch))
        cleaned_children = [c for c in cleaned_children if str(c.get("name", "")).strip()]
        return {"name": name, "children": cleaned_children}

    try:
        model_name = os.getenv("GEMINI_MINDMAP_MODEL", "") or "gemini-2.5-flash"
        model = genai.GenerativeModel(model_name)

        resp = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.1,
                "max_output_tokens": 3200,
                "response_mime_type": "application/json",
            },
        )

        raw = getattr(resp, "text", "") or ""
        if not raw and getattr(resp, "candidates", None):
            try:
                raw = resp.candidates[0].content.parts[0].text or ""
            except Exception:
                raw = ""

        raw = _strip_fences(raw)
        raw = _clean_json_text(raw)

        try:
            data = safe_json_parse(raw)
        except Exception as parse_err:
            # Save raw model output for debugging
            try:
                with open("mindmap_llm_raw.txt", "w", encoding="utf-8") as f:
                    f.write(raw if isinstance(raw, str) else str(raw))
                print("❌ GEMINI RAW SAVED -> mindmap_llm_raw.txt")
            except Exception:
                pass

            # Second-pass: ask Gemini to repair/complete JSON (handles missing commas/truncation)
            repaired = repair_json_with_gemini(raw)
            if repaired:
                try:
                    with open("mindmap_llm_repaired.txt", "w", encoding="utf-8") as f:
                        f.write(repaired)
                except Exception:
                    pass
                try:
                    data = safe_json_parse(repaired)
                except Exception:
                    raise parse_err
            else:
                raise parse_err

        if not isinstance(data, dict) or "name" not in data:
            raise ValueError("Invalid mindmap JSON")

        data["name"] = safe_title
        return _clean_labels(data)

    except Exception as e:
        print("❌ GEMINI MINDMAP FAILED:", repr(e))
        # If quota/rate-limit, surface to frontend (do not hide behind fallback mindmap)
        if is_quota_error(e):
            msg = str(e)
            ra = _extract_retry_after_seconds(msg)
            raise QuotaExceededError("Gemini quota exceeded. Please retry later or enable billing.", retry_after=ra)
        return build_fallback_mindmap(transcript, safe_title)


@app.route("/summarize", methods=["POST"])
def summarize():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        transcript = (data.get("transcript") or "").strip()
        summaryType = (data.get("type") or "").strip()
        video_url = (data.get("video_url") or "").strip()

        if not transcript:
            return jsonify({"error": "Transcript not provided"}), 400
        if not summaryType:
            summaryType = "Paragraph"  # default to avoid client-side empty value

        db = get_db()

        cache_q = {"uid": user["uid"], "summaryType": summaryType}
        if video_url:
            cache_q["video_url"] = video_url
        cached = db.summaries.find_one(cache_q, sort=[("updatedAt", -1)], projection={"_id": 0})
        if cached and cached.get("summary"):
            return jsonify({"summary": cached["summary"], "cached": True})

        prompt = f"""Summarize the following transcript in the user's preferred format: {summaryType}.
Keep it accurate and easy to study.

TRANSCRIPT:
{transcript}
"""

        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        summary = (response.text or "").strip()

        now = datetime.now(timezone.utc)
        doc = {"uid": user["uid"], "video_url": video_url or None, "summaryType": summaryType, "summary": summary, "updatedAt": now, "createdAt": now}
        db.summaries.insert_one(doc)

        return jsonify({"summary": summary, "cached": False})

    except Exception as e:
        if is_quota_error(e):
            return jsonify({"error": "Gemini quota exceeded. Please try later or enable billing."}), 429
        return jsonify({"error": str(e)}), 500

# ---------------- ROADMAP GENERATION ----------------
def generate_roadmap(formData):
    prmpt = f"""You are an expert curriculum designer and YouTube SEO specialist.

Create a personalized learning roadmap for a student with the following details:
- Age: {formData['age']}
- Subject: {formData['subject']}
- Current level: {formData['level']}
- Prior experience: {formData['experience']}
- Learning pace: {formData['pace']}
- Goal: {formData['goal']}
- Duration: {formData['duration']} months

IMPORTANT INSTRUCTIONS (STRICT):
1. Every topic MUST explicitly include the subject name "{formData['subject']}".
2. Topics must be written as precise YouTube search queries.
3. Avoid generic terms like "Introduction", "Basics", "Overview" alone.
4. Use phrases that a learner would type into YouTube.
5. Topics must be beginner-to-goal progressive and practical.
6. Do NOT include explanations, descriptions, or extra text.
7. Do NOT include commas inside a topic — each bullet must be ONE clean search query.
8. Do NOT reference any other programming language.

FORMAT:
Week 1:
- <YouTube search query>
- <YouTube search query>

Week 2:
- <YouTube search query>
- <YouTube search query>

Now generate the roadmap.
"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prmpt)
    return response.text

def parse_roadmap(roadmap: str):
    weeks = {}
    current_week = None
    for line in roadmap.split("\n"):
        line = line.strip()
        if line.lower().startswith("week"):
            current_week = line.split(":")[0].strip()
            weeks[current_week] = []
        elif line.startswith("-") and current_week:
            topic = line[1:].strip()
            weeks[current_week].append(topic)

    # ✅ print parsed roadmap dictionary (like your old code)
    print("\n================ PARSED ROADMAP ================")
    print(json.dumps(weeks, indent=2, ensure_ascii=False))
    print("================================================\n")

    return weeks

# ---------------- YOUTUBE HELPERS (WITH DETAILED PRINTS) ----------------
def get_video_details(query, max_results=2, retries=3, delay=2):
    """
    Returns list of videos for a query with like_count.
    ✅ Also prints fetching details similar to your old code.
    """
    video_details = []
    results_count = 0
    next_page_token = None
    page_no = 0

    print(f"\n========== YOUTUBE FETCH START ==========")
    print(f"Query: {query} | max_results={max_results}")
    print(f"========================================\n")

    while results_count < max_results:
        page_no += 1
        try:
            print(f"[SEARCH] page={page_no} results_so_far={results_count} token={next_page_token}")
            # Extra guard to reduce Shorts appearing in search results.
            # (We still hard-filter by duration + title in _looks_like_shorts.)
            safe_query = f"{query} -shorts -#shorts" if query else query
            search_response = youtube.search().list(
                q=safe_query,
                part="snippet",
                maxResults=min(50, max_results - results_count),
                type="video",
                pageToken=next_page_token
            ).execute()

            items = search_response.get("items", [])
            print(f"[SEARCH] page={page_no} returned_items={len(items)}")

            for item in items:
                video_id = item.get("id", {}).get("videoId")
                if not video_id:
                    print("[SKIP] Missing videoId in item")
                    continue

                title = item.get("snippet", {}).get("title", "")
                channel = item.get("snippet", {}).get("channelTitle", "")
                published = item.get("snippet", {}).get("publishedAt", "")
                video_url = f"https://www.youtube.com/watch?v={video_id}"

                # Defaults (in case API omits fields)
                duration_sec = 0
                iso_dur = ""
                _skip_shorts = False


                # Fetch video statistics with retry
                stats = {}
                last_err = None
                for attempt in range(1, retries + 1):
                    try:
                        print(f"  [STATS] attempt={attempt}/{retries} id={video_id}")
                        video_data = youtube.videos().list(
                            part="statistics,contentDetails",
                            id=video_id
                        ).execute()
                        if video_data.get("items"):
                            item0 = video_data["items"][0] or {}
                            stats = item0.get("statistics", {}) or {}
                            content_details = item0.get("contentDetails", {}) or {}
                            iso_dur = (content_details.get("duration") or "").strip()
                            duration_sec = _parse_iso8601_duration_to_seconds(iso_dur)

                            # Filter Shorts / very short videos
                            if _looks_like_shorts(title, duration_sec, video_url):
                                print(f"  [SKIP-SHORTS] id={video_id} dur={duration_sec}s title={title[:60]}")
                                stats = {}
                                _skip_shorts = True
                            else:
                                _skip_shorts = False
                            break
                    except HttpError as e:
                        last_err = e
                        if e.resp.status in [500, 503]:
                            sleep_s = delay * attempt
                            print(f"  [STATS-RETRY] status={e.resp.status} sleeping={sleep_s}s")
                            time.sleep(sleep_s)
                        else:
                            raise

                if last_err and not stats:
                    print(f"  [STATS-FAIL] id={video_id} err={str(last_err)}")
                    continue

                # If we skipped due to Shorts, continue
                if locals().get('_skip_shorts'):
                    continue

                like_count = int(stats.get("likeCount", 0) or 0)
                view_count = int(stats.get("viewCount", 0) or 0)
                comment_count = int(stats.get("commentCount", 0) or 0)

                video_details.append({
                    "url": video_url,
                    "video_id": video_id,
                    "title": title,
                    "channel": channel,
                    "publishedAt": published,
                    "like_count": like_count,
                    "view_count": view_count,
                    "comment_count": comment_count,
                    "duration_sec": duration_sec,
                    "duration_iso": iso_dur
                })

                # ✅ This print line matches your older style
                print(query, " [", results_count, "] : ", video_details[-1])

                results_count += 1
                if results_count >= max_results:
                    break

            next_page_token = search_response.get("nextPageToken")
            if not next_page_token:
                print("[SEARCH] No nextPageToken, ending.")
                break

        except HttpError as e:
            if e.resp.status in [500, 503]:
                print(f"[SEARCH-RETRY] status={e.resp.status} sleeping={delay}s")
                time.sleep(delay)
                continue
            raise

    print(f"\n========== YOUTUBE FETCH END ==========")
    print(f"Query: {query} | fetched={len(video_details)}")
    print(f"======================================\n")

    return video_details

def get_best_video(query: str, used_video_ids: set = None, blocked_keywords: set = None):
    """Pick a non-Shorts YouTube video for a query, avoiding repeats and topic overlap."""
    used_video_ids = used_video_ids or set()
    blocked_keywords = blocked_keywords or set()

    # pull more candidates so filters have room to work
    video_list = get_video_details(query, max_results=12)
    print("here")  # keep your old marker
    if not video_list:
        return None

    def _contains_blocked(title: str) -> bool:
        t = (title or "").lower()
        for kw in blocked_keywords:
            if kw and kw in t:
                return True
        return False

    for v in video_list:
        vid = v.get("video_id") or _yt_id(v.get("url", ""))
        if not vid:
            continue
        if vid in used_video_ids:
            continue
        if _contains_blocked(v.get("title", "")):
            continue
        return v

    # fallback: allow overlap but still avoid duplicates
    for v in video_list:
        vid = v.get("video_id") or _yt_id(v.get("url", ""))
        if vid and vid not in used_video_ids:
            return v
    return None

def build_weekly_json(roadmap):
    weeks = parse_roadmap(roadmap)
    result = []

    used_video_ids = set()      # avoid recommending same video across topics
    covered_keywords = set()    # reduce topic overlap between consecutive topics

    print("\n=========== BUILD WEEKLY JSON START ===========")
    for week, topics in weeks.items():
        print(f"\n[WEEK] {week} topics_count={len(topics)}")
        week_data = []
        for idx, topic in enumerate(topics, start=1):
            q = topic + " in english"
            print(f"  -> ({idx}) Topic: {topic}")
            print(f"     Query: {q}")

            cur_kw = _extract_topic_keywords(topic)
            blocked = covered_keywords - cur_kw

            best = get_best_video(q, used_video_ids=used_video_ids, blocked_keywords=blocked)

            if best and best.get("url"):
                url = best["url"]
                vid = best.get("video_id") or _yt_id(url)
                if vid:
                    used_video_ids.add(vid)
            else:
                url = None

            week_data.append({"topic": topic, "video": url if url else "No video found"})
            covered_keywords |= cur_kw

        result.append({week: week_data})

    print("\n=========== BUILD WEEKLY JSON END ===========\n")
    return result

@app.route("/generate-roadmap", methods=["POST"])
def generate():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        formData = request.get_json() or {}
        roadmap = generate_roadmap(formData)

        print("\n================ ROADMAP RAW TEXT ================")
        print(roadmap)
        print("==================================================\n")

        final_output = build_weekly_json(roadmap)

        with open("videos.json", "w", encoding="utf-8") as f:
            json.dump(final_output, f, indent=2, ensure_ascii=False)

        return jsonify({"roadmap": roadmap, "videos": final_output})

    except Exception as e:
        if is_quota_error(e):
            return jsonify({"error": "Gemini quota exceeded. Please try later or enable billing."}), 429
        return jsonify({"error": str(e)}), 500

# ---------------- TRANSCRIPT ----------------
@app.route("/get-transcript", methods=["POST"])
def get_transcript():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        url = (data.get("url") or "").strip()

        if not url:
            return jsonify({"error": "URL is required"}), 400

        db = get_db()

        cached = db.transcripts.find_one({"uid": user["uid"], "url": url}, projection={"_id": 0})
        if cached and cached.get("transcript"):
            return jsonify({"url": url, "transcript": cached["transcript"], "language": cached.get("language", "en"), "cached": True})

        transcript = supadata.transcript(
            url=url,
            lang="en",
            text=True,
            mode="auto"
        )

        # Supadata returns either a finished transcript with .content or a job_id
        if hasattr(transcript, "content"):
            db.transcripts.update_one(
                {"uid": user["uid"], "url": url},
                {"$set": {"uid": user["uid"], "url": url, "transcript": transcript.content, "language": getattr(transcript, "lang", "en"), "updatedAt": datetime.now(timezone.utc)},
                 "$setOnInsert": {"createdAt": datetime.now(timezone.utc)}},
                upsert=True,
            )
            return jsonify({"url": url, "transcript": transcript.content, "language": transcript.lang, "cached": False})
        else:
            return jsonify({"message": "Transcript is being processed", "job_id": transcript.job_id})

    except Exception as e:
        if is_quota_error(e):
            payload = {"error": "Transcript API quota exceeded. Please retry later or upgrade your plan.", "code": "quota_exceeded"}
            ra = _extract_retry_after_seconds(str(e))
            if ra is not None:
                payload["retry_after"] = ra
            return jsonify(payload), 429
        return jsonify({"error": str(e)}), 500

# ---------------- QUIZ ----------------
def generate_mcq(transcript: str):
    prmpt = f"""You are an expert educational consultant.
A student has given a video transcript and wants to create a multiple-choice quiz based on it.

IMPORTANT LANGUAGE RULE:
- The quiz MUST be written in ENGLISH ONLY.
- If the transcript is not in English, translate the content to English first (internally), then generate the quiz.
- Do NOT output any non-English text.

Transcript:
{transcript}

Create a multiple-choice quiz with 10 questions. Each question should have 4 options.

If a question requires a code snippet (because the transcript includes code), include a **short** snippet inline in the question using single backticks, like: `for i in range(n): ...`.
Keep the snippet on the SAME LINE as the question (do not use multi-line code blocks).
Format:
Question 1: Question
a) Option 1 
b) Option 2
c) Option 3
d) Option 4
Correct Answer: Answer

Only the quiz is required.
"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prmpt)
    return response.text

def _normalize_text(s: str) -> str:
    if s is None:
        return ""
    s = str(s).strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s

def parse_to_json(quiz_text: str):
    """
    Robust MCQ parser.

    Handles:
    - Correct Answer: b
    - Correct Answer: b) Option text
    - Correct Answer: Option text
    """
    quiz_data = []
    blocks = re.split(r"Question\s*\d+\s*:", quiz_text, flags=re.IGNORECASE)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if len(lines) < 3:
            continue

        question = lines[0]

        # collect options a-d
        options = []
        for line in lines[1:]:
            m = re.match(r"^([a-dA-D])[\)\.]\s*(.+)$", line)
            if m:
                options.append(m.group(2).strip())
            if len(options) == 4:
                break

        if len(options) < 2:
            continue

        answer_line = next((l for l in lines if l.lower().startswith("correct answer")), "")
        raw = answer_line.split(":", 1)[1].strip() if ":" in answer_line else ""
        raw_norm = _normalize_text(raw)

        correct_option = ""

        # Case 1: raw starts with a/b/c/d (like "b")
        m = re.match(r"^([a-d])\b", raw_norm)
        if m and len(options) >= 4:
            idx = ord(m.group(1)) - ord("a")
            if 0 <= idx < len(options):
                correct_option = options[idx]

        # Case 2: raw like "b) option text" or "b. option text"
        if not correct_option:
            m2 = re.match(r"^([a-d])[\)\.]\s*(.+)$", raw_norm)
            if m2 and len(options) >= 4:
                idx = ord(m2.group(1)) - ord("a")
                if 0 <= idx < len(options):
                    correct_option = options[idx]

        # Case 3: raw is option text -> fuzzy match
        if not correct_option:
            def canon(x):
                x = _normalize_text(x)
                x = re.sub(r"[^a-z0-9 ]+", "", x)
                return x.strip()

            raw_c = canon(raw)
            for opt in options:
                oc = canon(opt)
                if not oc:
                    continue
                if oc == raw_c or oc in raw_c or raw_c in oc:
                    correct_option = opt
                    break

        if correct_option:
            quiz_data.append({"question": question, "options": options, "answer": correct_option})

    return quiz_data

@app.route("/generate-mcq", methods=["POST"])
def generate_mcq_api():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        transcript = (data.get("transcript") or "").strip()
        video_url = (data.get("video_url") or "").strip()
        courseTitle = (data.get("courseTitle") or "").strip()

        if not video_url:
            return jsonify({"error": "video_url is required"}), 400

        # ✅ Hold enforcement
        if not courseTitle:
            courseTitle = derive_course_title_from_video(user.get("uid"), video_url) or ""
        if courseTitle:
            blocked = block_if_held(user.get("uid"), courseTitle)
            if blocked:
                return blocked

        db = get_db()

        # ✅ 1) If quiz already exists for this user+video, return it (no transcript needed)
        cached = db.quizzes.find_one(
            {"uid": user["uid"], "video_url": video_url},
            sort=[("updatedAt", -1)],
            projection={"_id": 0}
        )
        if cached and cached.get("quiz_id") and cached.get("questions_only"):
            return jsonify({
                "quiz_id": cached["quiz_id"],
                "questions": cached["questions_only"],
                "cached": True
            })

        # ✅ 2) If transcript not sent (common after logout/login), fetch it from DB cache
        if not transcript:
            tdoc = db.transcripts.find_one(
                {"uid": user["uid"], "url": video_url},
                sort=[("updatedAt", -1)],
                projection={"_id": 0, "transcript": 1}
            )
            if tdoc and tdoc.get("transcript"):
                transcript = tdoc["transcript"].strip()

        # ✅ 3) If still no transcript, tell frontend to fetch transcript first
        if not transcript:
            return jsonify({
                "error": "Transcript not available. Please fetch transcript first.",
                "needs_transcript": True
            }), 400

        # ✅ 4) Generate new quiz now
        quiz_text = generate_mcq(transcript)
        quiz_data = parse_to_json(quiz_text)
        quiz_data = quiz_data[:10]

        if not quiz_data:
            return jsonify({"error": "Failed to parse quiz. Try again."}), 500

        quiz_id = str(uuid.uuid4())
        questions_only = [{"question": q["question"], "options": q["options"]} for q in quiz_data]

        now = datetime.now(timezone.utc)
        db.quizzes.insert_one({
            "uid": user["uid"],
            "video_url": video_url,
            "quiz_id": quiz_id,
            "full_quiz": quiz_data,
            "questions_only": questions_only,
            "attempts_used": 0,
            "createdAt": now,
            "updatedAt": now,
        })

        return jsonify({"quiz_id": quiz_id, "questions": questions_only, "cached": False})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/generate-mindmap", methods=["POST"])
def generate_mindmap_api():
    """Generate (or fetch cached) mindmap JSON tree for a given video.

    Body: { video_url: str, transcript?: str, title?: str, courseTitle?: str, force?: bool }
    Returns: { tree: {...}, cached: bool }
    """
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        video_url = (data.get("video_url") or "").strip()
        transcript = (data.get("transcript") or "").strip()
        title = (data.get("title") or "Mindmap").strip() or "Mindmap"
        courseTitle = (data.get("courseTitle") or "").strip()
        force = bool(data.get("force"))

        if not video_url:
            return jsonify({"error": "video_url is required"}), 400

        # ✅ Hold enforcement
        if not courseTitle:
            courseTitle = derive_course_title_from_video(user.get("uid"), video_url) or ""
        if courseTitle:
            blocked = block_if_held(user.get("uid"), courseTitle)
            if blocked:
                return blocked

        db = get_db()

        if not force:
            cached = db.mindmaps.find_one(
                {"uid": user["uid"], "video_url": video_url},
                sort=[("updatedAt", -1)],
                projection={"_id": 0, "tree": 1}
            )
            if cached and cached.get("tree"):
                return jsonify({"tree": cached["tree"], "cached": True})

        # If transcript not sent, try DB transcript cache
        if not transcript:
            tdoc = db.transcripts.find_one(
                {"uid": user["uid"], "url": video_url},
                sort=[("updatedAt", -1)],
                projection={"_id": 0, "transcript": 1}
            )
            if tdoc and tdoc.get("transcript"):
                transcript = tdoc["transcript"].strip()

        if not transcript:
            return jsonify({
                "error": "Transcript not available. Please fetch transcript first.",
                "needs_transcript": True
            }), 400

        tree = generate_mindmap_tree(transcript, title=title)
        now = datetime.now(timezone.utc)
        db.mindmaps.update_one(
            {"uid": user["uid"], "video_url": video_url},
            {"$set": {"uid": user["uid"], "video_url": video_url, "title": title, "tree": tree, "updatedAt": now},
             "$setOnInsert": {"createdAt": now}},
            upsert=True
        )

        return jsonify({"tree": tree, "cached": False})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ---------------- NOTES (Mindmap) ----------------
@app.route("/notes/list", methods=["GET"])
def list_notes():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        db = get_db()
        notes = list(
            db.notes.find(
                {"uid": user["uid"]},
                projection={"_id": 0}
            ).sort("updatedAt", -1).limit(200)
        )
        return jsonify({"notes": notes})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/notes/save-mindmap", methods=["POST"])
def save_mindmap_note():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        note_type = "mindmap"
        title = (data.get("title") or "Mindmap").strip()[:120]
        courseTitle = (data.get("courseTitle") or "").strip()[:200]
        video_url = (data.get("video_url") or "").strip()
        tree = data.get("tree") or {}

        if not video_url:
            return jsonify({"error": "video_url is required"}), 400
        if not isinstance(tree, dict) or "name" not in tree:
            return jsonify({"error": "tree is required"}), 400

        db = get_db()
        now = datetime.now(timezone.utc)
        note_id = (data.get("note_id") or "").strip() or str(uuid.uuid4())

        db.notes.update_one(
            {"uid": user["uid"], "note_id": note_id},
            {
                "$set": {
                    "uid": user["uid"],
                    "note_id": note_id,
                    "type": note_type,
                    "title": title,
                    "courseTitle": courseTitle,
                    "video_url": video_url,
                    "tree": tree,
                    "updatedAt": now,
                },
                "$setOnInsert": {"createdAt": now},
            },
            upsert=True,
        )

        return jsonify({"ok": True, "note_id": note_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/notes/delete", methods=["POST"])
def delete_note():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        note_id = (data.get("note_id") or "").strip()
        if not note_id:
            return jsonify({"error": "note_id is required"}), 400
        db = get_db()
        db.notes.delete_one({"uid": user["uid"], "note_id": note_id})
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _norm(s: str) -> str:
    if not s:
        return ""
    return re.sub(r"^[a-dA-D][\)\.]\s*", "", str(s)).strip().lower()

@app.route("/submit-quiz", methods=["POST"])
def submit_quiz():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    try:
        data = request.get_json() or {}
        quiz_id = data.get("quiz_id")
        courseTitle = (data.get("courseTitle") or "").strip()
        answers = data.get("answers")

        if not quiz_id:
            return jsonify({"error": "quiz_id required"}), 400
        if not isinstance(answers, list):
            return jsonify({"error": "answers must be a list"}), 400

        db = get_db()

        quiz_doc = db.quizzes.find_one({"uid": user["uid"], "quiz_id": quiz_id})
        if not quiz_doc:
            return jsonify({"error": "Invalid quiz_id"}), 400

        # derive course title if not provided
        if not courseTitle:
            courseTitle = derive_course_title_from_video(
                user.get("uid"), (quiz_doc.get("video_url") or "")
            ) or ""

        # block if course is held
        if courseTitle:
            blocked = block_if_held(user.get("uid"), courseTitle)
            if blocked:
                return blocked

        used = int(quiz_doc.get("attempts_used", 0))
        full_quiz = quiz_doc.get("full_quiz", [])

        score = 0
        t0 = time.time()
        results = []
        for i, q in enumerate(full_quiz):
            user_ans = answers[i] if i < len(answers) else ""
            correct = q.get("answer", "")
            ok = (_normalize_text(user_ans) == _normalize_text(correct))
            if ok:
                score += 1
            results.append({
                "question": q.get("question"),
                "selected": user_ans,
                "correct": correct,
                "isCorrect": ok
            })

        used += 1
        db.quizzes.update_one(
            {"uid": user["uid"], "quiz_id": quiz_id},
            {"$set": {"attempts_used": used, "updatedAt": datetime.now(timezone.utc)}}
        )

        required = max(1, math.ceil(len(full_quiz) * PASS_PERCENT))
        passed = score >= required

        # --- Persist per-video quiz attempt summary (used by Admin drilldowns) ---
        # We store a compact doc keyed by (uid, courseTitle, videoNo) with bestScore + lastScore.
        try:
            video_url = (quiz_doc.get("video_url") or "").strip()

            # prefer frontend-provided video number (avoids ordering mismatch)
            req_video_no = data.get("video_no") or data.get("videoNo")
            req_total_videos = data.get("total_videos") or data.get("totalVideos")
            try:
                req_video_no = int(req_video_no) if req_video_no is not None else None
            except Exception:
                req_video_no = None
            try:
                req_total_videos = int(req_total_videos) if req_total_videos is not None else None
            except Exception:
                req_total_videos = None

            video_no, total_videos, video_title = derive_video_meta_from_course(
                db=db,
                uid=user.get("uid"),
                course_title=courseTitle,
                video_url=video_url,
            )

            if req_video_no:
                video_no = req_video_no
            if req_total_videos:
                total_videos = req_total_videos

            if courseTitle and video_no:
                now_utc = datetime.now(timezone.utc)
                db.quiz_attempts.update_one(
                    {"uid": user.get("uid"), "courseTitle": courseTitle, "videoNo": int(video_no)},
                    {
                        "$setOnInsert": {
                            "uid": user.get("uid"),
                            "courseTitle": courseTitle,
                            "videoNo": int(video_no),
                            "totalVideos": int(total_videos) if total_videos else None,
                            "video_url": video_url,
                            "createdAt": now_utc,
                        },
                        "$set": {
                            "quiz_id": quiz_id,
                            "lastScore": int(score),
                            "totalQuestions": int(len(full_quiz)),
                            "required": int(required),
                            "passed": bool(passed),
                            "attemptsUsed": int(used),
                            "updatedAt": now_utc,
                            "lastAttemptAt": now_utc,
                        },
                        "$max": {"bestScore": int(score)},
                    },
                    upsert=True,
                )
        except Exception:
            pass

        # ✅ SEND QUIZ RESULT EMAIL (VIDEO = COURSE VIDEO NUMBER like 4/32, not YouTube id)
        try:
            to_email = (user.get("email") or "").strip()
            if to_email:
                status_line = "PASSED ✅" if passed else "NEEDS REATTEMPT ⚠️"
                video_url = (quiz_doc.get("video_url") or "").strip()

                # Derive video number (prefers the same logic we used above)
                # (If persistence block failed, compute here as fallback)
                try:
                    video_no = int(video_no) if "video_no" in locals() and video_no else None
                except Exception:
                    video_no = None
                try:
                    total_videos = int(total_videos) if "total_videos" in locals() and total_videos else None
                except Exception:
                    total_videos = None
                if not video_no:
                    video_no, total_videos, video_title = derive_video_meta_from_course(
                        db=db,
                        uid=user.get("uid"),
                        course_title=courseTitle,
                        video_url=video_url,
                    )

                # fallback
                youtube_id = _yt_id(video_url)

                video_label = f"#{video_no}" if video_no else (youtube_id or "N/A")
                video_label_full = (
                    f"{video_no}/{total_videos}"
                    if (video_no and total_videos)
                    else (str(video_no) if video_no else "N/A")
                )

                body = f"""
                <p style="margin:0 0 10px 0;">Hi <b>{user.get('name') or 'there'}</b>,</p>
                <p style="margin:0 0 10px 0;">Your quiz has been evaluated and your result is ready.</p>

                <ul style="margin:10px 0 10px 20px;">
                  <li><b>Course:</b> {courseTitle or 'N/A'}</li>
                  <li><b>Status:</b> <b>{status_line}</b></li>
                  <li><b>Score:</b> {score}/{len(full_quiz)} (Pass mark: {required})</li>
                  <li><b>Quiz Result Time (IST):</b> {_now_ist_str()}</li>
                  <li><b>Attempts used:</b> {used}</li>
                  <li><b>Video No:</b> {video_label_full}</li>
                  <li><b>Video:</b> {video_label}</li>
                  <li><b>Video Link:</b> {video_url or 'N/A'}</li>
                </ul>
                """

                html = _brand_email(
                    f"Quiz completed — {status_line}",
                    f"Course: {courseTitle or 'N/A'} • Score: {score}/{len(full_quiz)} • Video: {video_label}",
                    body,
                    primary_cta={"label": "View My Courses", "url": _safe_public_url("/my-courses")},
                    secondary_cta={"label": "Open Contact", "url": _safe_public_url("/contact")}
                )

                _send_email(
                    to_email,
                    f"Zenith Learning — Quiz Result ({status_line})",
                    html,
                    kind="courses"
                )
        except Exception:
            pass

        return jsonify({
            "score": score,
            "total": len(full_quiz),
            "required": required,
            "passed": passed,
            "attempts_used": used,
            "results": results
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ---------------- CHAT (CONTINUOUS) ----------------
# NOTE:
# - Persisted in MongoDB so it survives restarts.
# - Still cached in-memory for speed.
# - Stored per (uid, courseTitle, conversation_id).
CHAT_SESSIONS = {}   # session_key (uid::courseTitle::conversation_id) -> list of messages
MAX_TURNS = 20       # keep last N messages to control token usage

# ---------------- PDF CHAT (Upload + Q&A) ----------------
# NOTE:
# - PDF text/chunks are persisted in MongoDB (db.pdf_store).
# - PDF chat history is persisted in MongoDB (db.pdf_chat_sessions).
# - We still keep a small in-memory cache (PDF_STORE) per server instance.
PDF_STORE = {}  # pdf_id -> {uid, courseTitle, filename, text, chunks, created_at}
MAX_PDF_CHARS = 250_000  # safety cap for extracted text


def _pdf_cache_put(obj: dict):
    try:
        pdf_id = obj.get("pdf_id") or obj.get("id")
        if pdf_id:
            PDF_STORE[pdf_id] = obj
    except Exception:
        pass


def _pdf_get(uid: str, pdf_id: str, course_title: str = None):
    """Fetch PDF from in-memory cache first, then MongoDB."""
    if not pdf_id:
        return None

    obj = PDF_STORE.get(pdf_id)
    if obj and obj.get("uid") == uid:
        return obj

    db = get_db()
    if db is None:
        return None
    q = {"pdf_id": pdf_id, "uid": uid}
    if course_title:
        q["courseTitle"] = course_title
    doc = db.pdf_store.find_one(q, projection={"_id": 0})
    if doc:
        _pdf_cache_put(doc)
    return doc


def _chunk_text(s: str, chunk_size: int = 1200, overlap: int = 200):
    s = (s or '').strip()
    if not s:
        return []
    out=[]
    i=0
    n=len(s)
    while i < n:
        j=min(n, i+chunk_size)
        out.append(s[i:j])
        if j==n:
            break
        i=max(0, j-overlap)
    return out


def _simple_retrieve(chunks, query: str, k: int = 6):
    q = (query or '').lower()
    toks = [t for t in re.findall(r"[a-z0-9]{3,}", q)][:30]
    if not toks:
        return chunks[:k]
    scored=[]
    for c in chunks:
        cl=c.lower()
        score=sum(cl.count(t) for t in toks)
        if score>0:
            scored.append((score, c))
    scored.sort(key=lambda x: x[0], reverse=True)
    best=[c for _,c in scored[:k]]
    if len(best)<k:
        # backfill with early chunks to keep some context
        for c in chunks:
            if c not in best:
                best.append(c)
            if len(best)>=k:
                break
    return best

def _trim_history(history, max_turns=MAX_TURNS):
    # Each turn is usually user+assistant. Keep last ~max_turns*2 messages.
    if len(history) > max_turns * 2:
        return history[-max_turns * 2:]
    return history

@app.route("/chat", methods=["POST"])
def chat_bot():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    subject = (data.get("subject") or "").strip()
    user_input = (data.get("message") or "").strip()
    conversation_id = (data.get("conversation_id") or "").strip()

    # NEW: keep chat history per-course
    course_title = (data.get("courseTitle") or data.get("course") or subject or "").strip()

    if not subject or not user_input:
        return jsonify({"error": "Both 'subject' and 'message' are required"}), 400

    if not conversation_id:
        return jsonify({"error": "conversation_id is required"}), 400

    if not course_title:
        course_title = subject

    db = get_db()

    # Composite key in memory; persisted in Mongo so it survives restarts
    session_key = f"{user['uid']}::{course_title}::{conversation_id}"

    # Load history
    history = CHAT_SESSIONS.get(session_key)
    if history is None:
        history = []
        if db is not None:
            doc = db.chat_sessions.find_one(
                {"uid": user["uid"], "courseTitle": course_title, "conversation_id": conversation_id},
                sort=[("updatedAt", -1)],
                projection={"_id": 0, "history": 1}
            )
            history = (doc.get("history") if doc else []) or []

    # Add user message
    history.append({"role": "user", "content": user_input})
    history = _trim_history(history)

    system_prompt = f"""
You are a helpful tutor specializing in {subject}.
Rules:
1) If the question is NOT related to {subject}, reply exactly: Out of scope
2) Keep response clear like ChatGPT.
3) Use numbered points for explanations.
4) If code is needed, put it in a separate fenced code block using triple backticks with language.
5) Do NOT use stars (*) or dash bullets (-). Use only numbers (1., 2., 3.)
6) Be concise but helpful.
"""

    convo_text = ""
    for msg in history:
        role = msg["role"]
        content = msg["content"]
        if role == "user":
            convo_text += f"User: {content}\n"
        else:
            convo_text += f"Assistant: {content}\n"

    prompt = f"""{system_prompt}

Conversation so far:
{convo_text}

Now respond to the latest user message only.
"""

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        reply_text = (getattr(response, 'text', None) or '').strip()

        history.append({"role": "assistant", "content": reply_text})
        history = _trim_history(history)

        CHAT_SESSIONS[session_key] = history
        if db is not None:
            now = datetime.now(timezone.utc)
            db.chat_sessions.update_one(
                {"uid": user["uid"], "courseTitle": course_title, "conversation_id": conversation_id},
                {"$set": {"history": history, "updatedAt": now}, "$setOnInsert": {"createdAt": now}},
                upsert=True
            )

        return jsonify({
            "conversation_id": conversation_id,
            "courseTitle": course_title,
            "reply": reply_text,
            "history_size": len(history)
        })
    except Exception as e:
        if is_quota_error(e):
            return jsonify({"error": "Gemini quota exceeded. Please try later or enable billing."}), 429
        return jsonify({"error": str(e)}), 500


@app.route("/chat/history", methods=["GET"])
def chat_history():
    """Return stored chat history for a specific course + conversation."""
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    course_title = (request.args.get("courseTitle") or request.args.get("course") or "").strip()
    conversation_id = (request.args.get("conversation_id") or "").strip()
    if not course_title or not conversation_id:
        return jsonify({"error": "courseTitle and conversation_id are required"}), 400

    session_key = f"{user['uid']}::{course_title}::{conversation_id}"
    history = CHAT_SESSIONS.get(session_key)

    db = get_db()
    if history is None and db is not None:
        doc = db.chat_sessions.find_one(
            {"uid": user["uid"], "courseTitle": course_title, "conversation_id": conversation_id},
            projection={"_id": 0, "history": 1}
        ) or {}
        history = (doc.get("history") or [])
        CHAT_SESSIONS[session_key] = history

    return jsonify({
        "courseTitle": course_title,
        "conversation_id": conversation_id,
        "history": history or []
    })


def _delete_pair_from_history(history: list, pair_index: int) -> list:
    """Treat one Q+A as a pair (user then assistant). Delete by pair index (0-based)."""
    if not isinstance(history, list):
        return []
    try:
        pi = int(pair_index)
    except Exception:
        return history
    if pi < 0:
        return history

    i = pi * 2
    if i >= len(history):
        return history

    new_hist = list(history)
    del new_hist[i : min(len(new_hist), i + 2)]
    return new_hist


@app.route("/chat/history/pair", methods=["DELETE"])
def chat_history_delete_pair():
    """Delete a Q+A pair (one set) for a course chat."""
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    course_title = (data.get("courseTitle") or data.get("course") or "").strip()
    conversation_id = (data.get("conversation_id") or "").strip()
    pair_index = data.get("pair_index")

    if not course_title or not conversation_id or pair_index is None:
        return jsonify({"error": "courseTitle, conversation_id, and pair_index are required"}), 400

    session_key = f"{user['uid']}::{course_title}::{conversation_id}"
    db = get_db()

    history = CHAT_SESSIONS.get(session_key)
    if history is None and db is not None:
        doc = db.chat_sessions.find_one(
            {"uid": user["uid"], "courseTitle": course_title, "conversation_id": conversation_id},
            projection={"_id": 0, "history": 1},
        ) or {}
        history = (doc.get("history") or [])

    history = _delete_pair_from_history(history or [], pair_index)
    CHAT_SESSIONS[session_key] = history

    if db is not None:
        now = datetime.now(timezone.utc)
        db.chat_sessions.update_one(
            {"uid": user["uid"], "courseTitle": course_title, "conversation_id": conversation_id},
            {"$set": {"history": history, "updatedAt": now}},
            upsert=True,
        )

    return jsonify({
        "courseTitle": course_title,
        "conversation_id": conversation_id,
        "history": history,
    })

@app.route("/pdf/upload", methods=["POST"])
def pdf_upload():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    if "file" not in request.files:
        return jsonify({"error": "No file provided (field name: file)"}), 400

    f = request.files["file"]
    # Optional: store per-course
    course_title = (request.form.get("courseTitle") or request.form.get("course") or "").strip() or "Course"
    filename = (getattr(f, 'filename', '') or '').strip() or 'document.pdf'
    if not filename.lower().endswith('.pdf'):
        return jsonify({"error": "Only PDF files are supported"}), 400

    try:
        reader = PdfReader(f)
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ''
            except Exception:
                t = ''
            if t:
                parts.append(t)

        text = "\n".join(parts).strip()
        if not text:
            return jsonify({"error": "Could not extract text from this PDF (it may be scanned). Try a text-based PDF."}), 400

        if len(text) > MAX_PDF_CHARS:
            text = text[:MAX_PDF_CHARS]

        chunks = _chunk_text(text)
        pdf_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc)
        doc = {
            "pdf_id": pdf_id,
            "uid": user["uid"],
            "courseTitle": course_title,
            "filename": filename,
            "text": text,
            "chunks": chunks,
            "createdAt": now,
            "updatedAt": now,
        }

        # cache + persist
        PDF_STORE[pdf_id] = {
            **doc,
            "created_at": now.isoformat(),
        }
        db = get_db()
        if db is not None:
            db.pdf_store.insert_one(doc)

        return jsonify({"pdf_id": pdf_id, "filename": filename, "chunks": len(chunks), "courseTitle": course_title})
    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route("/pdf/list", methods=["GET"])
def pdf_list():
    """List PDFs uploaded by the user for a specific course (used to restore multi-PDF history)."""
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    course_title = (request.args.get("courseTitle") or request.args.get("course") or "").strip() or "Course"

    db = get_db()
    pdfs = []
    if db is not None:
        cur = db.pdf_store.find(
            {"uid": user["uid"], "courseTitle": course_title},
            projection={"_id": 0, "pdf_id": 1, "filename": 1, "createdAt": 1},
        ).sort("createdAt", -1)
        for d in cur:
            pid = d.get("pdf_id")
            if not pid:
                continue
            pdfs.append({"id": pid, "name": d.get("filename") or "document.pdf"})

    return jsonify({"courseTitle": course_title, "pdfs": pdfs})


@app.route("/pdf/chat", methods=["POST"])
def pdf_chat():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    pdf_id = (data.get("pdf_id") or '').strip()
    question = (data.get("question") or '').strip()
    conversation_id = (data.get("conversation_id") or '').strip()
    course_title = (data.get("courseTitle") or data.get("course") or '').strip() or "Course"

    if not pdf_id or not question:
        return jsonify({"error": "pdf_id and question are required"}), 400
    if not conversation_id:
        return jsonify({"error": "conversation_id is required"}), 400

    obj = _pdf_get(user["uid"], pdf_id, course_title)
    if not obj or obj.get("uid") != user["uid"]:
        return jsonify({"error": "PDF not found (upload again)"}), 404


    # Enforce PDF belongs to this course
    if (obj.get("courseTitle") or "Course") != (course_title or "Course"):
        return jsonify({"error": "PDF does not belong to this course"}), 400



    chunks = obj.get("chunks") or []
    best = _simple_retrieve(chunks, question, k=6)
    context = "\n\n---\n\n".join(best)

    system_prompt = """
You are a helpful assistant.
You must answer ONLY using the provided PDF excerpts.
If the answer is not in the excerpts, reply exactly: Not found in the document
Rules:
1) Keep response clear like ChatGPT.
2) Use numbered points only (1., 2., 3.).
3) If code is needed, put it in a separate fenced code block using triple backticks with language.
4) Do NOT use stars (*) or dash bullets (-).
"""

    prompt = f"""{system_prompt}

PDF filename: {obj.get('filename')}

PDF excerpts:
{context}

User question:
{question}
"""

    # Persist per-pdf chat history
    db = get_db()
    session_key = f"{user['uid']}::{course_title}::{pdf_id}::{conversation_id}"
    history = []
    if db is not None:
        doc = db.pdf_chat_sessions.find_one(
            {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id, "conversation_id": conversation_id},
            projection={"_id": 0, "history": 1}
        ) or {}
        history = (doc.get("history") or [])
    history.append({"role": "user", "content": question})
    history = _trim_history(history)

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        reply_text = (getattr(response, 'text', None) or '').strip()

        history.append({"role": "assistant", "content": reply_text})
        history = _trim_history(history)
        if db is not None:
            now = datetime.now(timezone.utc)
            db.pdf_chat_sessions.update_one(
                {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id, "conversation_id": conversation_id},
                {"$set": {"history": history, "updatedAt": now}, "$setOnInsert": {"createdAt": now}},
                upsert=True
            )

        return jsonify({
            "reply": reply_text,
            "pdf_id": pdf_id,
            "courseTitle": course_title,
            "conversation_id": conversation_id,
            "history_size": len(history)
        })
    except Exception as e:
        if is_quota_error(e):
            return jsonify({"error": "Gemini quota exceeded. Please try later or enable billing."}), 429
        return jsonify({"error": str(e)}), 500


@app.route("/pdf/chat/history", methods=["GET"])
def pdf_chat_history():
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    pdf_id = (request.args.get("pdf_id") or '').strip()
    conversation_id = (request.args.get("conversation_id") or '').strip()
    course_title = (request.args.get("courseTitle") or request.args.get("course") or '').strip() or "Course"
    if not pdf_id or not conversation_id:
        return jsonify({"error": "pdf_id and conversation_id are required"}), 400

    db = get_db()
    history = []
    if db is not None:
        doc = db.pdf_chat_sessions.find_one(
            {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id, "conversation_id": conversation_id},
            projection={"_id": 0, "history": 1}
        ) or {}
        history = (doc.get("history") or [])

    return jsonify({
        "pdf_id": pdf_id,
        "courseTitle": course_title,
        "conversation_id": conversation_id,
        "history": history or []
    })


@app.route("/pdf/chat/history/pair", methods=["DELETE"])
def pdf_chat_history_delete_pair():
    """Delete a Q+A pair (one set) for a PDF chat."""
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    pdf_id = (data.get("pdf_id") or "").strip()
    course_title = (data.get("courseTitle") or data.get("course") or "").strip() or "Course"
    conversation_id = (data.get("conversation_id") or "").strip()
    pair_index = data.get("pair_index")

    if not pdf_id or not conversation_id or pair_index is None:
        return jsonify({"error": "pdf_id, conversation_id, and pair_index are required"}), 400

    db = get_db()
    history = []
    if db is not None:
        doc = db.pdf_chat_sessions.find_one(
            {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id, "conversation_id": conversation_id},
            projection={"_id": 0, "history": 1},
        ) or {}
        history = (doc.get("history") or [])

    history = _delete_pair_from_history(history or [], pair_index)

    if db is not None:
        now = datetime.now(timezone.utc)
        db.pdf_chat_sessions.update_one(
            {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id, "conversation_id": conversation_id},
            {"$set": {"history": history, "updatedAt": now}},
            upsert=True,
        )

    return jsonify({
        "pdf_id": pdf_id,
        "courseTitle": course_title,
        "conversation_id": conversation_id,
        "history": history,
    })




@app.route("/pdf/delete", methods=["POST"])
def pdf_delete():
    """Delete a PDF (pdf_store) and all its chat sessions (pdf_chat_sessions / pdf_chats) for this user + course."""
    user, err = require_user()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    data = request.get_json() or {}
    pdf_id = (data.get("pdf_id") or "").strip()
    course_title = (data.get("courseTitle") or data.get("course") or "").strip() or "Course"

    if not pdf_id:
        return jsonify({"error": "pdf_id is required"}), 400

    db = get_db()
    if db is None:
        return jsonify({"error": "Database not available"}), 500

    q = {"uid": user["uid"], "courseTitle": course_title, "pdf_id": pdf_id}

    # 1) Delete PDF content/metadata
    r_store = db.pdf_store.delete_one(q)

    # 2) Delete all chat sessions for this PDF
    r_sessions = db.pdf_chat_sessions.delete_many(q)

    # 3) Backward-compat: if an older collection name exists, delete there too
    deleted_pdf_chats = 0
    try:
        r_legacy = db["pdf_chats"].delete_many(q)
        deleted_pdf_chats = getattr(r_legacy, "deleted_count", 0) or 0
    except Exception:
        pass

    # 4) Remove from in-memory cache
    try:
        if pdf_id in PDF_STORE:
            del PDF_STORE[pdf_id]
    except Exception:
        pass

    return jsonify({
        "ok": True,
        "pdf_id": pdf_id,
        "courseTitle": course_title,
        "deleted": {
            "pdf_store": getattr(r_store, "deleted_count", 0) or 0,
            "pdf_chat_sessions": getattr(r_sessions, "deleted_count", 0) or 0,
            "pdf_chats": deleted_pdf_chats,
    }
    })

# ===================== RESUME-BASED MOCK INTERVIEW (InterviewSense-style) =====================
# Collections:
# - interview_sessions: per-user interview chat sessions with resume context
# Notes:
# - Uses existing Firebase auth (require_user) and Gemini (get_gemini_response)
# - Stores resume text + chat history in MongoDB for persistence

def _extract_pdf_text_bytes(pdf_bytes: bytes, max_chars: int = 250_000) -> str:
    """Extract text from PDF bytes using PyPDF2; cap size to avoid huge prompts."""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if t:
                parts.append(t)
            if sum(len(p) for p in parts) > max_chars:
                break
        text = "\n".join(parts).strip()
        return text[:max_chars]
    except Exception:
        return ""



def _is_docx_filename(name: str) -> bool:
    name = (name or "").lower()
    return name.endswith(".docx") or name.endswith(".docm")


def _extract_docx_text_bytes(docx_bytes: bytes, max_chars: int = 250_000) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
            if sum(len(x) for x in parts) > max_chars:
                break
        text = "\n".join(parts).strip()
        return text[:max_chars]
    except Exception:
        return ""


def _extract_resume_text(resume_file):
    """Return (resume_text, kind, raw_bytes) where kind in {pdf, docx, unknown}."""
    filename = getattr(resume_file, "filename", "") or ""
    ctype = getattr(resume_file, "content_type", "") or ""
    data = resume_file.read() or b""

    kind = "unknown"
    text = ""
    if "pdf" in ctype or filename.lower().endswith(".pdf"):
        kind = "pdf"
        text = _extract_pdf_text_bytes(data)
    elif "word" in ctype or _is_docx_filename(filename):
        kind = "docx"
        text = _extract_docx_text_bytes(data)
    else:
        text = _extract_pdf_text_bytes(data)
        if text:
            kind = "pdf"
        else:
            text = _extract_docx_text_bytes(data)
            if text:
                kind = "docx"
    return (text or ""), kind, data


def _normalize_keyword(k: str) -> str:
    k = (k or "").strip().lower()
    k = re.sub(r"[^a-z0-9\+\#\.\- ]+", " ", k)
    k = re.sub(r"\s+", " ", k).strip()
    return k


def _filter_unwanted_keywords(keys):
    """Remove noisy tokens: names, emails, phones, urls, pure numbers, etc."""
    out = []
    seen = set()
    for k in (keys or []):
        kk = _normalize_keyword(k)
        if not kk:
            continue
        if len(kk) <= 2:
            continue
        if re.fullmatch(r"\d{3,}", kk):
            continue
        if "@" in kk or "http" in kk or "www" in kk:
            continue
        if re.search(r"\b(gmail|yahoo|outlook|linkedin|github)\b", kk):
            continue
        if kk in {"vinoth", "kumar", "t s", "ts"}:
            continue
        if kk in seen:
            continue
        seen.add(kk)
        out.append(k.strip())
    return out


def _docx_replace_section(doc: Document, heading_patterns, new_paras, bullet: bool = False):
    """Replace content after a heading until the next heading-like paragraph, preserving template."""
    paras = doc.paragraphs
    head_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        for hp in (heading_patterns or []):
            if re.fullmatch(hp, t, flags=re.I) or re.search(hp, t, flags=re.I):
                head_idx = i
                break
        if head_idx is not None:
            break
    if head_idx is None:
        return False

    def is_next_heading(text: str) -> bool:
        tt = (text or "").strip()
        if not tt:
            return False
        if len(tt) <= 40 and tt.upper() == tt and re.search(r"[A-Z]", tt):
            return True
        if re.fullmatch(r"(summary|objective|skills|experience|projects|education|certifications|achievements)", tt, flags=re.I):
            return True
        return False

    start = head_idx + 1
    end = start
    while end < len(paras) and not is_next_heading(paras[end].text):
        end += 1

    base_style = None
    for j in range(start, end):
        if (paras[j].text or "").strip():
            base_style = paras[j].style
            break

    for j in range(start, end):
        paras[j].text = ""

    lines = [ln.strip() for ln in (new_paras or []) if (ln or "").strip()]
    if not lines:
        return True

    target = start
    for ln in lines:
        if target < end:
            paras[target].text = ln
            if base_style is not None:
                try: paras[target].style = base_style
                except Exception: pass
            if bullet:
                try: paras[target].style = "List Bullet"
                except Exception: pass
            target += 1
        else:
            newp = doc.add_paragraph(ln)
            if base_style is not None:
                try: newp.style = base_style
                except Exception: pass
            if bullet:
                try: newp.style = "List Bullet"
                except Exception: pass
    return True


def _pdf_bytes_to_docx_document(pdf_bytes: bytes):
    """Best-effort PDF -> DOCX conversion using pdf2docx.
    Returns a python-docx Document if conversion succeeds, else None.
    NOTE: Only text-based PDFs convert well; scanned PDFs may lose layout.
    """
    if not pdf_bytes:
        return None
    try:
        from pdf2docx import Converter  # type: ignore
    except Exception:
        return None

    pdf_path = None
    docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as fpdf:
            fpdf.write(pdf_bytes)
            pdf_path = fpdf.name
        docx_path = pdf_path + ".docx"

        cv = Converter(pdf_path)
        try:
            cv.convert(docx_path, start=0, end=None)
        finally:
            try:
                cv.close()
            except Exception:
                pass

        with open(docx_path, "rb") as fdocx:
            docx_bytes = fdocx.read()
        return Document(io.BytesIO(docx_bytes))
    except Exception:
        return None
    finally:
        for p in (docx_path, pdf_path):
            try:
                if p and os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass


def _build_gemini_tailor_prompt(jd_struct: dict, resume_struct: dict, ats_struct: dict, jd_text: str, resume_text: str) -> str:
    """Ask Gemini for ATS scoring + section-wise tailored rewrite (strict JSON).

    NOTE: This function avoids Python f-string brace issues by building JSON schema via json.dumps.
    """
    schema = {
        "ats_score": 0,
        "reasoning_summary": "",
        "missing_skills_priority": [],
        "improvements": [],
        "tailored_sections": {
            "professional_summary": [],
            "skills": [],
            "experience_bullets": [],
            "projects_bullets": [],
            "education_lines": []
        },
        "tailored_diff": [
            {
                "section": "Professional Summary | Skills | Experience | Projects | Education",
                "replace_instruction": "Tell the user which section to replace in the resume.",
                "old_content": "",
                "new_content": ""
            }
        ],
        "missing_requirements": [
            {"type": "education|work_experience|achievements|other", "requirement": "", "details": ""}
        ]
    }
    schema_json = json.dumps(schema, ensure_ascii=False, indent=2)
    parts = [
        "You are an ATS evaluator and resume tailoring assistant.",
        "",
        "You will be given:",
        "1) Parsed JD structure (skills/phrases)",
        "2) Parsed Resume structure (skills, projects, experience, education, years estimate)",
        "3) A preliminary ATS struct (matched/missing)",
        "4) Raw JD text and Resume text (truncated)",
        "",
        "Your job:",
        "A) Provide an ATS score 0-100 with a brief reasoning summary.",
        "B) Provide 'missing_skills_priority' (top 10) and 'improvements' (bullet list).",
        "C) Provide 'tailored_diff' as OLD → NEW rewrites so we can keep the user's original DOCX template.",
        "D) Provide 'missing_requirements' for non-skill requirements (education/years/achievements) using semantic understanding (NOT word matching).",
        "",
        "STRICT RULES:",
        "- Do NOT output names, phone numbers, emails, URLs as keywords.",
        "- Only treat true technical/business skills and relevant domain terms as keywords.",
        "- Do NOT hallucinate certifications, companies, or degrees not present in RESUME_STRUCT.",
        "- If resume already satisfies a requirement, do NOT mark it missing.",
        "",
        "Respond as VALID JSON ONLY with this schema:",
        schema_json,
        "",
        "JD_STRUCT:",
        json.dumps(jd_struct, ensure_ascii=False)[:12000],
        "",
        "RESUME_STRUCT:",
        json.dumps(resume_struct, ensure_ascii=False)[:12000],
        "",
        "ATS_STRUCT:",
        json.dumps(ats_struct, ensure_ascii=False)[:12000],
        "",
        "JD_TEXT:",
        (jd_text or "")[:12000],
        "",
        "RESUME_TEXT:",
        (resume_text or "")[:20000],
    ]
    return "\n".join(parts).strip()
def _interview_sessions_col(db):
    return db["interview_sessions"]

def _interview_session_public(doc):
    if not doc:
        return None
    report = doc.get("report") if isinstance(doc.get("report"), dict) else None
    overall = None
    try:
        overall = int(report.get("overall_score")) if report else None
    except Exception:
        overall = None
    return {
        "session_id": doc.get("session_id"),
        "created_at": (doc.get("created_at") or datetime.utcnow()).isoformat(),
        "updated_at": (doc.get("updated_at") or doc.get("created_at") or datetime.utcnow()).isoformat(),
        "role_target": doc.get("role_target") or "",
        "resume_uploaded": bool(doc.get("resume_text")),
        "message_count": int(doc.get("message_count") or 0),
        "resume_filename": doc.get("resume_filename") or "",
        "jd_title": doc.get("jd_title") or "",
        "jd_source": doc.get("jd_source") or "",
        "has_report": bool(report),
        "overall_score": overall,
    }

@app.post("/interview/session/start")
def interview_start_session():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]

    uid = user.get("uid")
    payload = request.get_json(silent=True) or {}
    t0 = time.time()
    role_target = (payload.get("role_target") or "").strip()

    db = get_db()
    col = _interview_sessions_col(db)

    session_id = str(uuid.uuid4())
    doc = {
        "session_id": session_id,
        "uid": uid,
        "created_at": datetime.utcnow(),
        "role_target": role_target,
        "resume_filename": "",
        "resume_text": "",
        "chat_history": [],  # [{role:"candidate"|"hr", content:"...", ts:iso}]
        "message_count": 0,
        "updated_at": datetime.utcnow(),
    }
    col.insert_one(doc)
    return jsonify({"ok": True, "session": _interview_session_public(doc)})

@app.get("/interview/sessions")
def interview_list_sessions():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _interview_sessions_col(db)
    items = list(col.find({"uid": uid}).sort("created_at", -1).limit(50))
    return jsonify({"ok": True, "sessions": [_interview_session_public(d) for d in items]})

@app.get("/interview/sessions/<session_id>/history")
def interview_get_history(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _interview_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({
        "ok": True,
        "session": _interview_session_public(doc),
        "chat_history": doc.get("chat_history", []),
    })

@app.post("/interview/sessions/<session_id>/reset")
def interview_reset_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _interview_sessions_col(db)
    r = col.update_one(
        {"uid": uid, "session_id": session_id},
        {"$set": {"chat_history": [], "message_count": 0, "updated_at": datetime.utcnow()}}
    )
    if r.matched_count == 0:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({"ok": True})

@app.delete("/interview/sessions/<session_id>")
def interview_delete_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _interview_sessions_col(db)
    r = col.delete_one({"uid": uid, "session_id": session_id})
    if r.deleted_count == 0:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({"ok": True})

@app.post("/interview/resume/upload")
def interview_upload_resume():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")

    # multipart/form-data: file + optional session_id + role_target
    if "file" not in request.files:
        return jsonify({"error": "Missing file"}), 400

    file = request.files["file"]
    filename = (file.filename or "").strip()
    if not filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF resumes are supported"}), 400

    session_id = (request.form.get("session_id") or "").strip() or str(uuid.uuid4())
    role_target = (request.form.get("role_target") or "").strip()

    pdf_bytes = file.read() or b""
    resume_text = _extract_pdf_text_bytes(pdf_bytes)
    if not resume_text:
        return jsonify({"error": "Could not extract text from this PDF. Try a text-based PDF (not scanned)."}), 400

    db = get_db()
    col = _interview_sessions_col(db)

    now = datetime.utcnow()
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        # create session
        doc = {
            "session_id": session_id,
            "uid": uid,
            "created_at": now,
            "role_target": role_target,
            "resume_filename": filename,
            "resume_text": resume_text,
            "chat_history": [],
            "message_count": 0,
            "updated_at": now,
        }
        col.insert_one(doc)
    else:
        # update session
        update = {
            "resume_filename": filename,
            "resume_text": resume_text,
            "updated_at": now,
        }
        if role_target:
            update["role_target"] = role_target
        col.update_one({"_id": doc["_id"]}, {"$set": update})
        doc.update(update)

    return jsonify({
        "ok": True,
        "message": "Resume uploaded and processed",
        "session": _interview_session_public(doc),
        "resume_chars": len(resume_text),
    })

# ---------------- MOCK INTERVIEW: JD + REPORT ----------------
_INTERVIEW_PREDEFINED_ROLES = [
    {"key": "custom", "title": "Custom Job Description"},
    {"key": "business_analyst", "title": "Business Analyst"},
    {"key": "product_manager", "title": "Product Manager"},
    {"key": "software_engineer", "title": "Software Engineer"},
    {"key": "marketing_specialist", "title": "Marketing Specialist"},
    {"key": "data_analyst", "title": "Data Analyst"},
    {"key": "customer_service_rep", "title": "Customer Service Representative"},
    {"key": "sales_rep", "title": "Sales Representative"},
    {"key": "hr_specialist", "title": "Human Resources Specialist"},
    {"key": "ux_ui_designer", "title": "UX/UI Designer"},
    {"key": "qa_engineer", "title": "QA Engineer"},
]

def _role_title_from_key(k: str) -> str:
    k = (k or "").strip().lower()
    for r in _INTERVIEW_PREDEFINED_ROLES:
        if r["key"] == k:
            return r["title"]
    return ""

def _generate_job_description(role_title: str) -> str:
    role_title = (role_title or "").strip()
    if not role_title:
        return ""
    prompt = f"""Create a detailed Job Description for the role: {role_title}.
Include:
- Job Title
- Role Summary (2-4 sentences)
- Responsibilities (8-12 bullet points)
- Requirements (8-12 bullet points)
- Nice-to-haves (4-6 bullet points)
Keep it realistic and ATS-friendly. Limit to ~4500 characters."""
    try:
        return get_gemini_response(prompt)
    except Exception:
        return ""

def _build_interview_prompt(resume_text: str, jd_text: str, role_target: str, chat_history: list, user_message: str) -> str:
    resume_text = (resume_text or "")[:30_000]
    jd_text = (jd_text or "")[:18_000]

    turns = (chat_history or [])[-14:]
    hist_lines = []
    for t in turns:
        r = (t.get("role") or "").lower()
        c = (t.get("content") or "").strip()
        if not c:
            continue
        label = "Candidate" if r == "candidate" else "HR"
        hist_lines.append(f"{label}: {c}")
    history_block = "\n".join(hist_lines).strip()
    role_line = f"Target role: {role_target}\n" if role_target else ""

    return f"""You are an expert HR + technical interviewer.
Your job: conduct a realistic mock interview for the given Job Description (JD), using the candidate's resume.
Rules:
- Ask ONE strong question at a time.
- Mix behavioral + technical + project deep-dives aligned to the JD.
- If the answer is weak/unclear, ask a follow-up.
- Keep responses concise (2-6 sentences). No long essays.
- Do NOT invent resume details; only use what is present in the resume text.
- If the candidate asks for feedback, give actionable feedback and a better sample answer.

{role_line}
JOB DESCRIPTION (excerpt):
{jd_text}

RESUME (excerpt):
{resume_text}

CHAT SO FAR:
{history_block}

Candidate's latest message:
{user_message}

Now respond as HR interviewer.
""".strip()

def _generate_first_question(resume_text: str, jd_text: str, role_target: str) -> str:
    prompt = f"""You are an interviewer. Based on the Job Description and Resume below, ask the FIRST interview question.
Rules:
- Ask exactly ONE question.
- Prefer a strong opening question referencing either a key JD responsibility or a resume project.
- Keep it under 2 sentences.
Target role (optional): {role_target}

JOB DESCRIPTION:
{(jd_text or '')[:12000]}

RESUME:
{(resume_text or '')[:20000]}
"""
    try:
        out = get_gemini_response(prompt).strip()
        out = re.sub(r"^\s*(\d+\.|[-*])\s*", "", out).strip()
        return out or "Tell me about yourself and walk me through the most relevant parts of your resume for this role."
    except Exception:
        return "Tell me about yourself and walk me through the most relevant parts of your resume for this role."

@app.get("/interview/jd/templates")
def interview_jd_templates():
    # Public list (no auth needed)
    return jsonify({"ok": True, "roles": _INTERVIEW_PREDEFINED_ROLES})

@app.post("/interview/session/create")
def interview_create_session():
    """Create a session with JD + Resume in one shot (multipart/form-data)."""
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")

    # multipart: resume_file (required), jd_mode, role_key, jd_text OR jd_file, role_target
    resume_file = request.files.get("resume_file")
    if not resume_file:
        return jsonify({"error": "resume_file is required (PDF)"}), 400
    resume_filename = (resume_file.filename or "").strip()
    if not resume_filename.lower().endswith(".pdf"):
        return jsonify({"error": "Resume must be a PDF"}), 400
    resume_text = _extract_pdf_text_bytes(resume_file.read() or b"")
    if not resume_text:
        return jsonify({"error": "Could not extract text from resume PDF. Try a text-based PDF (not scanned)."}), 400

    jd_mode = (request.form.get("jd_mode") or "predefined").strip().lower()  # predefined|custom
    role_key = (request.form.get("role_key") or "").strip().lower()
    role_target = (request.form.get("role_target") or "").strip()

    jd_title = ""
    jd_source = ""
    jd_text = ""

    if jd_mode == "custom":
        jd_source = "custom"
        jd_title = "Custom Job Description"
        jd_text = (request.form.get("jd_text") or "").strip()
        jd_file = request.files.get("jd_file")
        if (not jd_text) and jd_file:
            jd_text = _extract_pdf_text_bytes(jd_file.read() or b"")
        if not jd_text:
            return jsonify({"error": "Custom JD: provide jd_text or jd_file (PDF)"}), 400
    else:
        jd_source = "predefined"
        jd_title = _role_title_from_key(role_key) or (role_target or "Job Description")
        jd_text = (request.form.get("jd_text") or "").strip()
        if not jd_text:
            jd_text = _generate_job_description(jd_title)
        if not jd_text:
            return jsonify({"error": "Failed to generate JD. Please try again or use Custom JD."}), 400

    # Create session
    db = get_db()
    col = _interview_sessions_col(db)
    session_id = str(uuid.uuid4())
    now = datetime.utcnow()

    first_q = _generate_first_question(resume_text, jd_text, role_target)

    doc = {
        "session_id": session_id,
        "uid": uid,
        "created_at": now,
        "updated_at": now,
        "role_target": role_target or jd_title,
        "resume_filename": resume_filename,
        "resume_text": resume_text,
        "jd_title": jd_title,
        "jd_source": jd_source,
        "jd_text": jd_text,
        "chat_history": [{"role": "hr", "content": first_q, "ts": now.isoformat()}],
        "message_count": 1,
        "report": None,
        "report_generated_at": None,
    }
    col.insert_one(doc)

    return jsonify({
        "ok": True,
        "session": _interview_session_public(doc),
        "first_question": first_q,
    })

def _interview_report_prompt(resume_text: str, jd_text: str, chat_history: list) -> str:
    # Keep prompt size controlled
    resume_text = (resume_text or "")[:20_000]
    jd_text = (jd_text or "")[:12_000]
    turns = (chat_history or [])[-30:]
    convo = []
    for t in turns:
        r = (t.get("role") or "").lower()
        c = (t.get("content") or "").strip()
        if not c:
            continue
        label = "Candidate" if r == "candidate" else "Interviewer"
        convo.append(f"{label}: {c}")
    convo_text = "\n".join(convo)

    return f"""You are a strict interview evaluator.
Given the Job Description, Resume, and Interview Transcript, produce a JSON report with:
- overall_score (0-100 integer)
- strengths (array of 4-7 bullets)
- gaps (array of 4-7 bullets)
- improvements (array of 6-10 actionable bullets)
- category_scores: object with keys ["communication","technical","problem_solving","project_depth","role_fit"] each 0-100
- summary (2-4 sentences)
- next_steps (array of 3-5 bullets)

Return ONLY valid JSON.

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}

INTERVIEW TRANSCRIPT:
{convo_text}
""".strip()

@app.get("/interview/sessions/<session_id>/report")
def interview_get_report(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")

    db = get_db()
    col = _interview_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404

    # Generate if missing
    if doc.get("report") and isinstance(doc.get("report"), dict):
        return jsonify({"ok": True, "session": _interview_session_public(doc), "report": doc.get("report")})

    if not doc.get("chat_history"):
        return jsonify({"error": "No interview transcript found for this session"}), 400

    try:
        raw = get_gemini_response(_interview_report_prompt(doc.get("resume_text",""), doc.get("jd_text",""), doc.get("chat_history", [])))
        # best-effort JSON parse
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(json)?", "", raw).strip()
            raw = raw.strip("`").strip()
        report = json.loads(raw)
    except Exception as e:
        return jsonify({"error": f"Failed to generate report: {e}"}), 500

    now = datetime.utcnow()
    col.update_one({"_id": doc["_id"]}, {"$set": {"report": report, "report_generated_at": now, "updated_at": now}})
    doc["report"] = report
    doc["report_generated_at"] = now
    return jsonify({"ok": True, "session": _interview_session_public(doc), "report": report})

# ---------------- ATS / RESUME TAILORING (AFFINDA + GEMINI) ----------------
# Goal:
# - deterministically parse resume/JD into structured signals (skills/projects/experience/education)
# - compute a clean ATS score without "unwanted keywords" (names/emails/phones/etc.)
# - ask Gemini to produce improvements + a tailored resume draft based on the parsed signals

_STOPWORDS = set([
    "the","and","for","with","from","this","that","to","of","in","on","a","an","as","is","are","was","were",
    "be","by","or","at","it","we","you","your","our","their","they","he","she","them","his","her","will",
    "can","may","should","must","have","has","had","do","does","did","not","but","if","then","than","into",
    "about","over","under","within","across","using","use","used","also","etc","per","via"
])

_SECTION_ALIASES = {
    "summary": ["summary","professional summary","objective","profile"],
    "skills": ["skills","technical skills","skill set","toolbox","technologies"],
    "experience": ["experience","work experience","professional experience","employment","internships","internship"],
    "projects": ["projects","project","academic projects","personal projects"],
    "education": ["education","academics","academic background","qualifications"],
    "certifications": ["certifications","certificates","courses","training"],
}

# A lightweight lexicon of common ATS-relevant skills/tools.
# (You can extend this list anytime; it is intentionally broad but curated.)
_SKILL_LEXICON = [
    # languages
    "python","java","javascript","typescript","c","c++","c#","go","rust","kotlin","swift","sql","r","matlab","bash",
    # web
    "react","angular","vue","node.js","node","express","next.js","flask","django","spring","spring boot","rest","graphql",
    "html","css","tailwind","bootstrap",
    # data / ml
    "pandas","numpy","scikit-learn","sklearn","tensorflow","pytorch","keras","xgboost","lightgbm",
    "nlp","llm","generative ai","prompt engineering","rag","embeddings",
    "data analysis","data analytics","statistics","hypothesis testing","a/b testing","experimentation",
    "power bi","tableau","excel","spreadsheets","spss","sas",
    # cloud / devops
    "aws","gcp","azure","docker","kubernetes","git","github","ci/cd","linux",
    # db
    "mysql","postgresql","postgres","mongodb","firebase","redis","sqlite","oracle",
    # testing / qa
    "unit testing","integration testing","jest","pytest","selenium",
    # mobile
    "flutter","dart","android","ios",
    # misc
    "microservices","system design","oauth","jwt","api","agile","scrum"
]

# normalize lexicon (map token -> canonical label)
def _canon_skill(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace("nodejs", "node.js")
    s = s.replace("sklearn", "scikit-learn")
    s = re.sub(r"\s+", " ", s)
    return s

_SKILL_CANON = {_canon_skill(x): x for x in _SKILL_LEXICON}

_EMAIL_RE = re.compile(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", re.I)
_URL_RE = re.compile(r"https?://\S+|www\.\S+", re.I)
_PHONE_RE = re.compile(r"(\+?\d[\d\s\-()]{7,}\d)")

# ---------------- ATS (AFFINDA + GEMINI) HELPERS ----------------
# We use Affinda (if configured) for high-quality structured extraction,
# then Gemini for resume-specific improvements + section rewrites (Option 1: only weak/missing sections).

def _affinda_config():
    return {
        "api_key": os.environ.get("AFFINDA_API_KEY", "").strip(),
        "workspace": os.environ.get("AFFINDA_WORKSPACE", "").strip(),
        "doctype_resume": os.environ.get("AFFINDA_DOCTYPE_RESUME", "").strip(),
        "base_url": (os.environ.get("AFFINDA_BASE_URL", "https://api.affinda.com").strip().rstrip("/")),
    }

def _affinda_parse_resume(file_bytes: bytes, filename: str = "resume.pdf") -> dict:
    """Call Affinda Resume Parser (best-effort). Returns dict with keys: ok, skills(list[str]), raw(json)."""
    cfg = _affinda_config()
    if not cfg["api_key"]:
        return {"ok": False, "error": "AFFINDA_API_KEY not set"}
    try:
        url = f"{cfg['base_url']}/v2/resumes"
        files = {"file": (filename or "resume.pdf", file_bytes)}
        data = {}
        if cfg["workspace"]:
            data["workspace"] = cfg["workspace"]
        if cfg["doctype_resume"]:
            data["document_type"] = cfg["doctype_resume"]

        # Affinda commonly accepts either Bearer or Token auth; try Bearer first, then Token.
        headers_list = [
            {"Authorization": f"Bearer {cfg['api_key']}"},
            {"Authorization": f"Token {cfg['api_key']}"},
        ]

        last = None
        for headers in headers_list:
            try:
                resp = requests.post(url, headers=headers, files=files, data=data, timeout=45)
                last = resp
                if resp.status_code in (401, 403):
                    continue
                if resp.status_code >= 400:
                    break
                j = resp.json()
                skills = []
                try:
                    data_obj = j.get("data") or {}
                    raw_sk = data_obj.get("skills") or []
                    for s in raw_sk:
                        if isinstance(s, str):
                            skills.append(s)
                        elif isinstance(s, dict):
                            nm = s.get("name") or s.get("skill") or s.get("value")
                            if nm:
                                skills.append(str(nm))
                    skills = sorted({str(x).strip() for x in skills if str(x).strip()}, key=lambda x: x.lower())
                except Exception:
                    skills = []
                return {"ok": True, "skills": skills, "raw": j}
            except Exception as e:
                last = e
                continue
        if hasattr(last, "text"):
            return {"ok": False, "error": f"Affinda error: HTTP {getattr(last,'status_code', '?')} {last.text[:300]}"}
        return {"ok": False, "error": f"Affinda error: {last}"}
    except Exception as e:
        return {"ok": False, "error": f"Affinda exception: {e}"}

def _merge_resume_struct_with_affinda(resume_struct: dict, aff: dict) -> dict:
    if not isinstance(resume_struct, dict):
        resume_struct = {}
    if isinstance(aff, dict) and aff.get("ok") and isinstance(aff.get("skills"), list):
        merged = set(resume_struct.get("skills") or [])
        for s in aff.get("skills") or []:
            ss = str(s).strip()
            if ss:
                merged.add(ss)
        resume_struct["skills"] = sorted(merged, key=lambda x: x.lower())
        resume_struct["affinda_used"] = True
    else:
        resume_struct["affinda_used"] = False
    return resume_struct

def _pick_weak_sections(resume_sections: dict, ats_struct: dict, missing_requirements: list) -> list:
    """Option 1: only weak/missing sections. Returns list of section keys: skills, experience, projects, education."""
    resume_sections = resume_sections or {}
    missing_skills = (ats_struct or {}).get("missing_skills") or []
    out = []

    if missing_skills:
        out.append("skills")

    if (resume_sections.get("experience") or "").strip():
        out.append("experience")
    elif (resume_sections.get("projects") or "").strip():
        out.append("projects")

    try:
        for r in (missing_requirements or []):
            typ = (r.get("type") or "").lower() if isinstance(r, dict) else ""
            if "education" in typ and (resume_sections.get("education") or "").strip():
                out.append("education")
                break
    except Exception:
        pass

    uniq = []
    for k in out:
        if k not in uniq and (resume_sections.get(k) or "").strip():
            uniq.append(k)
    return uniq[:3]

def _gemini_ats_section_rewrites(jd_text: str, resume_sections: dict, missing_skills_priority: list, weak_sections: list) -> dict:
    """Ask Gemini to produce resume-specific improvements + OLD->NEW rewrites for weak sections only.

    Option 1: Only rewrite the weak sections detected by our analyzer.
    """
    sec_payload = {}
    for k in (weak_sections or []):
        sec_payload[k] = (resume_sections.get(k) or "")[:6000]

    # Give Gemini a strict schema to reduce empty / generic output.
    schema = {
        "improvements": [
            "Example: Projects section: add REST API + MongoDB keywords where relevant; your current bullets mention Flask but not API design."
        ],
        "tailored_diff": [
            {
                "section": "skills|experience|projects|education",
                "replace_instruction": "Tell the user which section to replace.",
                "old_content": "Exact old section text from RESUME_SECTIONS.",
                "new_content": "Rewritten ATS-friendly content (copy/paste)."
            }
        ]
    }
    schema_json = json.dumps(schema, ensure_ascii=False, indent=2)

    # IMPORTANT: keep this prompt deterministic & grounded in provided text
    prompt = "\n".join([
        "You are an ATS resume optimizer. Your output MUST be grounded in the provided resume text.",
        "",
        "INPUTS:",
        "1) JOB DESCRIPTION (JD)",
        "2) RESUME_SECTIONS (ONLY weak sections; exact text)",
        "3) MISSING_SKILLS_PRIORITY (top missing JD keywords)",
        "",
        "STRICT RULES:",
        "- Output VALID JSON ONLY. No markdown, no extra text.",
        "- Do NOT invent companies, degrees, certifications, job titles, years, metrics or achievements not present.",
        "- You MAY add missing keywords ONLY as 'familiarity' / 'worked with' style language, unless the resume already shows usage.",
        "- Preserve the user's tone and structure (bullets/lines).",
        "- NEW content must be meaningfully different from OLD (not copy-paste).",
        "",
        "TASK A (Improvements):",
        "- Provide 8–12 resume-specific improvements.",
        "- Each improvement MUST include a section name (Skills/Experience/Projects/Education) and mention at least one missing keyword from MISSING_SKILLS_PRIORITY.",
        "- Each improvement MUST reference what is currently written (quote a short phrase from OLD content if possible).",
        "",
        "TASK B (Tailored rewrites):",
        "- For EACH weak section in RESUME_SECTIONS, return an OLD→NEW rewrite the user can copy-paste.",
        "- Naturally incorporate up to 3–6 of the most relevant missing keywords for that section.",
        "- If a section is empty, generate a compact ATS-friendly version consistent with the resume context.",
        "",
        "Respond as JSON with EXACT schema:",
        schema_json,
        "",
        "JD:",
        (jd_text or "")[:14000],
        "",
        "MISSING_SKILLS_PRIORITY:",
        json.dumps((missing_skills_priority or [])[:22], ensure_ascii=False),
        "",
        "RESUME_SECTIONS:",
        json.dumps(sec_payload, ensure_ascii=False),
    ])

    return _gemini_json(prompt) or {}
def _clean_text_basic(text: str) -> str:
    t = (text or "")
    t = _EMAIL_RE.sub(" ", t)
    t = _URL_RE.sub(" ", t)
    t = _PHONE_RE.sub(" ", t)
    t = t.replace("\x00", " ")
    return t

def _tokenize(text: str) -> list:
    t = re.sub(r"[^a-z0-9\+\#\.\s]", " ", (text or "").lower())
    t = re.sub(r"\s+", " ", t).strip()
    if not t:
        return []
    return t.split()

def _extract_skill_mentions(text: str) -> list:
    """
    Return a de-duplicated list of skills (canonicalized) found in the text.
    Matches:
      - exact lexicon phrases (including multi-word)
      - common variants (node, node.js, scikit-learn/sklearn)
    """
    t = _clean_text_basic(text).lower()
    # phrase match for multi-word skills
    found = set()
    for canon, pretty in _SKILL_CANON.items():
        if " " in canon:
            if canon in t:
                found.add(pretty)
    # token-level match for single words (and short tokens like c++)
    toks = set(_tokenize(t))
    for canon, pretty in _SKILL_CANON.items():
        if " " in canon:
            continue
        if canon in toks:
            found.add(pretty)
    # also map "node" -> node.js if present
    if "node" in toks and ("node.js" in _SKILL_CANON):
        found.add(_SKILL_CANON["node.js"])
    # consistent ordering
    return sorted(found, key=lambda x: x.lower())

def _split_into_sections(text: str) -> dict:
    """Best-effort section split based on common resume/JD headings."""
    raw = _clean_text_basic(text)
    lines = [ln.rstrip() for ln in raw.splitlines()]

    heading_to_key = {}
    for key, aliases in _SECTION_ALIASES.items():
        for a in aliases:
            heading_to_key[a.lower()] = key

    sections = {}
    cur_key = "other"
    buf = []

    def flush():
        nonlocal buf, cur_key
        s = "\n".join(buf).strip()
        if s:
            sections[cur_key] = (sections.get(cur_key, "") + "\n" + s).strip()
        buf = []

    for ln in lines:
        norm = re.sub(r"[^a-z0-9\s]", "", ln.lower()).strip()
        if 1 <= len(norm) <= 35 and norm in heading_to_key:
            flush()
            cur_key = heading_to_key[norm]
            continue
        if ln.strip() and ln.strip().isupper():
            norm2 = re.sub(r"[^a-z0-9\s]", "", ln.lower()).strip()
            if norm2 in heading_to_key:
                flush()
                cur_key = heading_to_key[norm2]
                continue
        buf.append(ln)
    flush()
    return sections

def _estimate_experience_years(exp_text: str) -> float:
    """Estimate experience duration from year ranges in Experience section."""
    t = _clean_text_basic(exp_text)
    years = [int(y) for y in re.findall(r"\b(19\d{2}|20\d{2})\b", t)]
    years = [y for y in years if 1980 <= y <= datetime.utcnow().year + 1]
    years = sorted(set(years))
    if len(years) >= 2:
        return max(0.0, float(max(years) - min(years)))
    intern_mentions = len(re.findall(r"\bintern\b|\binternship\b", t, re.I))
    if intern_mentions:
        return 0.5 * intern_mentions
    return 0.0

def _extract_education_lines(edu_text: str) -> list:
    out = []
    for ln in (edu_text or "").splitlines():
        s = ln.strip()
        if not s:
            continue
        if re.search(r"\b(b\.?e|b\.?tech|bachelor|m\.?s|m\.?tech|master|ph\.?d|university|college)\b", s, re.I):
            out.append(s)
    return out[:12]

def _extract_project_summaries(proj_text: str) -> list:
    """Split projects into chunks and attach detected tech stack."""
    t = (proj_text or "").strip()
    if not t:
        return []
    chunks = re.split(r"\n\s*\n+", t)
    out = []
    for ch in chunks:
        s = ch.strip()
        if not s:
            continue
        skills = _extract_skill_mentions(s)
        first_line = s.splitlines()[0].strip()
        title = first_line[:120]
        out.append({"title": title, "skills": skills[:12], "snippet": s[:600]})
        if len(out) >= 6:
            break
    return out

def _parse_resume_structured(resume_text: str) -> dict:
    sections = _split_into_sections(resume_text)
    skills = set(_extract_skill_mentions(sections.get("skills","") + "\n" + resume_text))
    edu_lines = _extract_education_lines(sections.get("education",""))
    projects = _extract_project_summaries(sections.get("projects",""))
    exp_years = _estimate_experience_years(sections.get("experience",""))
    return {
        "skills": sorted(skills, key=lambda x: x.lower()),
        "experience_years_est": round(exp_years, 2),
        "education_lines": edu_lines,
        "projects": projects,
        "sections_present": sorted([k for k,v in sections.items() if (v or "").strip()]),
    }

def _parse_jd_structured(jd_text: str) -> dict:
    skills = set(_extract_skill_mentions(jd_text))
    t = _clean_text_basic(jd_text).lower()
    phrases = set()
    for m in re.finditer(r"\b([a-z][a-z0-9]+(?:\s+[a-z][a-z0-9]+){1,2})\b", t):
        ph = m.group(1).strip()
        if any(w in _STOPWORDS for w in ph.split()):
            continue
        if len(ph) < 6 or len(ph) > 35:
            continue
        if _EMAIL_RE.search(ph) or _URL_RE.search(ph) or _PHONE_RE.search(ph):
            continue
        if re.search(r"\b(analysis|analytics|testing|engineering|reporting|visualization|database|cloud|api|automation|ml|ai|data)\b", ph):
            phrases.add(ph)
    return {
        "skills": sorted(skills, key=lambda x: x.lower()),
        "skill_phrases": sorted(list(phrases))[:60],
    }

def _extract_jd_requirements(jd_text: str) -> dict:
    """Best-effort extraction of non-skill requirements from a JD."""
    t = _clean_text_basic(jd_text or "")
    low = t.lower()

    # Education requirements
    edu = set()
    edu_patterns = [
        r"\b(bachelor(?:'s)?|b\.?\s*e\.?|b\.?\s*tech|btech|b\.\s*sc|bsc)\b",
        r"\b(master(?:'s)?|m\.?\s*e\.?|m\.?\s*tech|mtech|m\.\s*sc|msc|mba)\b",
        r"\b(ph\.?d|doctorate)\b",
        r"\b(computer science|information technology|electronics|electrical|statistics|mathematics|data science)\b",
    ]
    for pat in edu_patterns:
        for m in re.finditer(pat, low, flags=re.IGNORECASE):
            edu.add(m.group(0).strip())

    # Experience requirement (years)
    years_min = None
    for m in re.finditer(r"\b(\d{1,2})\s*\+?\s*(?:years|yrs)\b", low):
        try:
            y = int(m.group(1))
            years_min = y if years_min is None else min(years_min, y)
        except Exception:
            pass

    # Certifications / achievements signals
    certs = set()
    for m in re.finditer(r"\b(aws|azure|gcp)\s+(certified|certification)\b", low):
        certs.add(m.group(0).strip())
    for m in re.finditer(r"\b(certification|certified)\b[^\n\.]{0,60}", low):
        frag = m.group(0).strip()
        if len(frag) >= 10:
            certs.add(frag)

    achievements = set()
    ach_patterns = [
        r"\b(award|awarded|recognition|accomplishment|achievement)\b",
        r"\b(publication|published|patent)\b",
        r"\b(hackathon|winner|winning)\b",
        r"\b(leadership|led|mentored|managed)\b",
    ]
    for pat in ach_patterns:
        if re.search(pat, low):
            achievements.add(re.search(pat, low).group(0))

    # Responsibilities (non-skill) - keep short: ownership, collaboration, communication
    non_skill = set()
    for m in re.finditer(r"\b(communication|collaboration|stakeholder|ownership|leadership|mentoring)\b", low):
        non_skill.add(m.group(0))

    return {
        "education_tokens": sorted(edu),
        "min_years_experience": years_min,
        "certification_hints": sorted(certs)[:15],
        "achievement_hints": sorted(achievements),
        "non_skill_requirements": sorted(non_skill),
    }


def _compute_missing_requirements(jd_req: dict, resume_struct: dict) -> list:
    """Compare extracted JD requirements against parsed resume struct."""
    out = []

    # Education: if JD clearly asks for a degree/field and resume has no education lines
    edu_lines = resume_struct.get("education_lines") or []
    if jd_req.get("education_tokens") and not edu_lines:
        out.append({
            "type": "education",
            "requirement": "Education details mentioned in JD, but resume has no Education section detected.",
            "details": jd_req.get("education_tokens")[:10],
        })

    # Experience: compare min years
    min_years = jd_req.get("min_years_experience")
    try:
        resume_years = float(resume_struct.get("experience_years_est") or 0)
    except Exception:
        resume_years = 0.0
    if isinstance(min_years, int) and min_years > 0 and resume_years + 0.49 < float(min_years):
        out.append({
            "type": "work_experience",
            "requirement": f"JD mentions ~{min_years}+ years experience.",
            "details": {"resume_years_est": round(resume_years, 2), "jd_min_years": min_years},
        })

    # Certifications: if JD hints certifications and resume doesn't have certifications section
    sections_present = set((resume_struct.get("sections_present") or []))
    has_certs = ("certifications" in sections_present)
    if jd_req.get("certification_hints") and not has_certs:
        out.append({
            "type": "certifications",
            "requirement": "JD references certifications; resume has no Certifications section detected.",
            "details": jd_req.get("certification_hints")[:10],
        })

    # Achievements: heuristic - JD hints achievements and resume doesn't contain achievement-like words
    if jd_req.get("achievement_hints"):
        raw = " ".join([str(x) for x in edu_lines]) + " " + " ".join([str(x) for x in (resume_struct.get("projects") or [])])
        low = raw.lower()
        if not re.search(r"\b(award|awarded|winner|publication|patent|achievement|accomplishment)\b", low):
            out.append({
                "type": "achievements",
                "requirement": "JD implies measurable achievements/recognitions; resume doesn't clearly highlight them.",
                "details": jd_req.get("achievement_hints")[:10],
            })
    return out

def _compute_ats_score_structured(jd: dict, resume: dict) -> dict:
    jd_skills = set([_canon_skill(x) for x in (jd.get("skills") or [])])
    res_skills = set([_canon_skill(x) for x in (resume.get("skills") or [])])

    matched = sorted({s for s in jd_skills if s in res_skills})
    missing = sorted({s for s in jd_skills if s not in res_skills})

    coverage = (len(matched) / max(1, len(jd_skills))) if jd_skills else 0.0

    skills_score = round(60.0 * min(1.0, coverage), 1)

    exp_years = float(resume.get("experience_years_est") or 0.0)
    exp_score = 0.0
    if exp_years >= 3:
        exp_score = 20.0
    elif exp_years >= 2:
        exp_score = 14.0
    elif exp_years >= 1:
        exp_score = 8.0

    proj_count = len(resume.get("projects") or [])
    proj_score = min(10.0, 3.0 * proj_count)
    edu_score = 10.0 if (resume.get("education_lines") or []) else 5.0

    total = skills_score + exp_score + proj_score + edu_score
    score = int(round(min(100.0, total)))

    return {
        "score": score,
        "coverage": round(coverage * 100, 1),
        "matched_skills": matched[:40],
        "missing_skills": missing[:40],
        "score_breakdown": {
            "skills": skills_score,
            "experience": exp_score,
            "projects": round(proj_score, 1),
            "education": round(edu_score, 1),
        }
    }


def _compute_ats_analytics(jd_text: str, jd_struct: dict, resume_struct: dict, ats_struct: dict) -> dict:
    """Return lightweight analytics (no radar):
    - donut inputs (matched vs missing)
    - category breakdown (JD vs resume match %)
    - weighted missing skills based on JD frequency
    """
    jd_text_l = _clean_text_basic(jd_text or "").lower()

    jd_skills = [(_canon_skill(x) or "") for x in (jd_struct.get("skills") or [])]
    res_skills = set([_canon_skill(x) for x in (resume_struct.get("skills") or [])])

    # Category rules are intentionally simple & transparent.
    buckets = {
        "Programming": ["python","java","javascript","typescript","c","c++","c#","go","rust","kotlin","swift","sql","r","matlab","bash"],
        "Web": ["react","angular","vue","node.js","node","express","next.js","flask","django","spring","spring boot","rest","graphql","html","css","tailwind","bootstrap"],
        "Data/ML": ["pandas","numpy","scikit-learn","tensorflow","pytorch","keras","xgboost","lightgbm","nlp","llm","generative ai","rag","embeddings","data analysis","data analytics","statistics","hypothesis testing","a/b testing","experimentation","power bi","tableau","excel"],
        "Cloud/DevOps": ["aws","gcp","azure","docker","kubernetes","git","github","ci/cd","linux","microservices","system design","oauth","jwt"],
        "Databases": ["mysql","postgresql","postgres","mongodb","firebase","redis","sqlite","oracle"],
        "Testing/QA": ["unit testing","integration testing","jest","pytest","selenium"],
        "Mobile": ["flutter","dart","android","ios"],
    }

    inv = {}
    for cat, lst in buckets.items():
        for s in lst:
            inv[_canon_skill(s)] = cat

    total = {k: 0 for k in buckets.keys()}
    matched = {k: 0 for k in buckets.keys()}

    for s in jd_skills:
        cat = inv.get(s)
        if not cat:
            continue
        total[cat] += 1
        if s in res_skills:
            matched[cat] += 1

    category_breakdown = []
    for cat in buckets.keys():
        if total[cat] <= 0:
            continue
        mp = round(100.0 * (matched[cat] / max(1, total[cat])), 1)
        category_breakdown.append({
            "category": cat,
            "jd_count": int(total[cat]),
            "matched_count": int(matched[cat]),
            "missing_count": int(max(0, total[cat] - matched[cat])),
            "match_pct": mp,
        })

    # Weighted missing skills (based on occurrences in JD)
    missing_skills = (ats_struct.get("missing_skills") or [])
    weights = []
    for sk in missing_skills:
        canon = _canon_skill(sk)
        if not canon:
            continue
        if " " in canon or "." in canon or "+" in canon or "#" in canon:
            w = jd_text_l.count(canon)
        else:
            w = len(re.findall(r"\b" + re.escape(canon) + r"\b", jd_text_l))
        if w <= 0:
            w = 1
        weights.append({"skill": sk, "weight": int(w)})
    weights.sort(key=lambda x: (-x["weight"], x["skill"].lower()))

    matched_n = len(ats_struct.get("matched_skills") or [])
    missing_n = len(missing_skills)
    denom = max(1, matched_n + missing_n)
    match_pct = round(100.0 * (matched_n / denom), 1)

    return {
        "match_pct": match_pct,
        "matched_count": matched_n,
        "missing_count": missing_n,
        "category_breakdown": category_breakdown,
        "top_missing_weighted": weights[:12],
    }


def _gemini_json(prompt: str) -> dict:
    """Call Gemini and coerce its output into a JSON object (best-effort).

    Gemini sometimes returns markdown fences or extra commentary; we try hard to recover JSON.
    """
    try:
        raw = (get_gemini_response(prompt) or "").strip()
    except Exception:
        return {}
    if not raw:
        return {}

    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = re.sub(r"^```(json)?", "", raw, flags=re.I).strip()
        raw = raw.strip("`").strip()

    # First attempt: direct JSON parse
    try:
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            raw_obj = raw[start:end+1]
            return json.loads(raw_obj)
        return json.loads(raw)
    except Exception:
        pass

    # Second attempt: regex-extract first JSON object
    try:
        obj = _extract_json_obj(raw)
        return obj or {}
    except Exception:
        return {}


def _extract_json_obj(text: str):
    """Extract the first JSON object from Gemini output (best-effort)."""
    if not text:
        return None
    raw = str(text).strip()
    # Strip markdown fences if present
    raw = re.sub(r"```(?:json)?", "", raw, flags=re.I).replace("```", "").strip()
    m = re.search(r"\{.*\}", raw, flags=re.DOTALL)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None


def _gemini_code_analysis(problem_text: str, source_code: str, language: str) -> dict:
    """Return {timeComplexity, spaceComplexity, reason} using Gemini (best-effort)."""
    try:
        pt = (problem_text or "").strip()
        sc = (source_code or "").strip()
        if not sc:
            return {}
        prompt = f"""You are a code reviewer.

Given:
- Problem statement (may be partial): {pt[:4000]}
- Language: {language}
- Candidate solution code:
{sc[:12000]}

Task:
Estimate the asymptotic time and space complexity in Big-O and give a SHORT reason (1-2 lines).

Return STRICT JSON only:
{{
  "timeComplexity": "O(...)",
  "spaceComplexity": "O(...)",
  "reason": "short reasoning"
}}
"""
        obj = _gemini_json(prompt) or {}
        out = {
            "timeComplexity": (obj.get("timeComplexity") or obj.get("time_complexity") or "").strip(),
            "spaceComplexity": (obj.get("spaceComplexity") or obj.get("space_complexity") or "").strip(),
            "reason": (obj.get("reason") or obj.get("explanation") or "").strip(),
        }
        return {k: v for k, v in out.items() if v}
    except Exception:
        return {}


def _resume_sections_for_gemini(resume_text: str) -> dict:
    """Return section texts for Gemini semantic comparison."""
    sections = _split_into_sections(resume_text or "")
    # Normalize keys we care about
    return {
        "education_text": (sections.get("education") or "").strip(),
        "experience_text": (sections.get("experience") or "").strip(),
        "achievements_text": (sections.get("achievements") or "").strip(),
        "certifications_text": (sections.get("certifications") or "").strip(),
        "projects_text": (sections.get("projects") or "").strip(),
        "skills_text": (sections.get("skills") or "").strip(),
        "full_text": (resume_text or "").strip()[:20000],
    }


def _gemini_requirements_and_improvements(jd_text: str, resume_sections: dict) -> dict:
    """Use Gemini to semantically compare JD non-skill requirements vs resume sections.

    Returns:
      {
        "missing_requirements": [ {type, requirement, status, evidence_resume, evidence_jd, fix_suggestion}, ... ],
        "improvements": [ "..." ]
      }
    """
    prompt = f"""You are an ATS evaluator.

Goal:
1) Extract non-skill requirements from the Job Description (JD)
2) Compare them with the resume sections SEMANTICALLY (not keyword match)
3) Return missing requirements and actionable improvements

IMPORTANT RULES:
- Do NOT do keyword-only matching.
- Consider equivalent formats: "B.Tech" == "Bachelor of Technology", "BE" == "B.E."
- If resume has the requirement but in a different format, mark as PRESENT.
- If something is unclear, mark it as "uncertain" with a note.

Return STRICT JSON ONLY. No markdown.

JSON schema:
{{
  "missing_requirements": [
    {{
      "type": "education|work_experience|certification|achievement|other",
      "requirement": "string (what JD asks)",
      "status": "missing|present|uncertain",
      "evidence_resume": "string (quote/summary from resume if present else empty)",
      "evidence_jd": "string (quote/summary from JD)",
      "fix_suggestion": "string (what to add/change in resume)"
    }}
  ],
  "improvements": [
    "string actionable bullet",
    "..."
  ]
}}

RESUME EDUCATION SECTION:
{(resume_sections or {}).get("education_text","")}

RESUME EXPERIENCE SECTION:
{(resume_sections or {}).get("experience_text","")}

RESUME ACHIEVEMENTS SECTION:
{(resume_sections or {}).get("achievements_text","")}

RESUME CERTIFICATIONS SECTION:
{(resume_sections or {}).get("certifications_text","")}

RESUME PROJECTS SECTION:
{(resume_sections or {}).get("projects_text","")}

RESUME SKILLS SECTION:
{(resume_sections or {}).get("skills_text","")}

FULL RESUME (optional):
{(resume_sections or {}).get("full_text","")}

JOB DESCRIPTION:
{(jd_text or "")[:12000]}
""".strip()

    try:
        raw = (get_gemini_response(prompt) or "").strip()
    except Exception:
        raw = ""

    data = _extract_json_obj(raw) or {}

    mr = data.get("missing_requirements") or []
    # Only keep missing/uncertain for UI (present items create noise)
    filtered = []
    for r in mr:
        if not isinstance(r, dict):
            continue
        st = str(r.get("status") or "").lower().strip()
        if st in ("missing", "uncertain"):
            filtered.append({
                "type": r.get("type") or "other",
                "requirement": r.get("requirement") or "",
                "status": st or "missing",
                "evidence_resume": r.get("evidence_resume") or "",
                "evidence_jd": r.get("evidence_jd") or "",
                "fix_suggestion": r.get("fix_suggestion") or "",
            })

    improvements = data.get("improvements") or []
    improvements = [str(x).strip() for x in improvements if str(x).strip()]

    return {
        "missing_requirements": filtered,
        "improvements": improvements,
    }





def _ats_sessions_col(db):
    return db["ats_sessions"]


def _ats_session_public(doc):
    if not doc:
        return None
    return {
        "session_id": doc.get("session_id"),
        "created_at": (doc.get("created_at") or datetime.utcnow()).isoformat(),
        "updated_at": (doc.get("updated_at") or doc.get("created_at") or datetime.utcnow()).isoformat(),
        "company": doc.get("company") or "",
        "role": doc.get("role") or "",
        "title": doc.get("title") or "",
        "resume_filename": doc.get("resume_filename") or "",
        "jd_source": doc.get("jd_source") or "",
        "has_tailored": bool(doc.get("tailored_sections")),
        "ats": doc.get("ats") or {},
        "analytics": doc.get("analytics") or {},
    }

def _gemini_tailor_prompt(jd_text: str, resume_text: str, jd_struct: dict, resume_struct: dict, ats_struct: dict) -> str:
    import json
    return f"""You are an ATS scoring engine + resume tailoring assistant.

You will be given:
- Job Description (JD) text
- Resume text
- Parsed structured signals (skills/projects/experience/education)
- A deterministic ATS score + matched/missing skills computed by our parser

TASKS:
1) Use the structured signals to produce an ATS score from 0-100 (can be same as deterministic score, or adjust slightly if justified).
2) Explain score briefly (2-4 lines).
3) Provide 10-12 prioritized improvements (bullet list).
4) Provide a tailored resume draft (TEXT) that the candidate can copy-paste:
   - Professional Summary (3-4 lines)
   - Skills (grouped)
   - Experience bullets (rewrite 6-10 bullets max, STAR + metrics placeholders)
   - Projects bullets (rewrite 4-6 bullets max)
   - Keep it truthful: DO NOT invent companies, degrees, or years not in the resume text.
   - You may rephrase existing work/projects and add missing keywords ONLY if they are reasonable to claim as \"familiarity\" (not years of experience).
Return ONLY JSON with keys:
ats_score, reasoning_summary, improvements, missing_skills_priority, tailored_resume_text

Deterministic signals (JSON):
JD_STRUCT={json.dumps(jd_struct, ensure_ascii=False)[:12000]}
RESUME_STRUCT={json.dumps(resume_struct, ensure_ascii=False)[:12000]}
ATS_STRUCT={json.dumps(ats_struct, ensure_ascii=False)[:12000]}

JD_TEXT:
{(jd_text or "")[:12000]}

RESUME_TEXT:
{(resume_text or "")[:20000]}
"""

@app.post("/ats/analyze")
def ats_analyze():
    """multipart/form-data: resume_file (PDF/DOCX) + jd_file (PDF) OR jd_text.
    Also accepts: company, role (strings). Creates an ATS session stored in DB.
    """
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]

    uid = user.get("uid")

    resume_file = request.files.get("resume_file")
    if not resume_file:
        return jsonify({"error": "resume_file is required"}), 400
    resume_text, resume_kind, resume_raw = _extract_resume_text(resume_file)
    if not resume_text:
        return jsonify({"error": "Could not extract text from resume PDF"}), 400

    jd_text = (request.form.get("jd_text") or "").strip()
    jd_file = request.files.get("jd_file")
    if (not jd_text) and jd_file:
        jd_text = _extract_pdf_text_bytes(jd_file.read() or b"")
    if not jd_text:
        return jsonify({"error": "Provide jd_text or jd_file (PDF)"}), 400

    company = (request.form.get("company") or "").strip()
    role = (request.form.get("role") or "").strip()
    title = f"{company}-{role}".strip("-").strip() if (company or role) else "ATS Session"

    jd_source = "text"
    try:
        if jd_file and getattr(jd_file, "filename", ""):
            jd_source = f"file:{jd_file.filename}"
    except Exception:
        pass

    resume_struct = _parse_resume_structured(resume_text)

    # Optional: Affinda extraction (for cleaner skills/structure). Uses requests (no SDK).
    try:
        aff = _affinda_parse_resume(resume_raw, filename=(getattr(resume_file, "filename", "") or "resume.pdf")) if resume_raw else {"ok": False}
    except Exception:
        aff = {"ok": False}
    resume_struct = _merge_resume_struct_with_affinda(resume_struct, aff)

    jd_struct = _parse_jd_structured(jd_text)
    ats_struct = _compute_ats_score_structured(jd_struct, resume_struct)
    analytics = _compute_ats_analytics(jd_text, jd_struct, resume_struct, ats_struct)

    # Semantic non-skill requirements (education / experience / achievements) via Gemini
    resume_sections = _split_into_sections(resume_text)
    req_eval = _gemini_requirements_and_improvements(jd_text, _resume_sections_for_gemini(resume_text))
    missing_requirements = req_eval.get("missing_requirements") or []
    req_improvements = req_eval.get("improvements") or []

    # Missing skill priority (weighted by JD frequency if available)
    missing_skills_priority = []
    try:
        msw = (analytics or {}).get("missing_skill_weights") or []
        if msw and isinstance(msw, list):
            msw2 = []
            for it in msw:
                if isinstance(it, dict):
                    sk = it.get("skill") or it.get("name") or it.get("key")
                    wt = it.get("weight") or it.get("freq") or 1
                    if sk:
                        msw2.append((str(sk), float(wt) if isinstance(wt, (int, float)) else 1.0))
            msw2.sort(key=lambda x: x[1], reverse=True)
            missing_skills_priority = [x[0] for x in msw2[:25]]
    except Exception:
        missing_skills_priority = []
    if not missing_skills_priority:
        missing_skills_priority = (ats_struct.get("missing_skills") or [])[:25]

    # Option 1: rewrite only weak sections
    weak_section_keys = _pick_weak_sections(resume_sections, ats_struct, missing_requirements)

    # Gemini: resume-specific improvements + OLD->NEW section rewrites
    gem = _gemini_ats_section_rewrites(jd_text, resume_sections, missing_skills_priority, weak_section_keys) or {}

    tips = {}
    tips["missing_skills_priority"] = missing_skills_priority
    tips["missing_requirements"] = missing_requirements
    tips["weak_sections"] = weak_section_keys

    # Improvements (resume-specific)
    tips["improvements"] = []
    if isinstance(gem, dict) and isinstance(gem.get("improvements"), list):
        tips["improvements"] = [str(x).strip() for x in gem.get("improvements") if str(x).strip()][:20]
    if not tips["improvements"]:
        tips["improvements"] = [str(x).strip() for x in (req_improvements or []) if str(x).strip()][:20]
    if not tips["improvements"]:
        tips["improvements"] = [
            "Add missing JD keywords naturally inside your existing bullets (avoid keyword stuffing).",
            "Quantify impact in 2–4 key bullets (use concrete numbers: [X%], [N users], [latency], etc.).",
            "Strengthen the Skills section by grouping tools by category and including the top missing skills you can honestly claim as familiarity.",
        ]

    # Tailored diff (copy/paste ready)
    tips["tailored_diff"] = []
    if isinstance(gem, dict) and isinstance(gem.get("tailored_diff"), list):
        # Normalize section titles
        key_to_title = {"skills": "Skills", "experience": "Experience", "projects": "Projects", "education": "Education"}
        td = []
        for b in gem.get("tailored_diff") or []:
            if not isinstance(b, dict):
                continue
            sec = (b.get("section") or "").strip()
            # if Gemini returned key names, map them
            sec_norm = sec.lower().strip()
            sec_title = key_to_title.get(sec_norm, sec or "Section")
            oldc = (b.get("old_content") or "").strip()
            newc = (b.get("new_content") or "").strip()
            if not newc:
                continue
            td.append({
                "section": sec_title,
                "replace_instruction": b.get("replace_instruction") or f"Replace your '{sec_title}' section with the NEW content below.",
                "old_content": oldc,
                "new_content": newc,
            })
        tips["tailored_diff"] = td

    # Deterministic fallback so UI is never empty
    if not tips["tailored_diff"]:
        td = []
        if "skills" in weak_section_keys and (resume_sections.get("skills") or "").strip():
            oldc = (resume_sections.get("skills") or "").strip()
            add = ", ".join(missing_skills_priority[:10])
            newc = oldc
            if add and add.lower() not in oldc.lower():
                newc = (oldc + "\n\nAdditional keywords to consider (only if truthful): " + add).strip()
            td.append({
                "section": "Skills",
                "replace_instruction": "Update your Skills section by adding the NEW content below (only include what you can honestly claim).",
                "old_content": oldc,
                "new_content": newc,
            })
        tips["tailored_diff"] = td

    # For API backward-compatibility
    jd_requirements = []
    tailored_sections = (tips.get("tailored_sections") or {}) if isinstance(tips, dict) else {}

    # Store a reusable DOCX template for download (keeps the user's template)
    fs = gridfs.GridFS(get_db())
    template_file_id = None
    try:
        if resume_kind == "docx":
            template_bytes = resume_raw
        else:
            converted = _pdf_bytes_to_docx_document(resume_raw)
            if converted is not None:
                bio_buf = io.BytesIO()
                converted.save(bio_buf)
                template_bytes = bio_buf.getvalue()
            else:
                template_bytes = None
        if template_bytes:
            template_file_id = fs.put(template_bytes, filename=f"template_{uid}_{int(time.time())}.docx", contentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception:
        template_file_id = None

    session_id = str(uuid.uuid4())
    now = datetime.utcnow()
    doc = {
        "session_id": session_id,
        "uid": uid,
        "created_at": now,
        "updated_at": now,
        "company": company,
        "role": role,
        "title": title,
        "resume_filename": getattr(resume_file, "filename", "") or "",
        "resume_kind": resume_kind,
        "resume_text": resume_text[:80_000],
        "jd_text": jd_text[:80_000],
        "jd_source": jd_source,
        "parsed_resume": resume_struct,
        "parsed_jd": jd_struct,
        "ats": ats_struct,
        "analytics": analytics,
        "tips": tips,
        "jd_requirements": jd_requirements,
        "missing_requirements": missing_requirements,
        "tailored_sections": tailored_sections,
        "template_file_id": template_file_id,
    }
    try:
        col = _ats_sessions_col(get_db())
        col.insert_one(doc)
    except Exception:
        # If DB insert fails, still return analysis without history.
        session_id = ""


    # ✅ Email the ATS summary to the user (non-blocking)
    try:
        to_email = (user.get("email") or "").strip()
        if not to_email:
            try:
                udoc = get_db().users.find_one({"uid": uid}) if uid else None
                to_email = ((udoc or {}).get("email") or "").strip()
            except Exception:
                to_email = ""

        if to_email:
            score = (ats_struct or {}).get("score")
            coverage = (ats_struct or {}).get("coverage")
            matched = (ats_struct or {}).get("matched_skills") or []
            missing = (ats_struct or {}).get("missing_skills") or []
            miss_req = missing_requirements or []

            def _li(items):
                # Render a simple <ul> list. Items may be strings or dicts.
                out = []
                for x in (items or []):
                    try:
                        if isinstance(x, dict):
                            # Best-effort flatten dicts (avoid JSON blobs in mail)
                            parts = []
                            for k in ("name", "skill", "requirement", "text", "title"):
                                v = x.get(k)
                                if v:
                                    parts.append(str(v))
                                    break
                            if not parts:
                                kv = []
                                for k, v in list(x.items())[:6]:
                                    if v is None or v == "":
                                        continue
                                    kv.append(f"{k}: {v}")
                                parts.append("; ".join(kv) if kv else "—")
                            out.append(f"<li>{_html_escape(' '.join(parts).strip())}</li>")
                        else:
                            out.append(f"<li>{_html_escape(str(x))}</li>")
                    except Exception:
                        out.append("<li>—</li>")
                return "".join(out) or "<li>—</li>"

            def _li_missing_requirements(items):
                # Pretty formatting for Gemini requirement objects:
                # {requirement, status, fix_suggestion, ...}
                out = []
                for x in (items or []):
                    try:
                        if isinstance(x, dict):
                            req = (x.get("requirement") or x.get("text") or x.get("title") or "Requirement")
                            req = (req or "Requirement").strip()
                            status = (x.get("status") or "").strip()
                            sugg = (x.get("fix_suggestion") or x.get("suggestion") or "").strip()
                            line = f"<strong>{_html_escape(req)}</strong>"
                            if status:
                                st = status.lower()
                                if "miss" in st:
                                    color = "#dc2626"
                                elif "met" in st or "pass" in st:
                                    color = "#16a34a"
                                else:
                                    color = "#6b7280"
                                line += f' <span style="color:{color};font-weight:600;">({_html_escape(status)})</span>'
                            if sugg:
                                line += f'<div style="margin-top:4px;color:#374151;font-size:12px;line-height:1.35;">{_html_escape(sugg)}</div>'
                            out.append(f"<li style=\"margin-bottom:10px;\">{line}</li>")
                        else:
                            out.append(f"<li>{_html_escape(str(x))}</li>")
                    except Exception:
                        out.append("<li>—</li>")
                return "".join(out) or "<li>—</li>"

            # Company/Role context (shown in mail subject and header)
            cr = " — ".join([x for x in [company.strip() if isinstance(company, str) else "", role.strip() if isinstance(role, str) else ""] if x]).strip()
            mail_title = "Zenith ATS Intelligence Report" + (f" — {cr}" if cr else "")
            mail_subtitle = (
                (f"Company: {_html_escape(company)} • Role: {_html_escape(role)}<br/>" if (company or role) else "")
                + f"ATS Score: {score if score is not None else '—'}/100 • Match: {coverage if coverage is not None else '—'}%"
            )
            mail_subject = "Zenith ATS Intelligence — " + (cr if cr else "Your ATS Score & Match %")
            
            body_html = f"""
              <p style="margin:0 0 10px 0;">Here is your ATS evaluation summary.</p>
              <div style="display:flex;gap:10px;flex-wrap:wrap;margin:10px 0 16px 0;">
                <div style="padding:10px 12px;border:1px solid #e5e7eb;border-radius:14px;background:#ffffff;">
                  <div style="font-size:12px;color:#6b7280;">ATS Score</div>
                  <div style="font-size:18px;font-weight:700;color:#111827;">{score if score is not None else "—"}/100</div>
                </div>
                <div style="padding:10px 12px;border:1px solid #e5e7eb;border-radius:14px;background:#ffffff;">
                  <div style="font-size:12px;color:#6b7280;">Match</div>
                  <div style="font-size:18px;font-weight:700;color:#111827;">{coverage if coverage is not None else "—"}%</div>
                </div>
              </div>

              <h3 style="margin:14px 0 6px 0;font-size:14px;">Matched skills</h3>
              <ul style="margin:0 0 10px 18px;color:#111827;">{_li(matched[:30])}</ul>

              <h3 style="margin:14px 0 6px 0;font-size:14px;">Missing skills</h3>
              <ul style="margin:0 0 10px 18px;color:#111827;">{_li(missing[:30])}</ul>

              <h3 style="margin:14px 0 6px 0;font-size:14px;">Missing requirements (from JD)</h3>
              <ul style="margin:0 0 10px 18px;color:#111827;">{_li_missing_requirements(miss_req[:20])}</ul>

              <p style="margin:14px 0 0 0;color:#6b7280;font-size:12px;">
                Tip: Add missing keywords naturally in your projects/experience bullets and keep the resume concise (1–2 pages).
              </p>
            """

            html = _brand_email(
                title=mail_title,
                subtitle=mail_subtitle,
                body_html=body_html,
                cta_url="/ats-intelligence",
                cta_text="Open ATS Intelligence",
            )

            _send_email(
                to_email,
                mail_subject,
                html,
                kind="ats",
                reply_to=(os.getenv("CONTACT_INBOX") or os.getenv("ADMIN_EMAIL")),
            )
    except Exception as e:
        try:
            import traceback
            print("[email] ATS report send failed:", repr(e))
            traceback.print_exc()
        except Exception:
            pass
    return jsonify({
        "ok": True,
        "session": _ats_session_public(doc) if session_id else None,
        "session_id": session_id,
        "ats": ats_struct,
        "analytics": analytics,
        "parsed_resume": resume_struct,
        "parsed_jd": jd_struct,
        "tips": tips,
        "jd_requirements": jd_requirements,
        "missing_requirements": missing_requirements,
    })


@app.get("/ats/sessions")
def ats_list_sessions():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    col = _ats_sessions_col(get_db())
    items = list(col.find({"uid": uid}).sort("created_at", -1).limit(100))
    return jsonify({"ok": True, "sessions": [_ats_session_public(d) for d in items]})


@app.get("/ats/sessions/<session_id>")
def ats_get_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    col = _ats_sessions_col(get_db())
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({
        "ok": True,
        "session": _ats_session_public(doc),
        "parsed_resume": doc.get("parsed_resume") or {},
        "parsed_jd": doc.get("parsed_jd") or {},
        "ats": doc.get("ats") or {},
        "analytics": doc.get("analytics") or {},
        "tips": doc.get("tips") or {},
    })


@app.delete("/ats/sessions/<session_id>")
def ats_delete_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _ats_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404

    # delete GridFS template if present
    try:
        fid = doc.get("template_file_id")
        if fid:
            fs = gridfs.GridFS(db)
            fs.delete(fid)
    except Exception:
        pass

    col.delete_one({"uid": uid, "session_id": session_id})
    return jsonify({"ok": True})


@app.get("/ats/sessions/<session_id>/tailored_docx")
def ats_download_tailored_docx(session_id):
    """Download tailored resume DOCX from a stored ATS session (preserves template)."""
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    db = get_db()
    col = _ats_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404

    fs = gridfs.GridFS(db)
    template_bytes = None
    try:
        fid = doc.get("template_file_id")
        if fid:
            template_bytes = fs.get(fid).read()
    except Exception:
        template_bytes = None

    # Fallback if template missing
    if template_bytes:
        try:
            docx_doc = Document(io.BytesIO(template_bytes))
        except Exception:
            docx_doc = Document()
    else:
        docx_doc = Document()
        docx_doc.add_paragraph("TAILORED RESUME")
        docx_doc.add_paragraph("")

    tailored = doc.get("tailored_sections") or {}
    summary_lines = tailored.get("professional_summary") or []
    skills_lines = tailored.get("skills") or []
    exp_bullets = tailored.get("experience_bullets") or []
    proj_bullets = tailored.get("projects_bullets") or []
    edu_lines = tailored.get("education_lines") or []

    _docx_replace_section(docx_doc, [r"summary", r"professional summary", r"objective"], summary_lines, bullet=False)
    _docx_replace_section(docx_doc, [r"skills", r"technical skills"], skills_lines, bullet=True)
    _docx_replace_section(docx_doc, [r"experience", r"work experience", r"employment"], exp_bullets, bullet=True)
    _docx_replace_section(docx_doc, [r"projects", r"project experience"], proj_bullets, bullet=True)
    _docx_replace_section(docx_doc, [r"education"], edu_lines, bullet=False)

    out = io.BytesIO()
    docx_doc.save(out)
    out.seek(0)
    safe_title = (doc.get("title") or "tailored_resume").strip().replace("/", "-").replace("\\", "-")
    fname = f"{safe_title or 'tailored_resume'}.docx"
    return send_file(out, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=fname)


def _build_interview_prompt(resume_text: str, jd_text: str, role_target: str, chat_history: list, user_message: str) -> str:
    # Keep resume/JD compact for prompt
    resume_text = (resume_text or "")[:30_000]
    jd_text = (jd_text or "")[:18_000]

    # Format last few turns
    turns = (chat_history or [])[-12:]
    hist_lines = []
    for t in turns:
        r = (t.get("role") or "").lower()
        c = (t.get("content") or "").strip()
        if not c:
            continue
        label = "Candidate" if r == "candidate" else "HR"
        hist_lines.append(f"{label}: {c}")
    history_block = "\n".join(hist_lines).strip()

    role_line = f"Target role: {role_target}\n" if role_target else ""

    return f"""
You are an expert HR + technical interviewer.
Your job: conduct a realistic mock interview based on the candidate's resume AND the provided job description (JD).
Rules:
- Ask one strong question at a time.
- Mix behavioral + technical + project deep-dives aligned to the JD.
- If the candidate's answer is weak/unclear, ask a follow-up.
- Keep responses concise (2-6 sentences). No long essays.
- Do NOT invent resume details; only use what is present in the resume text.
- If the candidate asks for feedback, give actionable feedback and a better sample answer.

{role_line}
JOB DESCRIPTION (excerpt):
{jd_text}

RESUME (excerpt):
{resume_text}

CHAT SO FAR:
{history_block}

Candidate's latest message:
{user_message}

Now respond as HR interviewer.
""".strip()




@app.post("/ats/tailor_docx")
def ats_tailor_docx():
    """Return a tailored DOCX. If the uploaded resume is DOCX, preserves the original template."""
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]

    resume_file = request.files.get("resume_file")
    if not resume_file:
        return jsonify({"error": "resume_file is required"}), 400

    jd_file = request.files.get("jd_file")
    jd_text = (request.form.get("jd_text") or "").strip()

    # JD text
    if jd_file and not jd_text:
        jd_bytes = jd_file.read() or b""
        jd_text = _extract_pdf_text_bytes(jd_bytes)
    if not jd_text:
        return jsonify({"error": "Provide jd_text or jd_file"}), 400

    resume_text, resume_kind, resume_raw = _extract_resume_text(resume_file)
    if not resume_text:
        return jsonify({"error": "Could not read resume. Upload a PDF or DOCX."}), 400

    # parse + score (local algorithm)
    jd_struct = _extract_jd_struct(jd_text)
    resume_struct = _extract_resume_struct(resume_text)
    ats_struct = _compute_ats(jd_struct, resume_struct)
    ats_struct["matched_skills"] = _filter_unwanted_keywords(ats_struct.get("matched_skills", []))
    ats_struct["missing_skills"] = _filter_unwanted_keywords(ats_struct.get("missing_skills", []))

    # Gemini tailored sections
    prompt = _build_gemini_tailor_prompt(jd_struct, resume_struct, ats_struct, jd_text, resume_text)
    tips = _gemini_json(prompt)

    tailored = (tips or {}).get("tailored_sections") or {}
    summary_lines = tailored.get("professional_summary") or []
    skills_lines = tailored.get("skills") or []
    exp_bullets = tailored.get("experience_bullets") or []
    proj_bullets = tailored.get("projects_bullets") or []
    edu_lines = tailored.get("education_lines") or []

    # Build DOCX
    if resume_kind == "docx":
        try:
            doc = Document(io.BytesIO(resume_raw))
        except Exception:
            doc = Document()
    else:
        # PDF: try to convert to DOCX first (best-effort), then preserve that converted template.
        converted = _pdf_bytes_to_docx_document(resume_raw)
        if converted is not None:
            doc = converted
        else:
            # Fallback: PDF cannot preserve exact template; create a clean docx
            doc = Document()
            doc.add_paragraph("TAILORED RESUME")
            doc.add_paragraph("")

    # Replace common sections (best-effort)
    _docx_replace_section(doc, [r"summary", r"professional summary", r"objective"], summary_lines, bullet=False)
    _docx_replace_section(doc, [r"skills", r"technical skills"], skills_lines, bullet=True)
    _docx_replace_section(doc, [r"experience", r"work experience", r"employment"], exp_bullets, bullet=True)
    _docx_replace_section(doc, [r"projects", r"project experience"], proj_bullets, bullet=True)
    _docx_replace_section(doc, [r"education"], edu_lines, bullet=False)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    fname = "tailored_resume.docx"
    return send_file(out, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=fname)
@app.post("/interview/chat")
def interview_chat():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")

    payload = request.get_json(silent=True) or {}
    session_id = (payload.get("session_id") or "").strip()
    message = (payload.get("message") or "").strip()
    if not session_id:
        return jsonify({"error": "session_id is required"}), 400
    if not message:
        return jsonify({"error": "message is required"}), 400

    db = get_db()
    col = _interview_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Session not found"}), 404
    if not doc.get("resume_text"):
        return jsonify({"error": "Please upload your resume first"}), 400
    if not doc.get("jd_text"):
        return jsonify({"error": "Please set a job description (JD) first"}), 400

    prompt = _build_interview_prompt(
        doc.get("resume_text", ""),
        doc.get("jd_text", ""),
        doc.get("role_target", ""),
        doc.get("chat_history", []),
        message
    )

    try:
        reply = get_gemini_response(prompt)
    except Exception as e:
        return jsonify({"error": f"Gemini error: {e}"}), 500

    now_iso = datetime.utcnow().isoformat()
    history = doc.get("chat_history", [])
    history.append({"role": "candidate", "content": message, "ts": now_iso})
    history.append({"role": "hr", "content": reply, "ts": now_iso})

    new_count = int(doc.get("message_count") or 0) + 1
    col.update_one(
        {"_id": doc["_id"]},
        {"$set": {"chat_history": history, "message_count": new_count, "updated_at": datetime.utcnow()}}
    )

    return jsonify({"ok": True, "response": reply, "session_id": session_id, "message_count": new_count})


# ===================== MOCK TEST (General Apti / Tech Apti / Coding) =====================
# Modes:
# - general aptitude: quant + verbal/logical
# - tech aptitude: DSA/OOP/OS/CN/DBMS basics
# - coding: DSA-style coding, validated via Judge0 (Java/C/C++/Python)
#
# Storage:


# ----------------- ADMIN: USER & COURSE DETAIL (DRILLDOWN) -----------------
@app.get("/admin/user/<uid>/summary")
def admin_user_summary(uid):
    """User drilldown summary used by admin UI."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    u = db.users.find_one({"uid": uid}, projection={"_id": 0}) or {}
    if not u:
        return jsonify({"ok": False, "error": "User not found"}), 404

    # courses studying (distinct)
    studying = sorted({(x.get("courseTitle") or "").strip() for x in db.course_states.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1}) if x.get("courseTitle")})

    # holds
    holds = sorted({(x.get("courseTitle") or "").strip() for x in db.course_holds.find({"uid": uid, "held": True}, projection={"_id": 0, "courseTitle": 1}) if x.get("courseTitle")})

    # mock tests
    mt_total = db.mocktest_sessions.count_documents({"uid": uid})
    mt_submitted = db.mocktest_sessions.count_documents({"uid": uid, "submittedAt": {"$exists": True}})

    # course progress
    cp_docs = list(db.course_progress.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1, "quizPassedMap": 1, "updatedAt": 1}))

    return jsonify({
        "ok": True,
        "user": {"uid": u.get("uid"), "name": u.get("name"), "email": u.get("email"), "photoURL": u.get("photoURL"), "isAdmin": bool(u.get("isAdmin"))},
        "studyingCourses": studying,
        "heldCourses": holds,
        "mocktests": {"total": int(mt_total), "submitted": int(mt_submitted)},
        "courseProgressCount": len(cp_docs),
    })


@app.get("/admin/user/<uid>/courses-studying")
def admin_user_courses_studying(uid):
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    titles = {}
    for st in db.course_states.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1, "updatedAt": 1}):
        ct = (st.get("courseTitle") or "").strip()
        if not ct:
            continue
        prev = titles.get(ct)
        if not prev or (st.get("updatedAt") and st.get("updatedAt") > (prev.get("updatedAt") or 0)):
            titles[ct] = {"courseTitle": ct, "updatedAt": st.get("updatedAt")}

    out = []
    for ct, info in titles.items():
        out.append({
            "courseTitle": ct,
            "held": is_course_held(uid, ct),
            "updatedAt": info.get("updatedAt"),
        })

    out.sort(key=lambda x: (x.get("held") is False, x.get("updatedAt") or 0), reverse=True)
    return jsonify({"ok": True, "courses": out})


@app.get("/admin/user/<uid>/course-progress")
def admin_user_course_progress(uid):
    """Per-course quiz progress for a single user."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()

    def count_videos(obj):
        if obj is None:
            return 0
        if isinstance(obj, list):
            return sum(count_videos(it) for it in obj)
        if isinstance(obj, dict):
            if isinstance(obj.get("videos"), list):
                return len(obj.get("videos"))
            if "video" in obj and ("topic" in obj or "title" in obj):
                return 1
            return sum(count_videos(v) for v in obj.values())
        return 0

    totals = {}
    for st in db.course_states.find({}, projection={"_id": 0, "courseTitle": 1, "videos": 1}):
        ct = st.get("courseTitle")
        if not ct:
            continue
        totals[ct] = max(totals.get(ct, 0), count_videos(st.get("videos")))

    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    rows = []
    for cp in db.course_progress.find({"uid": uid}, projection={"_id": 0, "courseTitle": 1, "quizPassedMap": 1, "updatedAt": 1}):
        ct = cp.get("courseTitle")
        total = int(totals.get(ct, 0))
        passed = count_true(cp.get("quizPassedMap") or {})
        pct = int(round((passed / total) * 100)) if total else 0
        rows.append({
            "courseTitle": ct,
            "totalQuizzes": total,
            "passedQuizzes": passed,
            "percent": pct,
            "held": is_course_held(uid, ct),
            "updatedAt": cp.get("updatedAt"),
        })

    rows.sort(key=lambda x: (x.get("held") is False, x.get("percent") or 0), reverse=True)
    return jsonify({"ok": True, "rows": rows})


@app.get("/admin/user/<uid>/course/<course_title>/quiz-results")
def admin_user_course_quiz_results(uid, course_title):
    """Video-wise quiz attempt summaries for a given user + course (Admin drilldown)."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    # normalize incoming title for tolerant DB lookup
    course_title = (course_title or "").strip()
    ct_re = f"^\\s*{re.escape(course_title)}\\s*$"

    # Pull compact per-video quiz attempt docs (written by /submit-quiz)
    # Sort by latest first. Legacy docs may not have `videoNo`, so sorting by it
    # can make us miss older stored results.
    _docs = list(
        db.quiz_attempts.find(
            {"uid": uid, "courseTitle": {"$regex": ct_re, "$options": "i"}},
            projection={"_id": 0},
        ).sort([("updatedAt", -1), ("createdAt", -1)])
    )

    # If quiz_attempts is missing (older data / video number couldn't be derived),
    # we still want to show a video-wise table using course_states + course_progress.
    # That way Admin gets per-video status + link even if scores are unavailable.
    def _get_url(obj):
        if obj is None:
            return ""
        if isinstance(obj, str):
            return obj.strip()
        if isinstance(obj, dict):
            return (obj.get("videoUrl")
                    or obj.get("video_url")
                    or obj.get("url")
                    or obj.get("video")
                    or "").strip()
        return str(obj).strip()

    def _flatten_video_urls(node):
        out = []
        if node is None:
            return out
        if isinstance(node, str):
            u = node.strip()
            if u:
                out.append(u)
            return out
        if isinstance(node, list):
            for it in node:
                out.extend(_flatten_video_urls(it))
            return out
        if isinstance(node, dict):
            # common shapes
            if isinstance(node.get("videos"), list):
                for it in (node.get("videos") or []):
                    u = _get_url(it)
                    if u:
                        out.append(u)
                return out
            # week-style dict: {"Week 1": [ ... ]}
            for v in node.values():
                out.extend(_flatten_video_urls(v))
            return out
        return out

    def _flatten_video_entries(node):
        """Return list of {title,url} in order, de-duped by url."""
        out = []

        def _push(title, url):
            u = (url or "").strip()
            if not u:
                return
            t = (title or "").strip() or None
            out.append({"title": t, "url": u})

        def _walk(n):
            if n is None:
                return
            if isinstance(n, str):
                _push(None, n)
                return
            if isinstance(n, list):
                for it in n:
                    _walk(it)
                return
            if isinstance(n, dict):
                u = n.get("url") or n.get("video_url") or n.get("videoUrl") or n.get("link") or n.get("video")
                t = n.get("title") or n.get("video_title") or n.get("name") or n.get("label") or n.get("topic")
                if u:
                    _push(t, u)
                # common containers
                for k in ["videos", "items", "lessons", "playlist", "week", "weeks", "modules"]:
                    if k in n:
                        _walk(n.get(k))
                # also try values
                if not u:
                    for v in n.values():
                        if isinstance(v, (dict, list)):
                            _walk(v)
                return

        _walk(node)
        seen = set()
        dedup = []
        for it in out:
            u = it.get("url")
            if not u or u in seen:
                continue
            seen.add(u)
            dedup.append(it)
        return dedup

    # Build index from stored attempts
    by_video_no = {}
    attempts_count = {}
    total_videos = None
    for d in _docs:
        try:
            tv = d.get("totalVideos")
            if tv:
                total_videos = int(tv)
        except Exception:
            pass

        # Backward compatibility: older rows may store the video number under
        # different field names.
        vn = (
            d.get("videoNo")
            or d.get("video_no")
            or d.get("videoIndex")
            or d.get("video_index")
            or d.get("video")
        )
        try:
            vn = int(vn) if vn is not None else None
        except Exception:
            vn = None
        if not vn:
            continue
        # count all attempt docs per videoNo (used for Admin display)
        attempts_count[vn] = int(attempts_count.get(vn, 0)) + 1
        # keep the newest doc per videoNo (cursor is sorted by updatedAt desc)
        if vn not in by_video_no:
            by_video_no[vn] = d

    # Pull course videos (title + url). Use this user's course_state to avoid cross-user mismatches.
    cs = db.course_states.find_one(
        {"uid": uid, "courseTitle": {"$regex": ct_re, "$options": "i"}},
        projection={"_id": 0, "videos": 1},
    ) or {}
    video_list = _flatten_video_entries(cs.get("videos"))
    # Fallback: in some versions, the canonical course structure is stored in a
    # shared `courses` collection (not per-user).
    if not video_list:
        try:
            cdoc = db.courses.find_one(
                {"courseTitle": {"$regex": ct_re, "$options": "i"}},
                projection={"_id": 0, "videos": 1},
            ) or {}
            video_list = _flatten_video_entries(cdoc.get("videos")) or video_list
        except Exception:
            pass
    video_urls = [it.get("url") for it in video_list if isinstance(it, dict) and isinstance(it.get("url"), str)]

    # Pull pass map so we can show status even when scores are missing
    cp = db.course_progress.find_one(
        {"uid": uid, "courseTitle": {"$regex": ct_re, "$options": "i"}},
        projection={"_id": 0, "quizPassedMap": 1, "quizSubmittedMap": 1, "quizCompletedMap": 1, "highestUnlockedId": 1},
    ) or {}
    passed_map = cp.get("quizPassedMap") or {}

    highest_unlocked = cp.get("highestUnlockedId")
    try:
        highest_unlocked = int(highest_unlocked) if highest_unlocked is not None else None
    except Exception:
        highest_unlocked = None

    # Decide video count.
    # ✅ Admin requirement: always show ALL video titles (even if locked/unlocked).
    # Prefer the full list from course_states when available.
    pm_max = 0
    try:
        for k in (passed_map or {}).keys():
            try:
                pm_max = max(pm_max, int(k))
            except Exception:
                pass
    except Exception:
        pm_max = 0

    attempts_max = 0
    try:
        attempts_max = max(by_video_no.keys()) if by_video_no else 0
    except Exception:
        attempts_max = 0

    if video_list:
        total_videos = len(video_list)
    else:
        total_videos = max(pm_max, attempts_max, len(video_urls), int(total_videos or 0))

    rows = []
    # Fill rows from 1..N so table doesn't come empty
    N = int(total_videos) if total_videos else len(video_urls)
    if N <= 0:
        N = 0
    for i in range(1, N + 1):
        d = by_video_no.get(i)
        link = video_urls[i - 1] if (video_urls and (i - 1) < len(video_urls)) else None
        title = None
        if video_list and (i - 1) < len(video_list):
            title = (video_list[i - 1] or {}).get("title")

        legacy_inferred = False
        attempted = False
        # Determine passed: prefer stored attempt doc; fallback to progress map (may contain false for failed)
        passed = None
        if d is not None and ("passed" in d):
            passed = bool(d.get("passed"))
            attempted = True
        if passed is None:
            key = str(i)
            attempted = key in (passed_map or {})
            if attempted:
                try:
                    passed = bool((passed_map or {}).get(key))
                except Exception:
                    passed = False
            else:
                # ✅ Legacy fallback (old courses): if marks/attempt docs aren't stored
                # and quizPassedMap doesn't contain the early videos, infer passed
                # based on highestUnlockedId.
                # Example: highestUnlockedId=3 implies quizzes for Video 1 & 2 were passed.
                legacy_inferred = False
                if highest_unlocked is not None and i < highest_unlocked:
                    attempted = True
                    passed = True
                    legacy_inferred = True
                else:
                    passed = False
                    legacy_inferred = False

        # attemptsUsed:
        # 1) if stored attempt doc exists, use its attemptsUsed
        # 2) else, use count of stored docs for that video
        # 3) else, fallback to course_progress flags
        attempts_used = (d.get("attemptsUsed") or d.get("attempts") or d.get("attemptCount")) if d else None
        if attempts_used is None:
            attempts_used = attempts_count.get(i)
        if attempts_used is None:
            # legacy course_progress often stores only boolean maps
            key = str(i)
            submitted_map = (cp.get("quizSubmittedMap") or {}) if isinstance(cp, dict) else {}
            completed_map = (cp.get("quizCompletedMap") or {}) if isinstance(cp, dict) else {}
            if key in (submitted_map or {}) or key in (completed_map or {}):
                attempts_used = 1
            elif highest_unlocked is not None and i < highest_unlocked:
                attempts_used = 1
            else:
                attempts_used = 0

        # whether this row has real stored marks
        has_score = False
        try:
            has_score = (d is not None) and (d.get("lastScore") is not None) and (d.get("totalQuestions") is not None)
        except Exception:
            has_score = False

        rows.append({
            "uid": uid,
            "courseTitle": course_title,
            "videoNo": i,
            "totalVideos": total_videos,
            "videoTitle": (title or (d.get("videoTitle") if d else None) or f"Video {i}"),
            "video_url": (d.get("video_url") if d else None) or link,
            "quiz_id": d.get("quiz_id") if d else None,
            "bestScore": (d.get("bestScore") or d.get("best") or d.get("best_score") or d.get("score") or d.get("marks") or d.get("lastScore")) if d else None,
            "lastScore": (d.get("lastScore") or d.get("last") or d.get("last_score") or d.get("score") or d.get("marks")) if d else None,
            "totalQuestions": (d.get("totalQuestions") or d.get("total") or d.get("max") or d.get("totalQ")) if d else None,
            "required": d.get("required") if d else None,
            "passed": bool(passed),
            "attempted": bool(attempted),
            "attemptsUsed": int(attempts_used or 0),
            "hasScore": bool(has_score),
            "legacyInferred": bool(legacy_inferred),
            "lastAttemptAt": (d.get("lastAttemptAt") or d.get("updatedAt")) if d else None,
            "updatedAt": d.get("updatedAt") if d else None,
        })

    return jsonify({
        "ok": True,
        "rows": rows,
        "meta": {"uid": uid, "courseTitle": course_title, "totalVideos": total_videos}
    })

@app.get("/admin/user/<uid>/mocktests")
def admin_user_mocktests(uid):
    """Recent mock test sessions for a single user."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    col = _mocktest_sessions_col(db)

    docs = list(col.find({"uid": uid}, projection={"_id": 0}).sort("created_at", -1).limit(200))

    out = []
    for d in docs:
        total_score = d.get("total_score")
        total_marks = d.get("total_marks")
        score_percent = None
        try:
            if total_score is not None and total_marks:
                score_percent = (float(total_score) / float(total_marks)) * 100.0
        except Exception:
            score_percent = None

        submitted_at = d.get("submittedAt") or d.get("submitted_at")
        status = (d.get("status") or "").lower()
        if submitted_at is None and status == "submitted":
            submitted_at = d.get("updated_at")

        passed = None
        if score_percent is not None:
            passed = bool(score_percent >= 60.0)

        out.append({
            "sessionId": d.get("session_id") or d.get("sessionId") or d.get("_id"),
            "mode": d.get("mode") or "unknown",
            "createdAt": d.get("created_at") or d.get("createdAt"),
            "submittedAt": submitted_at,
            "scorePercent": score_percent,
            "passed": bool(passed) if passed is not None else False,
        })

    return jsonify({"ok": True, "rows": out})


@app.get("/admin/user/<uid>/mocktests/<session_id>")
def admin_user_mocktest_detail(uid, session_id):
    """Full detail for a single mock test session (scores + analysis)."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    col = _mocktest_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id}, projection={"_id": 0})
    if not doc:
        # fallback: some older docs used sessionId
        doc = col.find_one({"uid": uid, "sessionId": session_id}, projection={"_id": 0})
    if not doc:
        return jsonify({"ok": False, "error": "Not found"}), 404

    # Keep payload compact: don't send full hidden tests.
    out = {
        "sessionId": doc.get("session_id") or doc.get("sessionId") or session_id,
        "title": doc.get("title") or "Mock Test",
        "mode": doc.get("mode") or "unknown",
        "pattern": doc.get("pattern") or {},
        "status": doc.get("status") or "",
        "createdAt": doc.get("created_at") or doc.get("createdAt"),
        "submittedAt": doc.get("submittedAt") or doc.get("submitted_at"),
        "scores": doc.get("scores") or {},
        "total_score": doc.get("total_score"),
        "total_marks": doc.get("total_marks"),
        "coding_total_marks": doc.get("coding_total_marks") or doc.get("coding_total_marks"),
        "coding_details": doc.get("coding_details") or {},
        "analysis": doc.get("analysis"),
        "updatedAt": doc.get("updated_at") or doc.get("updatedAt"),
    }

    return jsonify({"ok": True, "session": out})


@app.get("/admin/course/<course_title>/summary")
def admin_course_summary(course_title):
    """Course drilldown summary."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    ct = (course_title or "").strip()
    if not ct:
        return jsonify({"ok": False, "error": "Missing courseTitle"}), 400

    # users studying = distinct uid in course_states
    studying_uids = list(db.course_states.aggregate([
        {"$match": {"courseTitle": ct}},
        {"$group": {"_id": "$uid"}},
    ]))
    users_studying = len(studying_uids)

    users_on_hold = db.course_holds.count_documents({"courseTitle": ct, "held": True})

    # progress stats
    cps = list(db.course_progress.find({"courseTitle": ct}, projection={"_id": 0, "quizPassedMap": 1}))

    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    passed_total = sum(count_true(x.get("quizPassedMap") or {}) for x in cps)

    return jsonify({
        "ok": True,
        "courseTitle": ct,
        "usersStudying": int(users_studying),
        "usersOnHold": int(users_on_hold),
        "progressDocs": len(cps),
        "passedQuizCount": int(passed_total),
    })


@app.get("/admin/course/<course_title>/progress")
def admin_course_progress_detail(course_title):
    """User progress rows for a given course."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    db = get_db()
    ct = (course_title or "").strip()
    if not ct:
        return jsonify({"ok": False, "error": "Missing courseTitle"}), 400

    # total quizzes per course derived from course_states
    def count_videos(obj):
        if obj is None:
            return 0
        if isinstance(obj, list):
            return sum(count_videos(it) for it in obj)
        if isinstance(obj, dict):
            if isinstance(obj.get("videos"), list):
                return len(obj.get("videos"))
            if "video" in obj and ("topic" in obj or "title" in obj):
                return 1
            return sum(count_videos(v) for v in obj.values())
        return 0

    total = 0
    for st in db.course_states.find({"courseTitle": ct}, projection={"_id": 0, "videos": 1}):
        total = max(total, count_videos(st.get("videos")))

    def count_true(dct):
        if not isinstance(dct, dict):
            return 0
        return sum(1 for _, v in dct.items() if bool(v))

    users = {u["uid"]: u for u in db.users.find({}, projection={"_id": 0, "uid": 1, "email": 1, "name": 1})}

    rows = []
    for cp in db.course_progress.find({"courseTitle": ct}, projection={"_id": 0, "uid": 1, "quizPassedMap": 1, "updatedAt": 1}):
        uid = cp.get("uid")
        u = users.get(uid) or {}
        passed = count_true(cp.get("quizPassedMap") or {})
        pct = int(round((passed / total) * 100)) if total else 0
        rows.append({
            "uid": uid,
            "name": u.get("name"),
            "email": u.get("email"),
            "courseTitle": ct,
            "totalQuizzes": int(total),
            "passedQuizzes": int(passed),
            "percent": pct,
            "held": is_course_held(uid, ct),
            "updatedAt": cp.get("updatedAt"),
        })

    rows.sort(key=lambda x: (x.get("held") is False, x.get("percent") or 0), reverse=True)
    return jsonify({"ok": True, "rows": rows})


@app.get("/admin/course/<course_title>/quiz-attempts")
def admin_course_quiz_attempts(course_title):
    """Per-video quiz attempt summaries for a course (across users)."""
    _, err = require_admin()
    if err:
        msg, code = err
        return jsonify({"error": msg}), code

    ct = (course_title or "").strip()
    if not ct:
        return jsonify({"ok": False, "error": "Missing courseTitle"}), 400

    db = get_db()
    docs = list(
        db.quiz_attempts.find({"courseTitle": ct}, projection={"_id": 0})
        .sort([("videoNo", 1), ("lastAttemptAt", -1)])
        .limit(5000)
    )

    uids = sorted({d.get("uid") for d in docs if d.get("uid")})
    users = {}
    if uids:
        for u in db.users.find({"uid": {"$in": uids}}, projection={"_id": 0, "uid": 1, "name": 1, "email": 1}):
            users[u.get("uid")] = {"name": u.get("name"), "email": u.get("email")}

    rows = []
    for d in docs:
        uid = d.get("uid")
        u = users.get(uid) or {}
        rows.append({
            "uid": uid,
            "name": u.get("name"),
            "email": u.get("email"),
            "courseTitle": d.get("courseTitle"),
            "videoNo": d.get("videoNo"),
            "totalVideos": d.get("totalVideos"),
            "lastScore": d.get("lastScore"),
            "bestScore": d.get("bestScore"),
            "totalQuestions": d.get("totalQuestions"),
            "required": d.get("required"),
            "passed": bool(d.get("passed")),
            "attemptsUsed": d.get("attemptsUsed"),
            "lastAttemptAt": d.get("lastAttemptAt") or d.get("updatedAt"),
            "video_url": d.get("video_url"),
        })

    return jsonify({"ok": True, "rows": rows})

# - mocktest_sessions: one doc per test session (questions + results + analysis)

def _mocktest_sessions_col(db):
    return db.mocktest_sessions

# Judge0
# NOTE: RapidAPI endpoint requires headers; most self-host/public instances do not.
# Default to the official public CE instance so coding works out-of-the-box.
_JUDGE0_BASE_URL = (os.getenv("JUDGE0_BASE_URL") or "https://ce.judge0.com").rstrip("/")
_JUDGE0_RAPIDAPI_KEY = (os.getenv("JUDGE0_RAPIDAPI_KEY") or "").strip()
_JUDGE0_RAPIDAPI_HOST = (os.getenv("JUDGE0_RAPIDAPI_HOST") or "").strip()

# Judge0 language IDs (Judge0 CE common IDs; can differ by deployment)
# https://ce.judge0.com/#languages (verify for your instance)
_JUDGE0_LANG_IDS = {
    "c": 50,
    "cpp": 54,
    "c++": 54,
    "java": 62,
    "python": 71,
    "python3": 71,
}

def _judge0_headers():
    h = {"Content-Type": "application/json"}
    # RapidAPI style headers (optional)
    if _JUDGE0_RAPIDAPI_KEY:
        h["X-RapidAPI-Key"] = _JUDGE0_RAPIDAPI_KEY
    if _JUDGE0_RAPIDAPI_HOST:
        h["X-RapidAPI-Host"] = _JUDGE0_RAPIDAPI_HOST
    return h

def _judge0_run(source_code: str, stdin: str, language: str, expected: str = None, time_limit: float = None):
    """Run code on Judge0 and return dict with stdout/stderr/status. If expected provided, include pass bool."""
    lang_key = (language or "").strip().lower()
    lang_id = _JUDGE0_LANG_IDS.get(lang_key)
    if not lang_id:
        raise ValueError(f"Unsupported language: {language}. Allowed: Java, C, C++, Python")

    payload = {
        "language_id": lang_id,
        "source_code": source_code or "",
        "stdin": stdin or "",
        # Optional knobs: keep defaults unless needed
        "redirect_stderr_to_stdout": False,
    }

    # Judge0 CE supports per-submission time limits; RapidAPI proxy usually supports it too.
    # Only set when provided to avoid incompat issues across deployments.
    if time_limit is not None:
        payload["cpu_time_limit"] = float(time_limit)

    url = f"{_JUDGE0_BASE_URL}/submissions?base64_encoded=false&wait=true"
    r = requests.post(url, headers=_judge0_headers(), json=payload, timeout=60)
    r.raise_for_status()
    out = r.json() or {}

    stdout = (out.get("stdout") or "").strip()
    stderr = (out.get("stderr") or "").strip()
    compile_out = (out.get("compile_output") or "").strip()
    status = (out.get("status") or {}) or {}
    status_id = status.get("id")
    status_desc = status.get("description")

    res = {
        "stdout": stdout,
        "stderr": stderr,
        "compile_output": compile_out,
        "status_id": status_id,
        "status": status_desc,
        "time": out.get("time"),
        "memory": out.get("memory"),
    }
    if expected is not None:
        exp = (expected or "").strip()
        res["expected"] = exp
        res["passed"] = (stdout.strip() == exp.strip()) and (not stderr) and (not compile_out) and (status_id in [3])  # 3 = Accepted
    return res

# Piston (free code runner) fallback - works without API keys
_PISTON_BASE_URL = (os.getenv("PISTON_BASE_URL") or "http://127.0.0.1:2000/api/v2").rstrip("/")

# Optional auth token (mainly for hosted / public Piston instances).
# If set, it will be sent as the Authorization header value.
# Example: "Bearer <token>" (depends on your provider).
_PISTON_AUTH = (os.getenv("PISTON_AUTH") or "").strip()

def _piston_headers():
    h = {}
    if _PISTON_AUTH:
        h["Authorization"] = _PISTON_AUTH
    return h


# Optional Sphere Engine support (if you have credentials)
# Set SPHERE_ENGINE_ENDPOINT like: https://<your-subdomain>.compilers.sphere-engine.com
# Set SPHERE_ENGINE_TOKEN to your access token
_SPHERE_ENGINE_ENDPOINT = (os.getenv("SPHERE_ENGINE_ENDPOINT") or "").rstrip("/")
_SPHERE_ENGINE_TOKEN = (os.getenv("SPHERE_ENGINE_TOKEN") or "").strip()

_SPHERE_LANG = {
    "python": "python",
    "python3": "python",
    "java": "java",
    "c": "c",
    "cpp": "cpp",
    "c++": "cpp",
}

_PISTON_LANG = {
    "python": "python",
    "python3": "python",
    "java": "java",
    "c": "c",
    "cpp": "cpp",
    "c++": "cpp",
}

def _piston_run(source_code: str, stdin: str, language: str, expected: str = None, time_limit: float = None):
    lang_key = (language or "").strip().lower()
    lang = _PISTON_LANG.get(lang_key)
    if not lang:
        raise ValueError(f"Unsupported language: {language}. Allowed: Java, C, C++, Python")

    url = f"{_PISTON_BASE_URL}/execute"
    payload = {
        "language": lang,
        "version": "*",
        "files": [{"content": source_code or ""}],
        "stdin": stdin or "",
    }
    r = requests.post(url, headers=_piston_headers(), json=payload, timeout=60)
    # Helpful error for common misconfigurations (401 / wrong base path)
    if r.status_code == 401:
        raise RuntimeError(
            "Piston returned 401 Unauthorized. If you're using the public emkc.org Piston, it now requires an auth token. "
            "Set PISTON_AUTH (e.g., 'Bearer <token>') and PISTON_BASE_URL accordingly, or use your self-hosted Piston (recommended)."
        )
    r.raise_for_status()
    out = r.json() or {}

    run = out.get("run") or {}
    stdout = (run.get("stdout") or "").strip()
    stderr = (run.get("stderr") or "").strip()

    res = {
        "stdout": stdout,
        "stderr": stderr,
        "compile_output": "",
        "status_id": 3 if not stderr else 11,
        "status": "Accepted" if not stderr else "Runtime Error",
        "time": run.get("time"),
        "memory": run.get("memory"),
    }
    if expected is not None:
        exp = (expected or "").strip()
        res["expected"] = exp
        res["passed"] = (stdout.strip() == exp.strip()) and (not stderr)
    return res

def _sphere_engine_run(source_code: str, stdin: str, language: str, expected: str = None, time_limit: float = None):
    if not _SPHERE_ENGINE_ENDPOINT or not _SPHERE_ENGINE_TOKEN:
        raise RuntimeError("Sphere Engine is not configured")

    lang_key = (language or "").strip().lower()
    lang = _SPHERE_LANG.get(lang_key)
    if not lang:
        raise ValueError(f"Unsupported language: {language}. Allowed: Java, C, C++, Python")

    # Sphere Engine Compilers API uses problem id / compiler id depending on plan.
    # Since setups vary, we keep Sphere optional and recommend using Piston by default.
    raise RuntimeError("Sphere Engine integration placeholder: configure compiler IDs for your account")


def _estimate_complexity(source_code: str, language: str):
    """Very lightweight complexity estimator (heuristic).
    Users can override by adding a comment like: '# O(n log n), O(1)' or '// O(n), O(n)'.
    """
    code = source_code or ""

    # Explicit hint
    for line in code.splitlines()[:40]:
        if "O(" in line:
            parts = re.findall(r'O\([^\)]*\)', line)
            if len(parts) >= 2:
                return {"estimated_time": parts[0], "estimated_space": parts[1], "notes": "From inline hint in code comments."}
            if len(parts) == 1:
                return {"estimated_time": parts[0], "estimated_space": "O(1)", "notes": "From inline hint in code comments."}

    lower = code.lower()

    # Crude nesting estimate
    nest = 0
    cur = 0
    for raw in code.splitlines():
        line = raw.strip().lower()
        if line.startswith(("for ", "while ")):
            cur += 1
            nest = max(nest, cur)
        # naive close on braces/return
        if line == "}" or line.startswith("return"):
            cur = max(0, cur - 1)

    has_sort = ("sort" in lower) or ("sorted(" in lower)
    uses_hash = any(k in lower for k in ["dict", "hashmap", "unordered_map", "set<", "hashset", "map<"])
    uses_heap = ("heapq" in lower) or ("priorityqueue" in lower) or ("priority_queue" in lower)

    if nest >= 2:
        t = "O(n^2)"
    elif nest == 1 and has_sort:
        t = "O(n log n)"
    elif nest == 1:
        t = "O(n)"
    else:
        t = "O(1)" if not has_sort else "O(n log n)"

    s = "O(n)" if (uses_hash or uses_heap) else "O(1)"
    return {"estimated_time": t, "estimated_space": s, "notes": "Heuristic estimate. Add an 'O(...)' comment to override."}

def _code_run(source_code: str, stdin: str, language: str, expected: str = None, time_limit: float = None):
    """Robust code runner: try Judge0 first, then Piston as fallback."""
    last = None
    try:
        return _judge0_run(source_code, stdin, language, expected=expected, time_limit=time_limit)
    except Exception as e:
        last = e

    try:
        return _piston_run(source_code, stdin, language, expected=expected, time_limit=time_limit)
    except Exception as e:
        last = e

    raise last or RuntimeError("Code runner failed")


def _mocktest_generate_apti_questions(section: str, n: int, difficulty: str = "mixed"):
    """Generate aptitude questions using Gemini and self-verify with Gemini."""
    n = int(n or 0)
    if n <= 0:
        return []

    section = (section or "").strip().lower()
    if section not in ["general", "tech"]:
        section = "general"

    # Expose a 'subsection' field to support UI section menus.
    if section == "general":
        focus = "General Aptitude (company screening)"
        subsections = ["Quant", "Logical", "Verbal"]
        topics = "Percentages, Ratios, Time & Work, Time & Distance, Profit & Loss, Probability basics, Puzzles, Syllogisms, Reading comprehension short, Error spotting"
    else:
        focus = "Technical Aptitude (company screening)"
        subsections = ["DSA", "OOP", "OS", "CN", "DBMS"]
        topics = "Big-O, arrays/strings, stacks/queues, hashing, trees basics, OOP principles, threading basics, OS processes vs threads, deadlock, TCP/UDP, HTTP basics, normalization, indexing"

    recent_qs = list(RECENT_MCQ_HASHES)[-60:]
    prompt = f"""
You are generating aptitude multiple-choice questions for a mock test.

Return STRICT JSON only (no markdown), schema:
{{
  "questions": [
    {{
      "id": "Q1",
      "section": "{section}",
      "subsection": "{subsections[0]}",
      "difficulty": "{difficulty}",
      "topic": "...",
      "question": "...",
      "options": ["A ...","B ...","C ...","D ..."],
      "correctIndex": 0,
      "answerText": "A ...",
      "explanation": "1-2 lines"
    }}
  ]
}}

Rules:
- Avoid repeating any of these recent questions (do NOT reuse wording/hardly similar): {recent_qs}
- Generate exactly {n} questions.
- Each question must have 4 options only.
- correctIndex must be 0-3 and must match answerText.
- subsection must be one of: {subsections}.
- Avoid ambiguous questions. Use precise numerical values.
- 1 mark per question.
- Topic coverage: {topics}.
"""

    obj = _gemini_json(prompt) or {}
    qs = obj.get("questions") if isinstance(obj, dict) else None
    if not isinstance(qs, list):
        qs = []

    # Verification pass (Gemini: verify correctIndex consistency)
    verify_prompt = f"""
You will be given aptitude MCQ questions as JSON.
Verify each question has exactly 4 options, correctIndex 0-3, and answerText matches the option at correctIndex.
If any is wrong, fix it.
Return STRICT JSON with the same schema: {{ "questions": [ ... ] }} only.

INPUT JSON:
{json.dumps({"questions": qs}, ensure_ascii=False)}
"""
    verified = _gemini_json(verify_prompt) or {}
    vqs = verified.get("questions") if isinstance(verified, dict) else None
    if isinstance(vqs, list) and len(vqs) == len(qs):
        qs = vqs

    # --- Fallback bank ---
    # Gemini can occasionally return empty/invalid output. Ensure we ALWAYS return n MCQs.
    if not qs:
        if section == "general":
            bank = [
                {
                    "subsection": "Quant",
                    "topic": "Percentages",
                    "question": "If the price of an item increases by 20% and then decreases by 20%, what is the net change?",
                    "options": ["0%", "4% decrease", "4% increase", "20% decrease"],
                    "correctIndex": 1,
                    "explanation": "Take 100 → 120 → 96, net 4% decrease.",
                },
                {
                    "subsection": "Quant",
                    "topic": "Time & Work",
                    "question": "A can do a job in 10 days and B can do it in 15 days. In how many days can they finish together?",
                    "options": ["5", "6", "7", "8"],
                    "correctIndex": 1,
                    "explanation": "Rate = 1/10 + 1/15 = 1/6.",
                },
                {
                    "subsection": "Logical",
                    "topic": "Syllogism",
                    "question": "Statements: All pens are blue. Some blue are costly. Conclusions: (1) Some pens are costly. (2) All costly are pens.\\nWhich conclusion follows?",
                    "options": ["Only 1", "Only 2", "Both", "Neither"],
                    "correctIndex": 3,
                    "explanation": "No guarantee costly intersects pens; costly→pens not given.",
                },
                {
                    "subsection": "Verbal",
                    "topic": "Error spotting",
                    "question": "Choose the correct sentence:",
                    "options": ["He don't like coffee.", "He doesn't likes coffee.", "He doesn't like coffee.", "He didn't likes coffee."],
                    "correctIndex": 2,
                    "explanation": "With does/doesn't use base verb form.",
                },
                {
                    "subsection": "Logical",
                    "topic": "Calendar",
                    "question": "If today is Wednesday, what day will it be after 10 days?",
                    "options": ["Saturday", "Sunday", "Monday", "Tuesday"],
                    "correctIndex": 0,
                    "explanation": "10 mod 7 = 3; Wed + 3 = Saturday.",
                },
            ]
        else:
            bank = [
                {
                    "subsection": "DSA",
                    "topic": "Big-O",
                    "question": "What is the time complexity of binary search on a sorted array of size n?",
                    "options": ["O(n)", "O(log n)", "O(n log n)", "O(1)"],
                    "correctIndex": 1,
                    "explanation": "Each step halves the search space.",
                },
                {
                    "subsection": "OOP",
                    "topic": "Polymorphism",
                    "question": "Which OOP concept allows the same interface to represent different underlying forms?",
                    "options": ["Encapsulation", "Abstraction", "Inheritance", "Polymorphism"],
                    "correctIndex": 3,
                    "explanation": "Polymorphism enables one interface, many implementations.",
                },
                {
                    "subsection": "OS",
                    "topic": "Processes",
                    "question": "Which is true about a process vs a thread?",
                    "options": ["Threads have separate address spaces", "Processes share the same address space", "Threads share a process address space", "Processes cannot have multiple threads"],
                    "correctIndex": 2,
                    "explanation": "Threads within a process share address space/resources.",
                },
                {
                    "subsection": "CN",
                    "topic": "TCP/UDP",
                    "question": "Which protocol is connection-oriented?",
                    "options": ["UDP", "TCP", "IP", "ICMP"],
                    "correctIndex": 1,
                    "explanation": "TCP establishes a connection (3-way handshake).",
                },
                {
                    "subsection": "DBMS",
                    "topic": "Normalization",
                    "question": "Which normal form removes partial dependency?",
                    "options": ["1NF", "2NF", "3NF", "BCNF"],
                    "correctIndex": 1,
                    "explanation": "2NF removes partial dependency on a candidate key.",
                },
            ]

        qs = []
        for i in range(n):
            b = dict(bank[i % len(bank)])
            b["id"] = f"{section.upper()}-{i+1}"
            b["section"] = section
            b["difficulty"] = difficulty
            b["answerText"] = b["options"][int(b.get("correctIndex") or 0)]
            qs.append(b)

    # Normalize ids
    out = []
    for i, q in enumerate(qs, start=1):
        if not isinstance(q, dict):
            continue
        opts = q.get("options") or []
        if not isinstance(opts, list) or len(opts) != 4:
            continue
        ci = q.get("correctIndex")
        try:
            ci = int(ci)
        except Exception:
            ci = 0
        if ci < 0 or ci > 3:
            ci = 0
        out.append({
            "id": q.get("id") or f"{section.upper()}-{i}",
            "section": section,
            "subsection": q.get("subsection") or "",
            "difficulty": q.get("difficulty") or difficulty,
            "topic": q.get("topic") or "",
            "question": q.get("question") or "",
            "options": opts,
            "correctIndex": ci,
            "answerText": q.get("answerText") or opts[ci],
            "explanation": q.get("explanation") or "",
            "marks": 1,
        })
    return out


def _default_starter_code() -> dict:
    return {
        "python": "import sys\n\ndef solve():\n    data = sys.stdin.read().strip().split()\n    # TODO: parse input and implement\n    # print(result)\n    return\n\nif __name__ == '__main__':\n    solve()\n",
        "java": "import java.io.*;\nimport java.util.*;\n\npublic class Main {\n  static void solve(FastScanner fs, StringBuilder out) throws Exception {\n    // TODO: implement\n  }\n\n  public static void main(String[] args) throws Exception {\n    FastScanner fs = new FastScanner(System.in);\n    StringBuilder out = new StringBuilder();\n    solve(fs, out);\n    System.out.print(out.toString());\n  }\n\n  static class FastScanner {\n    private final InputStream in;\n    private final byte[] buffer = new byte[1 << 16];\n    private int ptr = 0, len = 0;\n    FastScanner(InputStream is){ in=is; }\n    private int read() throws IOException {\n      if (ptr >= len) {\n        len = in.read(buffer);\n        ptr = 0;\n        if (len <= 0) return -1;\n      }\n      return buffer[ptr++];\n    }\n    String next() throws IOException {\n      StringBuilder sb = new StringBuilder();\n      int c;\n      while ((c = read()) != -1 && c <= ' ') {}\n      if (c == -1) return null;\n      do { sb.append((char)c); } while ((c = read()) != -1 && c > ' ');\n      return sb.toString();\n    }\n    Integer nextInt() throws IOException {\n      String s = next();\n      return s == null ? null : Integer.parseInt(s);\n    }\n  }\n}\n",
        "c": "#include <stdio.h>\n#include <stdlib.h>\n#include <string.h>\n\nint main(void) {\n    // TODO: read input from stdin and implement\n    // printf(\"%d\\n\", ans);\n    return 0;\n}\n",
        "cpp": "#include <bits/stdc++.h>\nusing namespace std;\n\nint main(){\n  ios::sync_with_stdio(false);\n  cin.tie(nullptr);\n\n  // TODO: parse input and implement\n\n  return 0;\n}\n"
    }


def _merge_starter_code(model_value) -> dict:
    """Merge model-provided starterCode with defaults.

    UI expects keys: python, java, c, cpp.
    """
    defaults = _default_starter_code()
    if not model_value:
        return defaults
    if isinstance(model_value, str):
        s = model_value.strip()
        if s:
            defaults["python"] = s
        return defaults
    if not isinstance(model_value, dict):
        return defaults

    norm = {}
    for k, v in model_value.items():
        if not isinstance(v, str) or not v.strip():
            continue
        kk = str(k).strip().lower()
        if kk in ("py", "python3"):
            kk = "python"
        if kk in ("c++", "cpp17", "c++17", "cplusplus", "cxx"):
            kk = "cpp"
        if kk in ("c", "java", "python", "cpp"):
            norm[kk] = v

    defaults.update(norm)
    return defaults


def _mocktest_generate_coding_problems(n: int, difficulty: str = "mixed"):
    """Generate DSA coding problems + hidden tests + python reference solution, and validate reference via Judge0.

    Expected model output (STRICT JSON):
    { "problems": [ ... ] }
    """
    n = int(n or 0)
    if n <= 0:
        return []

    difficulty = (difficulty or "mixed").strip().lower()
    if difficulty not in ["easy", "mixed", "medium", "hard"]:
        difficulty = "mixed"


    # ✅ Total testcases per problem by difficulty (user requirement):
    # Easy: 15 (5 sample + 10 hidden)
    # Medium: 20 (5 sample + 15 hidden)
    # Hard: 25 (5 sample + 20 hidden)
    total_tests_map = {"easy": 15, "medium": 20, "hard": 25, "mixed": 20}
    want_total = int(total_tests_map.get(difficulty, 8))
    want_samples = 5
    want_hidden = max(0, want_total - want_samples)

    # prompt-level dedupe: ask model to avoid recently used titles
    avoid_titles = list(RECENT_CODING_TITLES)[-50:]

    prompt = (
        "You are generating coding problems for a mock test platform.\n"
        "Return STRICT JSON ONLY. No markdown.\n\n"
        "Return JSON with key 'problems' (array). Each problem must include: id/slug, title, difficulty (easy|medium|hard), topics (array), statement, inputFormat, outputFormat, constraints (array), starterCode (python/java/cpp) [UNSOLVED TEMPLATE ONLY - no working solution; include only skeleton + input reading + TODO comments], solutionHint, tests.\n\n"
        f"COMPLEXITY: Also include 'timeComplexity' and 'spaceComplexity' as Big-O strings (e.g., \"O(n log n)\", \"O(1)\").\n"
        f"TESTS: Provide EXACTLY {want_total} tests per problem in 'tests'. You MUST provide {want_samples} sample tests (isSample=true) and the remaining {want_hidden} tests must be hidden (isSample=false). Total tests must be {want_total} ({want_samples} sample + {want_hidden} hidden).\n"
        "CRITICAL: All testcases must be UNIQUE. Do NOT repeat the same input (even with different spacing/newlines). Hidden tests MUST NOT duplicate any sample test. If you are struggling, create new edge-case variations, but never repeat an input.\n"
        "Each test item: {stdin, expected, isSample}.\n"
        "expected must match exact output. Include edge cases.\n\n"
        + (f"Avoid repeating any of these titles (do NOT reuse): {avoid_titles}.\n" if avoid_titles else "")
        + f"Generate {n} problems.\n"
    )


    # IMPORTANT: If Gemini fails, do NOT use local fallbacks.
    # Retry Gemini a few times; if it still fails, raise so the client can ask the user to regenerate.
    probs = []
    last_err = None
    for _attempt in range(1, 5):  # up to 4 attempts
        try:
            obj = _gemini_json(prompt) or {}
            probs = obj.get("problems") if isinstance(obj, dict) else None
            if not isinstance(probs, list):
                probs = []
            if probs:
                last_err = None
                break
        except Exception as e:
            last_err = e
            probs = []

    if not probs:
        raise RuntimeError(f"Gemini failed to generate coding problems. Please regenerate. Details: {last_err}")

    def _coding_full_marks(diff: str) -> int:
        d = (diff or "").strip().lower()
        if d == "easy":
            return 15
        if d == "medium":
            return 20
        if d in ["hard", "difficult"]:
            return 25
        # mixed/unknown
        return 20

    out = []
    for i, p in enumerate(probs, start=1):
        if not isinstance(p, dict):
            continue

        # Accept tests in multiple formats.

        tests = p.get('tests') or p.get('testcases') or []

        samples = []

        hid = []

        if isinstance(tests, list) and tests:

            tagged_samples = [t for t in tests if isinstance(t, dict) and t.get('isSample') is True]

            tagged_hidden = [t for t in tests if isinstance(t, dict) and t.get('isSample') is False]

            if tagged_samples or tagged_hidden:

                samples = tagged_samples[:4]

                # ✅ Don't truncate hidden tests here.
                # We'll decide the final hidden count later (based on difficulty),
                # and we will enforce uniqueness so hidden never repeats sample inputs.
                hid = tagged_hidden

            else:

                samples = [t for t in tests[:4] if isinstance(t, dict)]

                hid = [t for t in tests[4:8] if isinstance(t, dict)]

        else:

            samples = p.get('samples') or []

            hid = p.get('hiddenTests') or []

        if not isinstance(samples, list): samples = []

        if not isinstance(hid, list): hid = []

        if len(samples) > 4:

            hid = (hid or []) + samples[4:]

            samples = samples[:4]

        samples = samples[:4]

        # Keep all hidden candidates for now; we'll decide final counts later.

        # If hidden tests are missing, we will regenerate them later.

        # Hard guarantee: always expose 4 sample tests and 11 hidden tests (total 15)
        # so the frontend never shows "no hidden tests configured".
        # If Gemini returns fewer tests, we pad by repeating existing ones.
        # Ensure at least 1 non-empty sample so UI never shows blank sample boxes.
        # Prefer promoting a hidden test to a sample if samples are missing/invalid.
        def _tc_valid(tc: dict) -> bool:
            if not isinstance(tc, dict):
                return False
            s = str(tc.get("stdin", "")).strip()
            # expected can be empty for some problems, but stdin must be present
            return len(s) > 0

        # drop empty samples
        samples = [x for x in samples if _tc_valid(x)]
        hid = [x for x in hid if _tc_valid(x)]

        if len(samples) == 0:
            if len(hid) > 0:
                # promote first hidden to sample
                samples = [hid.pop(0)]
            else:
                # last-resort fallback (non-empty stdin)
                samples = [{'stdin': '1\n', 'expected': ''}] 

        # normalize keys
        def _norm_tc(tc: dict) -> dict:
            if not isinstance(tc, dict):
                return {'stdin': '', 'expected': ''}
            return {
                'stdin': str(tc.get('stdin', '')),
                'expected': str(tc.get('expected', '')),
            }

        samples = [_norm_tc(x) for x in samples]
        hid = [_norm_tc(x) for x in hid]
        # Decide sample/hidden split
        # Requirement: difficulty-based totals (easy=8, medium=15, hard=20) with 4 samples.
        want_samples = 4
        total_tests_map = {"easy": 8, "medium": 15, "hard": 20, "mixed": 15}
        want_total = int(total_tests_map.get(difficulty, 8))

        # Build a compact problem text for testcase regeneration prompts.
        try:
            _cons = p.get('constraints') if isinstance(p.get('constraints'), list) else []
            problem_text = (
                f"{p.get('title') or ''}\n"
                f"{p.get('statement') or ''}\n\n"
                f"Input Format:\n{p.get('inputFormat') or ''}\n\n"
                f"Output Format:\n{p.get('outputFormat') or ''}\n\n"
                f"Constraints:\n" + "\n".join([str(x) for x in _cons])
            )
        except Exception:
            problem_text = (p.get('statement') or '')
        want_hidden = max(0, want_total - want_samples)

        # Ensure 4 samples: if short, promote hidden; else last-resort duplicate the first sample.
        while len(samples) < want_samples and len(hid) > 0:
            samples.append(hid.pop(0))
        while len(samples) < want_samples:
            samples.append(dict(samples[0]))

        samples = samples[:want_samples]

        # ✅ Ensure hidden testcases are UNIQUE (avoid duplicates across samples/hidden).
        # IMPORTANT: uniqueness is based on INPUT only (stdin), not expected output.
        def _tc_key(tc):
            try:
                raw_in = str((tc or {}).get("stdin") or "")
                # Normalize newlines then collapse all whitespace to single spaces.
                s = raw_in.replace("\r\n", "\n").replace("\r", "\n")
                s = " ".join(s.strip().split())
                return s
            except Exception:
                return ""

        seen_inputs = set()
        uniq_samples = []
        for tc in samples:
            if not isinstance(tc, dict):
                continue
            k = _tc_key(tc)
            if not k or k in seen_inputs:
                continue
            seen_inputs.add(k)
            uniq_samples.append(tc)
        samples = uniq_samples[:want_samples]

        uniq_hid = []
        for tc in hid:
            if not isinstance(tc, dict):
                continue
            k = _tc_key(tc)
            if not k or k in seen_inputs:
                continue
            seen_inputs.add(k)
            uniq_hid.append(tc)
        hid = uniq_hid

        need_hidden = max(0, want_hidden - len(hid))
        regen_attempts = 0
        while need_hidden > 0 and regen_attempts < 3:
            regen_attempts += 1
            avoid_list = [x for x in list(seen_inputs) if x][:80]
            regen_prompt = f"""You are generating additional HIDDEN testcases for this coding problem.

Problem:
{(problem_text or '').strip()[:2500]}

Generate EXACTLY {need_hidden} NEW hidden testcases.
Rules:
- Each testcase must be valid for the problem.
- Inputs MUST be different from all inputs in this avoid list (compare after whitespace normalization): {json.dumps(avoid_list)}
- Return STRICT JSON only:
{{ "hidden": [ {{ "stdin": "....", "expected": "...." }} ] }}
"""
            extra = _gemini_json(regen_prompt) or {}
            extra_hid = extra.get("hidden") or extra.get("testcases") or []
            added = 0
            for tc in extra_hid:
                if not isinstance(tc, dict):
                    continue
                k = _tc_key(tc)
                if not k or k in seen_inputs:
                    continue
                seen_inputs.add(k)
                hid.append(_norm_tc(tc))
                added += 1
                if len(hid) >= want_hidden:
                    break
            if added == 0:
                break
            need_hidden = max(0, want_hidden - len(hid))

        hid = hid[:want_hidden]

        ref = (p.get('referencePython') or '').strip()
        ref_ok = True

        test_results = []

        if ref:

            for t in hid[:6]:

                if not isinstance(t, dict):

                    ref_ok = False

                    break

                try:

                    jr = _code_run(ref, t.get('stdin') or '', 'python', expected=(t.get('expected') or ''))

                    test_results.append(jr)

                    if not jr.get('passed'):

                        ref_ok = False

                        break

                except Exception:

                    ref_ok = False

                    break

        out.append({

            'id': p.get('id') or (p.get('slug') or f'CODING-{i}'),

            'slug': p.get('slug') or '',

            'difficulty': p.get('difficulty') or difficulty,

            'topic': p.get('topic') or '',

            'topics': p.get('topics') if isinstance(p.get('topics'), list) else [],

            'title': p.get('title') or f'Coding Problem {i}',

            'statement': p.get('statement') or '',

            'inputFormat': p.get('inputFormat') or '',

            'outputFormat': p.get('outputFormat') or '',

            'constraints': p.get('constraints') if isinstance(p.get('constraints'), list) else [],

            'samples': samples,

            'hiddenTests': hid,

            'total_marks': (len(samples) + len(hid)),


            # Prefer Gemini-provided starterCode (per language), but always ensure
            # python/java/c/cpp exist so the UI can switch languages reliably.
            'starterCode': _merge_starter_code(p.get('starterCode')),
            # starterCodeFromBank: _default_starter_code(),

            'solutionHint': p.get('solutionHint') or '',

            'timeComplexity': p.get('timeComplexity') or '',
            'spaceComplexity': p.get('spaceComplexity') or '',

            'meta': {

                'referencePythonProvided': bool(ref),

                'referencePythonValidated': bool(ref and ref_ok),

                'sampleCount': len(samples),

                'hiddenCount': len(hid)

            }

        })

    # No fallback: if Gemini still failed to generate valid problems, ask the client to regenerate.
    if len(out) == 0:
        raise RuntimeError("Gemini did not return any valid coding problems. Please regenerate.")
    # Update recent coding titles cache
    for pr in out:
        tt = (pr.get('title') or pr.get('slug') or '').strip()
        if tt:
            RECENT_CODING_TITLES.append(tt)

    return out
def _mocktest_public_session(doc: dict):
    if not doc:
        return None
    return {
        "session_id": str(doc.get("session_id") or doc.get("_id") or ""),
        "title": doc.get("title") or "Mock Test",
        "mode": doc.get("mode") or "all",
        "difficulty": doc.get("difficulty") or "mixed",
        "pattern": doc.get("pattern") or {},
        "created_at": doc.get("created_at"),
        "updated_at": doc.get("updated_at"),
        "status": doc.get("status") or "draft",
        "scores": doc.get("scores") or {},
        "total_score": doc.get("total_score"),
        "total_marks": doc.get("total_marks"),
        "analysis": doc.get("analysis") or None,
        "coding_details": doc.get("coding_details") or {},
        "general_questions": doc.get("general_questions") or [],
        "tech_questions": doc.get("tech_questions") or [],
        "coding_problems": [
            {k: v for k, v in (p or {}).items() if k not in ["referencePython", "referenceJudge0"]}
            | {"hiddenTestCount": len(((p or {}).get("hiddenTests") or [])) or len(((p or {}).get("samples") or [])), "total_marks": (p or {}).get("total_marks")}
            for p in (doc.get("coding_problems") or [])
            if isinstance(p, dict)
        ],
        # include hidden tests only for evaluation endpoints, not in general fetch
    }

@app.post("/mocktest/session/create")
def mocktest_create_session():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]

    payload = request.get_json(silent=True) or {}
    mode = (payload.get("mode") or "all").strip().lower()
    difficulty = (payload.get("difficulty") or "mixed").strip().lower()
    pattern = payload.get("pattern") or {}
    gen_n = int(pattern.get("general") or 0)
    tech_n = int(pattern.get("tech") or 0)
    coding_n = int(pattern.get("coding") or 0)

    if mode == "general":
        tech_n = 0; coding_n = 0
    elif mode == "tech":
        gen_n = 0; coding_n = 0
    elif mode == "coding":
        gen_n = 0; tech_n = 0
    else:
        mode = "all" 

    if gen_n + tech_n + coding_n <= 0:
        return jsonify({"error": "Pattern must include at least one question/problem count."}), 400

    # Generate content
    try:
        general_qs = _mocktest_generate_apti_questions("general", gen_n, difficulty=difficulty) if gen_n else []
        tech_qs = _mocktest_generate_apti_questions("tech", tech_n, difficulty=difficulty) if tech_n else []
        coding_probs = _mocktest_generate_coding_problems(coding_n, difficulty=difficulty) if coding_n else []
    except Exception as e:
        # Gemini failure should not fall back silently; tell UI to regenerate.
        if is_quota_error(e):
            return jsonify({"error": "Quota exceeded for the Gemini API. Please try again later.", "code": "QUOTA_EXCEEDED"}), 429
        msg = str(e) or "Gemini generation failed. Please regenerate."
        return jsonify({"error": msg, "code": "GEMINI_GENERATION_FAILED"}), 503

    total_marks = len(general_qs) + len(tech_qs) + sum(int(p.get("total_marks") or 0) for p in coding_probs)

    doc = {
        "uid": user.get("uid"),
        "session_id": str(uuid.uuid4()),
        "title": (payload.get("title") or "Mock Test").strip() or "Mock Test",
        "mode": mode,
        "pattern": {"general": gen_n, "tech": tech_n, "coding": coding_n},
        "difficulty": difficulty,
        "status": "ready",
        "general_questions": general_qs,
        "tech_questions": tech_qs,
        "coding_problems": coding_probs,
        "answers": {"general": {}, "tech": {}, "coding": {}},
        "scores": {},
        "total_score": None,
        "total_marks": total_marks,
        "analysis": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow(),
        # camelCase mirrors (used by some admin screens / legacy code)
        "createdAt": datetime.utcnow(),
        "updatedAt": datetime.utcnow(),
    }

    col = _mocktest_sessions_col(get_db())
    col.insert_one(doc)
    return jsonify({"ok": True, "session": _mocktest_public_session(doc)})

@app.get("/mocktest/sessions")
def mocktest_list_sessions():
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    col = _mocktest_sessions_col(get_db())
    items = list(col.find({"uid": uid}).sort("created_at", -1).limit(50))
    # public summary list
    out = []
    for d in items:
        out.append({
            "session_id": str(d.get("session_id") or d.get("_id") or ""),
            "title": d.get("title") or "Mock Test",
            "mode": d.get("mode") or "all",
            "pattern": d.get("pattern") or {},
            "status": d.get("status") or "",
            "created_at": d.get("created_at"),
            "updated_at": d.get("updated_at"),
            "total_score": d.get("total_score"),
            "total_marks": d.get("total_marks"),
            "scores": d.get("scores") or {},
            "analysis": d.get("analysis") or None,
        })
    return jsonify({"ok": True, "sessions": out})

@app.get("/mocktest/sessions/<session_id>")
def mocktest_get_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    col = _mocktest_sessions_col(get_db())
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Not found"}), 404

    # Return full questions + coding metadata (no hidden tests)
    return jsonify({"ok": True, "session": _mocktest_public_session(doc), "answers": doc.get("answers") or {}})

@app.post("/mocktest/code/run")
def mocktest_code_run():
    """Run submitted code against up to 4 provided testcases (used for sample testcases in UI).

    Body: { language, source_code, tests: [{stdin, expected}...] }
    Returns: { ok, results:[{passed, stdout, stderr, compile_output, expected, status}] }
    """
    payload = request.get_json(silent=True) or {}
    t0 = time.time()
    language = (payload.get("language") or "python").strip().lower()
    source_code = payload.get("source_code") or ""
    tests = payload.get("tests") or []
    if not isinstance(tests, list) or not tests:
        return jsonify({"error": "No tests provided"}), 400

    t0 = time.time()

    results = []
    for t in tests[:4]:
        if not isinstance(t, dict):
            continue
        stdin = t.get("stdin") or ""
        expected = t.get("expected")
        try:
            r = _code_run(source_code, stdin, language, expected=expected)
        except Exception as e:
            r = {
                "stdout": "",
                "stderr": str(e),
                "compile_output": "",
                "status": "Error",
                "status_id": 13,
                "expected": (expected or "").strip() if expected is not None else "",
                "passed": False,
            }
        r["stdin"] = t.get("stdin","")
        results.append(r)

    wall_ms = int((time.time() - t0) * 1000)
    analysis = _gemini_code_analysis(payload.get("problem_text") or payload.get("problem") or "", source_code, language)
    # Build per-testcase details (all are public for Run)
    testcases = []
    passed_count = 0
    for i, r in enumerate(results):
        ok = bool(r.get("passed"))
        if ok:
            passed_count += 1
        testcases.append({
            "id": i + 1,
            "hidden": False,
            "passed": ok,
            "stdin": r.get("stdin", ""),
            "expected": r.get("expected", ""),
            "stdout": r.get("stdout", ""),
            "stderr": r.get("stderr", ""),
            "time_ms": r.get("time_ms"),
            "memory_kb": r.get("memory_kb"),
        })

    return jsonify({
        "ok": True,
        "results": results,   # backward compatible
        "testcases": testcases,
        "total_tests": len(testcases),
        "passed_tests": passed_count,
        "sample_total": len(testcases),
        "sample_passed": passed_count,
        "hidden_total": 0,
        "hidden_passed": 0,
        "hidden_failed": 0,
        "wall_time_ms": wall_ms,
        "timeComplexity": analysis.get("timeComplexity"),
        "spaceComplexity": analysis.get("spaceComplexity"),
        "analysisReason": analysis.get("reason"),
    }), 200


@app.post("/mocktest/sessions/<session_id>/coding/<pid>/submit")
def mocktest_submit_single_coding(session_id, pid):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]

    payload = request.get_json(silent=True) or {}
    language = (payload.get("language") or "python").strip().lower()
    source_code = payload.get("source_code") or ""

    if not source_code.strip():
        return jsonify({"error": "source_code is required"}), 400

    col = _mocktest_sessions_col(get_db())
    doc = col.find_one({"uid": user.get("uid"), "session_id": str(session_id)})
    if not doc:
        return jsonify({"error": "Session not found"}), 404

    if (doc.get("status") or "").lower() == "submitted":
        return jsonify({"error": "Session already submitted"}), 400

    prob = None
    for p in (doc.get("coding_problems") or []):
        if isinstance(p, dict) and str(p.get("id")) == str(pid):
            prob = p
            break

    if not prob:
        return jsonify({"error": "Problem not found"}), 404

    tests_samples = prob.get("samples") or []
    tests_hidden = prob.get("hiddenTests") or []

    def _norm_tests(arr):
        out = []
        for tt in (arr or []):
            if not isinstance(tt, dict):
                continue
            stdin = tt.get("stdin") if tt.get("stdin") is not None else tt.get("input")
            expected = tt.get("expected") if tt.get("expected") is not None else tt.get("output")
            if stdin is None or expected is None:
                continue
            out.append({"stdin": str(stdin), "expected": str(expected)})
        return out

    tests_samples_n = _norm_tests(tests_samples)
    tests_hidden_n = _norm_tests(tests_hidden)

    # Preserve whether a testcase is sample/hidden for frontend.
    tests = (
        [{**t, "hidden": False} for t in (tests_samples_n or [])]
        + [{**t, "hidden": True} for t in (tests_hidden_n or [])]
    )

    if not tests:
        return jsonify({"error": "No tests configured for this problem"}), 400

    t0 = time.time()
    results = []
    for i, t in enumerate(tests, start=1):
        stdin = t.get("stdin") or ""
        expected = t.get("expected")
        # Explicit flags set above
        hidden = bool(t.get("hidden", False))
        is_sample = not hidden

        try:
            r = _code_run(source_code, stdin, language, expected=expected)
        except Exception as e:
            r = {
                "stdout": "",
                "stderr": str(e),
                "compile_output": "",
                "status": "Error",
                "status_id": 13,
                "expected": (expected or "").strip() if expected is not None else "",
                "passed": False,
            }

        # Attach testcase metadata for frontend rendering
        r["id"] = i
        r["is_sample"] = is_sample
        r["hidden"] = hidden
        r["stdin"] = stdin
        r["expected"] = expected if expected is not None else ""
        results.append(r)

    passed = sum(1 for r in (results or []) if r.get("passed"))
    total = len(tests)
    sample_total = len(tests_samples_n)
    hidden_total = len(tests_hidden_n)
    sample_passed = sum(1 for r in (results or [])[:sample_total] if r.get("passed"))
    hidden_passed = sum(1 for r in (results or [])[sample_total:] if r.get("passed"))

    first_fail = next((r for r in (results or []) if not r.get("passed")), None)
    status = "Accepted" if passed == total else "Wrong Answer"
    if first_fail and (first_fail.get("compile_output") or "").strip():
        status = "Compilation Error"
    elif first_fail and (first_fail.get("stderr") or "").strip():
        status = "Runtime Error"

    # ✅ Marks logic (requested): Easy 15, Medium 20, Hard 25 per problem.
    full_marks = int(prob.get("total_marks") or 0)
    if full_marks <= 0:
        dd = (prob.get("difficulty") or "").strip().lower()
        if dd == "easy":
            full_marks = 15
        elif dd == "medium":
            full_marks = 20
        elif dd in ["hard", "difficult"]:
            full_marks = 25
        else:
            full_marks = 20

    marks_awarded = full_marks if (passed == total and total > 0) else int(round((full_marks * (passed / max(1, total))) if total else 0))

    # Build per-testcase details.
    # Hidden testcases: do NOT expose stdin/expected/stdout, only pass/fail (+stderr if any).
    testcases = []
    for c in (results or []):
        hide_io = False  # UI controls show/hide; always return full details
        testcases.append(_case_to_public_dict(c, hide_io=hide_io))

    # Best-effort complexity estimation (especially for fallback-bank problems
    # that may not include Big-O metadata).
    tc_analysis = None
    try:
        if not (prob.get("timeComplexity") or prob.get("time_complexity") or prob.get("spaceComplexity") or prob.get("space_complexity")):
            ptxt = (
                f"{prob.get('title') or ''}\n"
                f"{prob.get('statement') or ''}\n\n"
                f"Input Format:\n{prob.get('inputFormat') or ''}\n\n"
                f"Output Format:\n{prob.get('outputFormat') or ''}\n\n"
                f"Constraints:\n" + "\n".join([str(x) for x in (prob.get('constraints') or [])])
            )
            tc_analysis = _gemini_code_analysis(ptxt, source_code, language) or None
    except Exception:
        tc_analysis = None

    summary = {
        "passed": passed,
        "total": total,
        "passed_all": passed == total,
        "status": status,
        "sample_passed": sample_passed,
        "sample_total": sample_total,
        "hidden_passed": hidden_passed,
        "hidden_total": hidden_total,
        "marks_awarded": marks_awarded,
        "full_marks": full_marks,
        "at": datetime.utcnow().isoformat() + "Z",
        "wall_time_ms": int((time.time() - t0) * 1000),
        "timeComplexity": prob.get("timeComplexity") or prob.get("time_complexity") or (tc_analysis.get("timeComplexity") if isinstance(tc_analysis, dict) else None),
        "spaceComplexity": prob.get("spaceComplexity") or prob.get("space_complexity") or (tc_analysis.get("spaceComplexity") if isinstance(tc_analysis, dict) else None),
        "analysis": (tc_analysis.get("analysis") if isinstance(tc_analysis, dict) else None),
    }

    try:
        col.update_one(
            {"_id": doc.get("_id")},
            {
                "$set": {
                    f"answers.coding.{pid}.language": language,
                    f"answers.coding.{pid}.code": source_code,
                    f"answers.coding.{pid}.lastSubmit": summary,
                    "updated_at": datetime.utcnow(),
                }
            },
        )
    except Exception:
        pass


    # Update session-level score/marks so history shows correct totals.
    try:
        updated = col.find_one({"_id": doc.get("_id")}) or {}
        answers = updated.get("answers") or {}
        coding = (answers.get("coding") or {}) if isinstance(answers, dict) else {}
        c_score = 0
        c_total_marks = 0
        for _pid, c in (coding or {}).items():
            ls = (c or {}).get("lastSubmit") if isinstance(c, dict) else None
            if isinstance(ls, dict):
                c_score += int(ls.get("marks_awarded") or 0)
                c_total_marks += int(ls.get("full_marks") or 0)
        scores = updated.get("scores") or {}
        if not isinstance(scores, dict):
            scores = {}
        scores["coding"] = c_score

        # total_marks: for coding sessions use sum of per-problem marks; otherwise add aptitude MCQs too.
        mode2 = (updated.get("mode") or "").lower()
        base_q = int((updated.get("pattern") or {}).get("general") or 0) + int((updated.get("pattern") or {}).get("tech") or 0)
        total_marks = c_total_marks if mode2 == "coding" else base_q + c_total_marks
        total_score = (int(scores.get("general") or 0) + int(scores.get("tech") or 0) + int(scores.get("coding") or 0))

        col.update_one({"_id": updated.get("_id")}, {"$set": {"scores": scores, "total_score": total_score, "total_marks": total_marks, "updated_at": datetime.utcnow()}})
    except Exception:
        pass

    return jsonify({**summary, "testcases": testcases})


@app.post("/mocktest/sessions/<session_id>/submit")
def mocktest_submit(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    payload = request.get_json(silent=True) or {}
    answers = payload.get("answers") or {}
    coding_submissions = answers.get("coding") or {}
    gen_ans = answers.get("general") or {}
    tech_ans = answers.get("tech") or {}

    db = get_db()
    col = _mocktest_sessions_col(db)
    doc = col.find_one({"uid": uid, "session_id": session_id})
    if not doc:
        return jsonify({"error": "Not found"}), 404

    # Score aptitude
    gen_score = 0
    tech_score = 0
    gen_qs = doc.get("general_questions") or []
    tech_qs = doc.get("tech_questions") or []

    def _score_mcq(qs, ans_map):
        s = 0
        for q in qs:
            qid = q.get("id")
            try:
                sel = int(ans_map.get(qid)) if ans_map.get(qid) is not None else None
            except Exception:
                sel = None
            if sel is not None and sel == int(q.get("correctIndex") or 0):
                s += int(q.get("marks") or 1)
        return s

    gen_score = _score_mcq(gen_qs, gen_ans if isinstance(gen_ans, dict) else {})
    tech_score = _score_mcq(tech_qs, tech_ans if isinstance(tech_ans, dict) else {})

    # Score coding via Judge0
    # ✅ Marks logic: Easy 15, Medium 20, Hard 25 per coding problem.
    coding_score = 0
    coding_total_marks = 0
    coding_details = {}
    problems = doc.get("coding_problems") or []

    def _coding_full_marks(pdoc: dict) -> int:
        fm = int((pdoc or {}).get("total_marks") or 0)
        if fm > 0:
            return fm
        dd = ((pdoc or {}).get("difficulty") or "").strip().lower()
        if dd == "easy":
            return 15
        if dd == "medium":
            return 20
        if dd in ["hard", "difficult"]:
            return 25
        return 20

    for p in problems:
        if not isinstance(p, dict):
            continue
        pid = p.get("id")
        sub = (coding_submissions.get(pid) or {}) if isinstance(coding_submissions, dict) else {}
        # Fallback to DB-stored code/language (common when UI submits per-problem).
        db_sub = (((doc.get("answers") or {}).get("coding") or {}).get(pid) or {}) if isinstance(doc, dict) else {}
        lang = (sub.get("language") or db_sub.get("language") or "").strip().lower()
        code = (sub.get("code") or db_sub.get("code") or "")

        full_marks = _coding_full_marks(p)
        coding_total_marks += full_marks

        # Final submit should validate BOTH samples + hidden (Run => samples, Submit => samples + hidden)
        samples = p.get("samples") or []
        hidden = p.get("hiddenTests") or []
        tests = list(samples) + list(hidden)
        tests = [t for t in tests if isinstance(t, dict)]

        per = []
        passed = 0
        total = len(tests)

        # If final submit payload doesn't include code/lang, use cached lastSubmit from DB.
        if (not code or not lang) and isinstance(db_sub, dict):
            last_submit = db_sub.get("lastSubmit") or {}
            if isinstance(last_submit, dict):
                marks_awarded = int(last_submit.get("marks_awarded") or 0)
                passed_cached = int(last_submit.get("passed") or last_submit.get("passed_count") or 0)
                total_cached = int(last_submit.get("total") or total)
                coding_score += max(0, marks_awarded)
                coding_details[pid] = {
                    "passed": passed_cached,
                    "total": total_cached,
                    "marks_awarded": marks_awarded,
                    "full_marks": int(last_submit.get("full_marks") or full_marks),
                    "results": last_submit.get("results") or [],
                    "source": "cached",
                }
                continue

        if not code or not lang:
            coding_details[pid] = {
                "passed": 0,
                "total": total,
                "marks_awarded": 0,
                "full_marks": full_marks,
                "results": [],
            }
            continue

        for idx, t in enumerate(tests):
            try:
                jr = _code_run(code, t.get("stdin") or "", lang, expected=(t.get("expected") or ""))
            except Exception as e:
                jr = {"passed": False, "error": str(e)}
            item = {k: jr.get(k) for k in ["passed", "stdout", "expected", "stderr", "compile_output", "status", "time", "memory"] if k in jr}
            item["isSample"] = bool(idx < len(samples))
            per.append(item)
            if jr.get("passed"):
                passed += 1

        marks_awarded = full_marks if (passed == total and total > 0) else int(round((full_marks * (passed / max(1, total))) if total else 0))
        coding_score += marks_awarded
        coding_details[pid] = {
            "passed": passed,
            "total": total,
            "marks_awarded": marks_awarded,
            "full_marks": full_marks,
            "results": per,
        }

    total_score = gen_score + tech_score + coding_score
    # total marks: MCQs (1 each) + coding fixed marks per problem
    total_marks = len(gen_qs) + len(tech_qs) + coding_total_marks


    # Build topic-level performance to help Gemini give specific strong/weak topics
    def _topic_stats(qs, ans_map):
        stats = {}
        for q in qs:
            topic = (q.get("topic") or "general").strip()
            qid = q.get("id")
            try:
                sel = int(ans_map.get(qid)) if ans_map.get(qid) is not None else None
            except Exception:
                sel = None
            correct = (sel is not None and sel == int(q.get("correctIndex") or 0))
            s = stats.get(topic, {"correct": 0, "total": 0})
            s["total"] += 1
            if correct:
                s["correct"] += 1
            stats[topic] = s
        # derive strong/weak lists
        strong = []
        weak = []
        for t, s in stats.items():
            if s["total"] <= 0:
                continue
            acc = s["correct"] / max(1, s["total"])
            if s["total"] >= 2 and acc >= 0.7:
                strong.append({"topic": t, "acc": round(acc, 2), "total": s["total"]})
            if s["total"] >= 1 and acc < 0.5:
                weak.append({"topic": t, "acc": round(acc, 2), "total": s["total"]})
        return {"stats": stats, "strong": strong[:5], "weak": weak[:5]}

    gen_topics = _topic_stats(gen_qs, gen_ans if isinstance(gen_ans, dict) else {})
    tech_topics = _topic_stats(tech_qs, tech_ans if isinstance(tech_ans, dict) else {})

    # Coding topics: mark strong when >=60% tests passed for a problem topic
    coding_topic_summary = {}
    for p in problems:
        topic = (p.get("topic") or "coding").strip()
        pid = p.get("id")
        det = coding_details.get(pid) or {}
        total_t = int(det.get("total") or 0)
        pass_t = int(det.get("passed") or 0)
        s = coding_topic_summary.get(topic, {"passed": 0, "total": 0, "problems": 0})
        s["passed"] += pass_t
        s["total"] += total_t
        s["problems"] += 1
        coding_topic_summary[topic] = s

    # Gemini analysis for stronger/weaker sections
    analysis = None
    try:
        analysis_prompt = f"""
You are an interview coach analyzing a user's mock test performance.
Return STRICT JSON only:
{{
  "strong_sections": ["general|tech|coding"],
  "weak_sections": ["general|tech|coding"],
  "strong_topics": {{
    "general": ["..."],
    "tech": ["..."],
    "coding": ["..."]
  }},
  "weak_topics": {{
    "general": ["..."],
    "tech": ["..."],
    "coding": ["..."]
  }},
  "summary": "2-4 lines",
  "overall_feedback": "2-4 lines",
  "improve_knowledge": [
     {{"section":"general|tech|coding","topics":["...","..."],"action_plan":"short bullet plan","resources_suggestion":"short"}}
  ]
}}

User scores:
- General Aptitude: {gen_score}/{len(gen_qs)}
- Tech Aptitude: {tech_score}/{len(tech_qs)}
- Coding (testcases passed): {coding_score}/{sum(len((p.get("hiddenTests") or [])) for p in problems)}

Topic-level performance (use this to identify EXACT weak/strong topics; do NOT return generic lists):
General strong candidates: {json.dumps(gen_topics.get("strong"))}
General weak candidates: {json.dumps(gen_topics.get("weak"))}
Tech strong candidates: {json.dumps(tech_topics.get("strong"))}
Tech weak candidates: {json.dumps(tech_topics.get("weak"))}
Coding topic summary: {json.dumps(coding_topic_summary)}

Guidelines:
- Choose strong_topics / weak_topics based on the provided accuracy summaries.
- "Improve knowledge" should focus on weak topics first and be specific (what to study + practice type).
- Keep it concise and practical.
- Do NOT mention Gemini.
"""
        analysis = _gemini_json(analysis_prompt) or None
    except Exception:
        analysis = None

    update = {
        "$set": {
            "answers": {
                "general": gen_ans if isinstance(gen_ans, dict) else {},
                "tech": tech_ans if isinstance(tech_ans, dict) else {},
                "coding": coding_submissions if isinstance(coding_submissions, dict) else {},
            },
            "scores": {"general": gen_score, "tech": tech_score, "coding": coding_score},
            "coding_details": coding_details,
        "coding_total_marks": coding_total_marks,
            "total_score": total_score,
            "total_marks": total_marks,
            "analysis": analysis,
            "status": "submitted",
            "updated_at": datetime.utcnow(),
            "updatedAt": datetime.utcnow(),
            "submittedAt": datetime.utcnow(),
        }
    }
    col.update_one({"uid": uid, "session_id": session_id}, update)

    # Send report email (Mailgun when hosted, SMTP on localhost) if user's email is available.
    email_sent = False
    try:
        to_email = (user.get("email") or "").strip()
        if to_email:
            # ✅ Test Name + Difficulty (from session doc)
            test_name = (
                (doc.get("test_name") or "")
                or (doc.get("title") or "")
                or (doc.get("name") or "")
                or "Zenith Mock Test"
            ).strip()

            test_difficulty = (
                (doc.get("difficulty") or "")
                or (doc.get("level") or "")
                or (doc.get("selected_difficulty") or "")
            ).strip()

            test_name_display = f"{test_name} ({test_difficulty.title()})" if test_difficulty else test_name

            # ✅ Date formatting (IST)
            from datetime import timezone, timedelta

            ist = timezone(timedelta(hours=5, minutes=30))

            created_at = doc.get("created_at")
            if isinstance(created_at, datetime):
                # If stored as naive UTC datetime, treat as UTC
                if created_at.tzinfo is None:
                    created_at = created_at.replace(tzinfo=timezone.utc)
                created_str = created_at.astimezone(ist).strftime("%Y-%m-%d %I:%M %p IST")
            else:
                created_str = "N/A"

            submitted_dt = datetime.utcnow().replace(tzinfo=timezone.utc).astimezone(ist)
            submitted_str = submitted_dt.strftime("%Y-%m-%d %I:%M %p IST")

            subj = f"{test_name_display} • Report • Score {total_score}/{total_marks}"

            # Build section-wise score list ONLY for sections present in this test.
            score_items = [f"<li><b>🏆 Total Score:</b> {total_score}/{total_marks}</li>"]
            if len(gen_qs) > 0:
                score_items.append(f"<li><b>🧠 General Aptitude:</b> {gen_score}/{len(gen_qs)}</li>")
            if len(tech_qs) > 0:
                score_items.append(f"<li><b>💻 Technical Aptitude:</b> {tech_score}/{len(tech_qs)}</li>")
            if (coding_total_marks or 0) > 0:
                score_items.append(f"<li><b>👨‍💻 Coding:</b> {coding_score}/{coding_total_marks}</li>")

            # Gemini/AI report formatting with emojis (kept concise + practical).
            gemini_block = ""
            summ = ""
            if isinstance(analysis, dict):
                summ = analysis.get("summary")
                overall_fb = analysis.get("overall_feedback")
                strong_secs = analysis.get("strong_sections") or []
                weak_secs = analysis.get("weak_sections") or []
                strong_topics = analysis.get("strong_topics") or {}
                weak_topics = analysis.get("weak_topics") or {}
                improve = analysis.get("improve_knowledge") or []

                def _sec_label(s):
                    s = (s or "").strip().lower()
                    return {"general": "General Aptitude", "tech": "Technical Aptitude", "coding": "Coding"}.get(
                        s, s.title() or "Section"
                    )

                parts = []

                # ✅ Add test meta at top
                parts.append(f"<p><b>📝 Test Name:</b> {_html_escape(test_name_display)}</p>")
                parts.append(f"<p><b>🗓️ Test Created:</b> {_html_escape(created_str)}<br/>"
                            f"<b>⏰ Submitted:</b> {_html_escape(submitted_str)}</p>")
                parts.append("<hr/>")

                if summ:
                    parts.append(f"<p><b>🧾 Summary:</b> {_html_escape(str(summ))}</p>")
                if overall_fb:
                    parts.append(f"<p><b>💬 Feedback:</b> {_html_escape(str(overall_fb))}</p>")

                # Strong/weak sections
                if strong_secs:
                    parts.append(
                        f"<p><b>✅ Strong:</b> "
                        + ", ".join([_html_escape(_sec_label(x)) for x in strong_secs])
                        + "</p>"
                    )
                if weak_secs:
                    parts.append(
                        f"<p><b>⚠️ Needs work:</b> "
                        + ", ".join([_html_escape(_sec_label(x)) for x in weak_secs])
                        + "</p>"
                    )

                # Topics (only show for sections included in this test)
                def _topics_html(title, tmap):
                    blocks = []
                    for sec_key in ["general", "tech", "coding"]:
                        if sec_key == "general" and len(gen_qs) == 0:
                            continue
                        if sec_key == "tech" and len(tech_qs) == 0:
                            continue
                        if sec_key == "coding" and int(coding_total_marks or 0) <= 0:
                            continue
                        arr = tmap.get(sec_key) or []
                        arr = [str(x) for x in arr if str(x).strip()]
                        if not arr:
                            continue
                        blocks.append(
                            f"<li><b>{_html_escape(_sec_label(sec_key))}:</b> "
                            + ", ".join([_html_escape(x) for x in arr[:8]])
                            + "</li>"
                        )
                    if not blocks:
                        return ""
                    return f"<p><b>{title}</b></p><ul>" + "".join(blocks) + "</ul>"

                parts.append(_topics_html("💪 Strong Topics", strong_topics) or "")
                parts.append(_topics_html("🛠️ Focus Topics", weak_topics) or "")

                # Action plan
                if isinstance(improve, list) and improve:
                    plan_items = []
                    for item in improve[:4]:
                        if not isinstance(item, dict):
                            continue
                        sec = _sec_label(item.get("section"))
                        topics = item.get("topics") or []
                        topics = [str(x) for x in topics if str(x).strip()]
                        ap = (item.get("action_plan") or "").strip()
                        res = (item.get("resources_suggestion") or "").strip()
                        li = f"<li><b>🎯 {_html_escape(sec)}</b>"
                        if topics:
                            li += "<br/><span>📌 Topics: " + ", ".join([_html_escape(x) for x in topics[:8]]) + "</span>"
                        if ap:
                            li += "<br/><span>✅ Plan: " + _html_escape(ap) + "</span>"
                        if res:
                            li += "<br/><span>📚 Resources: " + _html_escape(res) + "</span>"
                        li += "</li>"
                        plan_items.append(li)
                    if plan_items:
                        parts.append("<p><b>📈 Improvement Plan</b></p><ul>" + "".join(plan_items) + "</ul>")

                gemini_block = "".join([p for p in parts if p])

            body_html = (
                "<p>Your mock test has been evaluated and your report is ready ✅</p>"
                + "<ul>" + "".join(score_items) + "</ul>"
                + (gemini_block if gemini_block else (
                    f"<p><b>📝 Test Name:</b> {_html_escape(test_name_display)}</p>"
                    f"<p><b>🗓️ Test Created:</b> {_html_escape(created_str)}<br/>"
                    f"<b>⏰ Submitted:</b> {_html_escape(submitted_str)}</p><hr/>"
                ))
                + "<p style='margin-top:12px'>Open Zenith to review your detailed analysis, mistakes, and recommendations.</p>"
            )

            html = _brand_email(
                title="Zenith Mock Test Report",
                subtitle=f"Score: {total_score}/{total_marks}",
                body_html=body_html,
                cta_url="/my-tests",
                cta_text="View Mock Test History",
            )

            email_sent = bool(
                _send_email(
                    to_email,
                    subj,
                    html,
                    text_body=(
                        f"Test Name: {test_name_display}\n"
                        f"Created: {created_str}\n"
                        f"Submitted: {submitted_str}\n\n"
                        f"Total: {total_score}/{total_marks}\n"
                        + (f"General: {gen_score}/{len(gen_qs)}\n" if len(gen_qs) > 0 else "")
                        + (f"Technical: {tech_score}/{len(tech_qs)}\n" if len(tech_qs) > 0 else "")
                        + (f"Coding: {coding_score}/{coding_total_marks}\n" if (coding_total_marks or 0) > 0 else "")
                        + (f"\nSummary: {summ}\n" if summ else "")
                    ),
                    kind="mocktest",
                    reply_to=(os.getenv("CONTACT_INBOX") or os.getenv("ADMIN_EMAIL")),
                )
            )
    except Exception as e:
        # Don't crash submission if email fails; log for debugging.
        try:
            import traceback
            print("[email] mock test report send failed:", repr(e))
            traceback.print_exc()
        except Exception:
            pass

    return jsonify({
        "ok": True,
        "email_sent": email_sent,
        "scores": {"general": gen_score, "tech": tech_score, "coding": coding_score},
        "total_score": total_score,
        "total_marks": total_marks,
        "analysis": analysis,
        "coding_details": coding_details,
        "coding_total_marks": coding_total_marks,
    })

@app.delete("/mocktest/sessions/<session_id>")
def mocktest_delete_session(session_id):
    user, err = require_user()
    if err:
        return jsonify({"error": err[0]}), err[1]
    uid = user.get("uid")
    col = _mocktest_sessions_col(get_db())
    res = col.delete_one({"uid": uid, "session_id": session_id})
    return jsonify({"ok": True, "deleted": int(res.deleted_count or 0)})

def _case_to_public_dict(case, hide_io=False):
    """Convert internal case result to API testcase dict.
    hide_io=True -> do not expose stdin/expected/stdout (used for hidden testcases).
    """
    d = {
        "id": case.get("id"),
        "hidden": bool(case.get("hidden", False)),
        "passed": bool(case.get("passed", False)),
        "time_ms": case.get("time_ms"),
        "memory_kb": case.get("memory_kb"),
    }
    if hide_io:
        d["stdin"] = ""
        d["expected"] = ""
        d["stdout"] = ""
        d["stderr"] = (case.get("stderr") or "")
    else:
        d["stdin"] = (case.get("stdin") or "")
        d["expected"] = (case.get("expected") or "")
        d["stdout"] = (case.get("stdout") or "")
        d["stderr"] = (case.get("stderr") or "")
    return d
if __name__ == "__main__":
    app.run(debug=True)